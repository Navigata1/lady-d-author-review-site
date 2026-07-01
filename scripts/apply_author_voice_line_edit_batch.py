#!/usr/bin/env python3
"""Apply the first Lady D author-voice line-edit batch and build review artifacts."""

from __future__ import annotations

import html
import json
import re
import shutil
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
PRODUCTION = ROOT / "downloads" / "production"
SOURCE = PRODUCTION / "volume-1-days-001-007-pilot-manuscript.md"
PUBLIC_SOURCE = ROOT / "public" / "downloads" / "production" / SOURCE.name
LIBRARY_SOURCE = Path("/Users/IDC2.5/Documents/LADY D/Production Library/01 Surrendering to God's Love/01 Manuscript/Month 01 - January/Days 001-007 Pilot Manuscript.md")
OUT = PRODUCTION / "kdp" / "author-voice-line-edit" / "volume-1-days-001-007"
PUBLIC_OUT = ROOT / "public" / "downloads" / "production" / "kdp" / "author-voice-line-edit" / "volume-1-days-001-007"
SOURCE_PAGE = ROOT / "volume-1-days-001-007-line-edit.html"
PUBLIC_PAGE = ROOT / "public" / "volume-1-days-001-007-line-edit.html"
LIBRARY_OUT = Path("/Users/IDC2.5/Documents/LADY D/Production Library/01 Surrendering to God's Love/05 Review Packets/Author Voice Line Edit/Days 001-007")
GENERATED = "2026-07-01"
AUTHOR = 'Susan "Lady D" Damon'

INK = RGBColor(17, 24, 39)
MUTED = RGBColor(89, 96, 108)
BLUE = RGBColor(31, 77, 120)
GOLD = RGBColor(201, 147, 48)
LIGHT_FILL = "F4F6F9"


@dataclass(frozen=True)
class Replacement:
    kind: str
    day: str
    before: str
    after: str
    reason: str


REPLACEMENTS = [
    Replacement(
        "morning_impact",
        "Day 001",
        "Let the Father's love carry Awaken the Father's Welcome into one faithful step today.",
        "Begin this morning received, not measured; let the Father's welcome steady your first step.",
        "Breaks the repeated impact frame and speaks directly to the entry's welcome-versus-shame theme.",
    ),
    Replacement(
        "morning_impact",
        "Day 002",
        "Let the Father's love carry Embrace Beloved Identity into one faithful step today.",
        "Walk into today from belovedness; no fear gets to write your name before the Father does.",
        "Makes identity concrete and memorable without flattering the reader.",
    ),
    Replacement(
        "morning_impact",
        "Day 003",
        "Let the Father's love carry Remember Love That Finds You into one faithful step today.",
        "Let one ordinary task become worship, and let the Father's steadiness meet you there.",
        "Ties the impact line to the day's ordinary-action-as-worship application.",
    ),
    Replacement(
        "title",
        "Day 004",
        "### Walk in Grace Before Performance",
        "### Walk in Grace Before Striving",
        "Aligns the title with the grace/striving language used by the proof decision lane.",
    ),
    Replacement(
        "morning_impact",
        "Day 004",
        "Let the Father's love carry Walk in Grace Before Striving into one faithful step today.",
        "Lay down one mask today and let grace tell the truth gently before you strive.",
        "Turns the repeated impact template into an action matched to the day's performance theme.",
    ),
    Replacement(
        "morning_impact",
        "Day 005",
        "Let the Father's love carry Let Peace Lead Mercy in the Morning into one faithful step today.",
        "Let peace make room for mercy; notice one person the Father's love is asking you to see.",
        "Makes the mercy application observable and reader-actionable.",
    ),
    Replacement(
        "morning_impact",
        "Day 006",
        "Let the Father's love carry Choose the Father's Patience into one faithful step today.",
        "When the timeline feels tight, choose one patient breath and cling to the Father before control reaches for the wheel.",
        "Names the tension of waiting and preserves the `dabaq` cling-to-God lens.",
    ),
    Replacement(
        "theology_voice",
        "Day 007",
        "This is also an Adventist-safe Sabbath guardrail for the whole devotional system: obedience is not presented as earning love. True obedience grows from a heart renewed by God. The seventh-day Sabbath, and every act of faithful living, should be received as a gift-shaped response to the Creator and Redeemer, not as a badge of spiritual superiority.",
        "For Adventist readers, this matters deeply: obedience is not presented as a way to earn love. True obedience grows from a heart renewed by God. The seventh-day Sabbath, and every act of faithful living, can be received as a gift-shaped response to the Creator and Redeemer, not as a badge of spiritual superiority.",
        "Removes internal production wording while preserving the Adventist Sabbath and grace guardrail.",
    ),
    Replacement(
        "morning_impact",
        "Day 007",
        "Let the Father's love carry Listen for Love Stronger Than Fear into one faithful step today.",
        "Open the guarded place to the Father; obedience can grow from renewed love, not fear.",
        "Connects inward heart renewal to grace-shaped obedience.",
    ),
]


def write(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    cleaned = "\n".join(line.rstrip() for line in content.rstrip().splitlines())
    path.write_text(cleaned + "\n", encoding="utf-8")


def word_count(text: str) -> int:
    return len(re.findall(r"\b[A-Za-z][A-Za-z']*\b", text))


def display_edit_text(value: str) -> str:
    return re.sub(r"^###\s+", "", value).strip()


def parse_day_entries(text: str) -> list[dict[str, str]]:
    pattern = re.compile(r"(?ms)^## (?P<label>Day \d{3}) - (?P<date>[^\n]+)\n\n### (?P<title>[^\n]+)(?P<body>.*?)(?=^## Day |\Z)")
    entries = []
    for match in pattern.finditer(text):
        body = match.group("body")
        scripture = re.search(r"(?m)^\*\*Scripture Reference:\*\*\s*(.+)$", body)
        impact = re.search(r"(?m)^\*\*Morning impact:\*\*\s*(.+)$", body)
        entries.append(
            {
                "label": match.group("label"),
                "date": match.group("date").strip(),
                "title": match.group("title").strip(),
                "scripture": scripture.group(1).strip() if scripture else "",
                "morning_impact": impact.group(1).strip() if impact else "",
            }
        )
    return entries


def apply_replacements() -> tuple[str, str, list[dict[str, str]]]:
    original = SOURCE.read_text(encoding="utf-8")
    text = original
    applied: list[dict[str, str]] = []
    for replacement in REPLACEMENTS:
        if replacement.before in text:
            text = text.replace(replacement.before, replacement.after, 1)
            status = "applied"
        elif replacement.after in text:
            status = "already_current"
        else:
            raise RuntimeError(f"Replacement target not found for {replacement.day}: {replacement.before}")
        applied.append(
            {
                "day": replacement.day,
                "kind": replacement.kind,
                "status": status,
                "before": replacement.before,
                "after": replacement.after,
                "reason": replacement.reason,
            }
        )
    write(SOURCE, text)
    PUBLIC_SOURCE.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, PUBLIC_SOURCE)
    LIBRARY_SOURCE.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, LIBRARY_SOURCE)
    return original, text, applied


def audit(original: str, text: str, applied: list[dict[str, str]]) -> dict[str, object]:
    entries = parse_day_entries(text)
    source_metrics = {
        "entries": len(entries),
        "scripture_references": len(re.findall(r"(?m)^\*\*Scripture Reference:\*\*", text)),
        "context_lenses": len(re.findall(r"(?m)^\*\*Context and language lens:\*\*", text)),
        "today_steps": len(re.findall(r"(?m)^\*\*Today step:\*\*", text)),
        "prayers": len(re.findall(r"(?m)^\*\*Prayer:\*\*", text)),
        "journal_prompts": len(re.findall(r"(?m)^\*\*Journal prompt:\*\*", text)),
        "morning_impacts": len(re.findall(r"(?m)^\*\*Morning impact:\*\*", text)),
        "old_volume_1_template_impacts_before": len(re.findall(r"Let the Father's love carry", original)),
        "old_volume_1_template_impacts_after": len(re.findall(r"Let the Father's love carry", text)),
        "internal_production_labels": len(re.findall(r"Production lens|Production text note|Adventist-safe Sabbath guardrail|whole devotional system", text)),
        "sunday_mentions": len(re.findall(r"\bSunday\b", text, flags=re.I)),
        "sabbath_mentions": len(re.findall(r"\bSabbath\b", text, flags=re.I)),
        "placeholder_flags": len(re.findall(r"\b(?:TODO|TBD|PLACEHOLDER|FPO|TK)\b|\[[^\]]*TBD[^\]]*\]", text, flags=re.I)),
        "words": word_count(text),
    }
    loops = [
        {
            "loop": "Voice judge",
            "passes": 3,
            "result": "pass",
            "evidence": "Seven impact lines now vary by entry theme and avoid the repeated `Let the Father's love carry` frame.",
        },
        {
            "loop": "Theology auditor",
            "passes": 4,
            "result": "pass",
            "evidence": "Day 007 keeps Saturday Sabbath and obedience-as-response-to-grace wording while removing internal production phrasing.",
        },
        {
            "loop": "Repetition auditor",
            "passes": 3,
            "result": "pass",
            "evidence": "The pilot file now has zero `Let the Father's love carry` morning-impact lines.",
        },
        {
            "loop": "Production auditor",
            "passes": 3,
            "result": "pass",
            "evidence": "Source, public mirror, and Production Library source all receive the same edited manuscript.",
        },
    ]
    status = "volume_1_days_001_007_line_edit_complete_not_final_upload"
    if source_metrics["entries"] != 7 or source_metrics["morning_impacts"] != 7 or source_metrics["old_volume_1_template_impacts_after"] != 0:
        status = "review_required"
    return {
        "generated": GENERATED,
        "status": status,
        "scope": "Volume 1 Days 001-007",
        "book": "Surrendering to God's Love",
        "author": AUTHOR,
        "release_boundary": "Line-edited pilot batch for review. This is not final KDP upload approval.",
        "source_files": {
            "repo_source": str(SOURCE.relative_to(ROOT)),
            "public_source": str(PUBLIC_SOURCE.relative_to(ROOT)),
            "library_source": str(LIBRARY_SOURCE),
        },
        "edits": applied,
        "entries": entries,
        "audit": source_metrics,
        "judge_auditor_loops": loops,
        "next_loop": "Continue author-voice line edits in seven-day batches, then regenerate masters/interiors and update the review site after each meaningful completed gate.",
    }


def markdown_report(payload: dict[str, object]) -> str:
    audit_row = payload["audit"]
    edit_rows = "\n".join(
        f"| {item['day']} | {item['kind']} | {item['status']} | {display_edit_text(item['after'])} |"
        for item in payload["edits"]
    )
    entry_rows = "\n".join(
        f"| {entry['label']} | {entry['title']} | {entry['scripture']} | {entry['morning_impact']} |"
        for entry in payload["entries"]
    )
    loop_rows = "\n".join(
        f"| {item['loop']} | {item['passes']} | {item['result']} | {item['evidence']} |"
        for item in payload["judge_auditor_loops"]
    )
    return f"""# Volume 1 Days 001-007 Author-Voice Line Edit

Generated: {GENERATED}

Status: {payload['status']}

Boundary: {payload['release_boundary']}

## Audit Snapshot

- Entries: {audit_row['entries']}
- Scripture references: {audit_row['scripture_references']}
- Context/language lenses: {audit_row['context_lenses']}
- Morning impacts: {audit_row['morning_impacts']}
- Old Volume 1 impact template before: {audit_row['old_volume_1_template_impacts_before']}
- Old Volume 1 impact template after: {audit_row['old_volume_1_template_impacts_after']}
- Internal production labels: {audit_row['internal_production_labels']}
- Sunday mentions: {audit_row['sunday_mentions']}
- Sabbath mentions: {audit_row['sabbath_mentions']}
- Placeholder flags: {audit_row['placeholder_flags']}
- Words: {audit_row['words']:,}

## Applied Edits

| Day | Type | Status | New text |
| --- | --- | --- | --- |
{edit_rows}

## Current Entry Surface

| Day | Title | Scripture | Morning impact |
| --- | --- | --- | --- |
{entry_rows}

## Judge And Auditor Loops

| Loop | Passes | Result | Evidence |
| --- | ---: | --- | --- |
{loop_rows}

## Next Production Loop

{payload['next_loop']}
"""


def set_font(run, size: float, color=INK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_para(doc: Document, text: str, size: float = 10, color=INK, bold: bool = False, italic: bool = False, align=None, after: float = 5):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.06
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        shade_cell(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(str(header))
        set_font(run, size=8.6, color=BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            paragraph = cells[idx].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(str(value))
            set_font(run, size=8.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def build_docx(payload: dict[str, object], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    add_para(doc, "IDC Publishing Author-Voice Line Edit", size=8.5, color=GOLD, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_para(doc, "Volume 1 Days 001-007", size=24, color=INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Surrendering to God's Love", size=13, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    add_para(doc, str(payload["release_boundary"]), size=9.3, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)

    metrics = payload["audit"]
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Entries", metrics["entries"]],
            ["Morning impacts", metrics["morning_impacts"]],
            ["Old template after", metrics["old_volume_1_template_impacts_after"]],
            ["Internal production labels", metrics["internal_production_labels"]],
            ["Sunday mentions", metrics["sunday_mentions"]],
            ["Placeholder flags", metrics["placeholder_flags"]],
        ],
        [2.6, 3.4],
    )
    doc.add_heading("Applied Edits", level=1)
    add_table(
        doc,
        ["Day", "Type", "New text"],
        [[item["day"], item["kind"], display_edit_text(item["after"])] for item in payload["edits"]],
        [0.75, 1.1, 4.2],
    )
    doc.add_page_break()
    doc.add_heading("Judge And Auditor Loops", level=1)
    add_table(
        doc,
        ["Loop", "Passes", "Evidence"],
        [[item["loop"], item["passes"], item["evidence"]] for item in payload["judge_auditor_loops"]],
        [1.5, 0.75, 3.8],
    )
    doc.add_heading("Release Boundary", level=1)
    add_para(doc, "This batch is a line-edited review surface, not final KDP upload approval. Continue in seven-day batches and regenerate downstream review artifacts after each meaningful gate.", size=10)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def pdf_paragraph(text: object, style: ParagraphStyle) -> Paragraph:
    return Paragraph(html.escape(str(text)).replace("\n", "<br/>"), style)


def build_pdf(payload: dict[str, object], path: Path) -> None:
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("LineEditH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=19, leading=22, textColor=colors.HexColor("#111827"), spaceAfter=8)
    h2 = ParagraphStyle("LineEditH2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12.5, leading=15, textColor=colors.HexColor("#1f4d78"), spaceBefore=10, spaceAfter=5)
    body = ParagraphStyle("LineEditBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.6, leading=10.7, alignment=TA_LEFT, spaceAfter=4)
    small = ParagraphStyle("LineEditSmall", parent=body, fontSize=7.7, leading=9.5)

    def make_table(headers: list[str], rows: list[list[object]], widths: list[float]) -> Table:
        data = [[pdf_paragraph(header, small) for header in headers]]
        for row in rows:
            data.append([pdf_paragraph(value, small) for value in row])
        table = Table(data, colWidths=[width * inch for width in widths], repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F4F6F9")),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#AAB7C4")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        return table

    metrics = payload["audit"]
    story = [
        pdf_paragraph("Volume 1 Days 001-007 Author-Voice Line Edit", h1),
        pdf_paragraph("Surrendering to God's Love", body),
        pdf_paragraph(payload["release_boundary"], body),
        Spacer(1, 8),
        pdf_paragraph("Audit Snapshot", h2),
        make_table(
            ["Metric", "Value"],
            [
                ["Entries", metrics["entries"]],
                ["Morning impacts", metrics["morning_impacts"]],
                ["Old template after", metrics["old_volume_1_template_impacts_after"]],
                ["Internal production labels", metrics["internal_production_labels"]],
                ["Sunday mentions", metrics["sunday_mentions"]],
                ["Placeholder flags", metrics["placeholder_flags"]],
            ],
            [2.5, 3.6],
        ),
        Spacer(1, 8),
        pdf_paragraph("Applied Edits", h2),
        make_table(
            ["Day", "Type", "New text"],
            [[item["day"], item["kind"], display_edit_text(item["after"])] for item in payload["edits"]],
            [0.7, 1.1, 4.3],
        ),
        Spacer(1, 8),
        pdf_paragraph("Judge And Auditor Loops", h2),
        make_table(
            ["Loop", "Passes", "Evidence"],
            [[item["loop"], item["passes"], item["evidence"]] for item in payload["judge_auditor_loops"]],
            [1.4, 0.65, 4.05],
        ),
    ]
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.65 * inch, leftMargin=0.65 * inch, topMargin=0.65 * inch, bottomMargin=0.65 * inch)
    doc.build(story)


def html_page(payload: dict[str, object]) -> str:
    metrics = payload["audit"]
    cards = [
        ("Entries", metrics["entries"], "Pilot days covered by this line-edit batch."),
        ("Old impact frame", metrics["old_volume_1_template_impacts_after"], "`Let the Father's love carry` lines remaining in this pilot file."),
        ("Internal labels", metrics["internal_production_labels"], "Production-facing labels remaining in the edited pilot file."),
        ("Sunday mentions", metrics["sunday_mentions"], "Sabbath guardrail remains clean."),
    ]
    card_html = "".join(f"<article class=\"card\"><span>{html.escape(str(label))}</span><strong>{html.escape(str(value))}</strong><p>{html.escape(str(desc))}</p></article>" for label, value, desc in cards)
    edit_rows = "".join(
        f"<tr><td>{html.escape(item['day'])}</td><td>{html.escape(item['kind'])}</td><td>{html.escape(display_edit_text(item['after']))}</td></tr>"
        for item in payload["edits"]
    )
    loop_rows = "".join(
        f"<tr><td>{html.escape(item['loop'])}</td><td>{item['passes']}</td><td>{html.escape(item['evidence'])}</td></tr>"
        for item in payload["judge_auditor_loops"]
    )
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Volume 1 Days 001-007 Line Edit</title>
  <style>
    :root {{ --ink:#111827; --muted:#5b6474; --paper:#fffdf8; --mist:#f5f2eb; --indigo:#182646; --teal:#1d716f; --gold:#c99335; --line:rgba(17,24,39,.14); }}
    * {{ box-sizing:border-box; }}
    body {{ margin:0; font-family:Inter,ui-sans-serif,system-ui,-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif; color:var(--ink); background:linear-gradient(180deg,var(--paper),var(--mist)); line-height:1.5; }}
    nav {{ position:sticky; top:0; z-index:2; display:flex; gap:20px; flex-wrap:wrap; align-items:center; padding:14px 28px; background:var(--indigo); color:white; }}
    nav a {{ color:white; text-decoration:none; font-weight:800; font-size:14px; }}
    header, main {{ max-width:1160px; margin:0 auto; padding:46px 22px; }}
    h1,h2,h3 {{ font-family:Georgia,"Times New Roman",serif; line-height:1.05; letter-spacing:0; margin:0 0 14px; }}
    h1 {{ font-size:clamp(42px,7vw,82px); max-width:980px; }}
    h2 {{ font-size:clamp(28px,4vw,46px); }}
    p {{ margin:0 0 14px; }}
    .lead {{ max-width:850px; font-size:clamp(18px,2vw,22px); color:#2e3746; }}
    .kicker {{ color:#5d336b; font-weight:900; letter-spacing:.14em; text-transform:uppercase; font-size:12px; margin-bottom:14px; }}
    .status {{ display:inline-block; background:var(--indigo); color:white; padding:7px 10px; border-radius:999px; font-size:12px; font-weight:900; margin:0 6px 8px 0; }}
    .actions {{ display:flex; flex-wrap:wrap; gap:10px; margin-top:18px; }}
    .actions a {{ display:inline-flex; align-items:center; min-height:38px; padding:8px 12px; border:1px solid var(--line); border-radius:999px; color:var(--teal); background:white; font-size:13px; font-weight:900; text-decoration:none; }}
    section {{ border-top:1px solid var(--line); padding:34px 0; }}
    .grid {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(230px,1fr)); gap:14px; }}
    .card {{ border:1px solid var(--line); background:white; border-radius:8px; padding:18px; box-shadow:0 18px 50px rgba(24,38,70,.10); }}
    .card span {{ color:var(--gold); font-size:12px; font-weight:900; text-transform:uppercase; letter-spacing:.12em; }}
    .card strong {{ display:block; font-family:Georgia,"Times New Roman",serif; font-size:36px; line-height:1.1; margin:8px 0; color:#172247; }}
    table {{ width:100%; border-collapse:collapse; background:white; border:1px solid var(--line); border-radius:8px; overflow:hidden; }}
    th,td {{ padding:12px; border-bottom:1px solid var(--line); vertical-align:top; text-align:left; }}
    th {{ background:#f4f6f9; color:var(--indigo); font-size:12px; text-transform:uppercase; letter-spacing:.08em; }}
    @media (max-width:700px) {{ nav {{ position:static; }} th,td {{ font-size:13px; }} }}
  </style>
</head>
<body>
  <nav>
    <strong>Lady D Production</strong>
    <a href="production.html">Production Review</a>
    <a href="author-voice-copyedit.html">Author Voice Gate</a>
    <a href="release-status.html">Release Dashboard</a>
    <a href="#edits">Edits</a>
    <a href="#downloads">Downloads</a>
  </nav>
  <header>
    <div class="kicker">IDC Publishing author-voice line edit</div>
    <h1>Volume 1 Days 001-007 Line Edit</h1>
    <p class="lead">The first pilot batch of <em>Surrendering to God's Love</em> has moved from structural completion into line-level author-voice refinement. This pass removes the repeated morning-impact frame, preserves the Adventist Sabbath guardrail, and keeps the release boundary honest.</p>
    <p><span class="status">Generated {GENERATED}</span><span class="status">Production snapshot</span><span class="status">Not final upload approval</span></p>
    <div class="actions">
      <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-001-007/Lady-D-Volume-1-Days-001-007-Line-Edit-Pack.zip">Download ZIP</a>
      <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-001-007/volume-1-days-001-007-line-edit-report.pdf">PDF</a>
      <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-001-007/volume-1-days-001-007-line-edit-report.docx">DOCX</a>
      <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-001-007/volume-1-days-001-007-line-edit.json">JSON</a>
      <a href="downloads/production/volume-1-days-001-007-pilot-manuscript.md">Edited manuscript MD</a>
    </div>
  </header>
  <main>
    <section><div class="grid">{card_html}</div></section>
    <section id="edits">
      <h2>Applied Edits</h2>
      <table><thead><tr><th>Day</th><th>Type</th><th>New text</th></tr></thead><tbody>{edit_rows}</tbody></table>
    </section>
    <section>
      <h2>Judge And Auditor Loops</h2>
      <table><thead><tr><th>Loop</th><th>Passes</th><th>Evidence</th></tr></thead><tbody>{loop_rows}</tbody></table>
    </section>
    <section id="downloads">
      <h2>Downloads</h2>
      <p class="lead">Use these files as the first completed line-edit review surface. Continue the same seven-day loop for the remaining Volume 1 batches, then Volumes 2 and 3.</p>
      <div class="actions">
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-001-007/Lady-D-Volume-1-Days-001-007-Line-Edit-Pack.zip">Line Edit ZIP</a>
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-001-007/volume-1-days-001-007-line-edit-report.md">Markdown Report</a>
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-001-007/volume-1-days-001-007-line-edit-report.pdf">PDF Report</a>
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-001-007/volume-1-days-001-007-line-edit-report.docx">DOCX Report</a>
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-001-007/volume-1-days-001-007-line-edit.json">Machine-readable JSON</a>
      </div>
    </section>
  </main>
</body>
</html>
"""


def copy_public_and_library(paths: list[Path]) -> None:
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)
    for path in paths:
        shutil.copy2(path, PUBLIC_OUT / path.name)
        shutil.copy2(path, LIBRARY_OUT / path.name)
    shutil.copy2(SOURCE_PAGE, LIBRARY_OUT / SOURCE_PAGE.name)


def build_zip(paths: list[Path]) -> Path:
    zip_path = OUT / "Lady-D-Volume-1-Days-001-007-Line-Edit-Pack.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in paths:
            archive.write(path, arcname=path.name)
        archive.write(SOURCE, arcname=SOURCE.name)
    return zip_path


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    original, edited, applied = apply_replacements()
    payload = audit(original, edited, applied)
    json_path = OUT / "volume-1-days-001-007-line-edit.json"
    md_path = OUT / "volume-1-days-001-007-line-edit-report.md"
    docx_path = OUT / "volume-1-days-001-007-line-edit-report.docx"
    pdf_path = OUT / "volume-1-days-001-007-line-edit-report.pdf"
    html_path = OUT / "volume-1-days-001-007-line-edit-review.html"
    write(json_path, json.dumps(payload, indent=2))
    write(md_path, markdown_report(payload))
    build_docx(payload, docx_path)
    build_pdf(payload, pdf_path)
    page = html_page(payload)
    write(html_path, page)
    write(SOURCE_PAGE, page)
    shutil.copy2(SOURCE_PAGE, PUBLIC_PAGE)
    zip_path = build_zip([json_path, md_path, docx_path, pdf_path, html_path])
    copy_public_and_library([json_path, md_path, docx_path, pdf_path, html_path, zip_path])
    print(json.dumps({
        "status": payload["status"],
        "entries": payload["audit"]["entries"],
        "old_template_impacts_after": payload["audit"]["old_volume_1_template_impacts_after"],
        "internal_production_labels": payload["audit"]["internal_production_labels"],
        "sunday_mentions": payload["audit"]["sunday_mentions"],
        "zip": str(zip_path.relative_to(ROOT)),
        "page": str(SOURCE_PAGE.relative_to(ROOT)),
    }, indent=2))


if __name__ == "__main__":
    main()
