#!/usr/bin/env python3
"""Build the Lady D author decision sheet and source-control gate pack."""

from __future__ import annotations

import hashlib
import html
import json
import re
import shutil
import subprocess
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "downloads" / "production" / "kdp" / "author-decision-sheet"
PUBLIC_OUT = ROOT / "public" / "downloads" / "production" / "kdp" / "author-decision-sheet"
SOURCE_PAGE = ROOT / "author-decision-sheet.html"
PUBLIC_PAGE = ROOT / "public" / "author-decision-sheet.html"
LIBRARY_ROOT = Path("/Users/IDC2.5/Documents/LADY D/Production Library")
LIBRARY_OUT = LIBRARY_ROOT / "_Shared" / "KDP Readiness" / "Author Decision Sheet"
GENERATED = "2026-07-08"
AUTHOR = 'Susan "Lady D" Damon'

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(89, 96, 108)
GOLD = RGBColor(122, 90, 0)
LIGHT_FILL = "E8EEF5"
SOFT_FILL = "F4F6F9"


@dataclass(frozen=True)
class Product:
    key: str
    volume: int
    kind: str
    title: str
    subtitle: str
    pages: int
    file_path: str
    current_cover_direction: str
    proof_role: str


PRODUCTS = [
    Product(
        "volume-1-devotional",
        1,
        "Devotional",
        "Surrendering to God's Love",
        "A 365-Day Devotional Journey into the Father's Heart",
        369,
        "downloads/production/kdp/interior-drafts/volume-1/volume-1-full-6x9-interior-draft.pdf",
        "warm Father-heart visual lane with refined path/light imagery",
        "primary launch candidate",
    ),
    Product(
        "volume-1-journal",
        1,
        "Companion Journal",
        "Surrendering to God's Love Companion Journal",
        "A Companion Journal for Receiving the Father's Heart",
        470,
        "downloads/production/kdp/companion-journal-drafts/volume-1/volume-1-companion-journal-6x9-draft.pdf",
        "paired journal cover, quieter and writing-forward",
        "paired support product",
    ),
    Product(
        "volume-2-devotional",
        2,
        "Devotional",
        "Walking with Jesus",
        "A 365-Day Devotional Journey with the Son",
        369,
        "downloads/production/kdp/interior-drafts/volume-2/volume-2-full-6x9-interior-draft.pdf",
        "Jesus-following visual lane with road, presence, and discipleship light",
        "second devotional candidate",
    ),
    Product(
        "volume-2-journal",
        2,
        "Companion Journal",
        "Walking with Jesus Companion Journal",
        "A Companion Journal for Following the Son",
        477,
        "downloads/production/kdp/companion-journal-drafts/volume-2/volume-2-companion-journal-6x9-draft.pdf",
        "paired journal cover, clearer writing space and devotional continuity",
        "paired support product",
    ),
    Product(
        "volume-3-devotional",
        3,
        "Devotional",
        "Filled with the Holy Spirit",
        "A 365-Day Devotional Journey of Power, Comfort, and Fire",
        369,
        "downloads/production/kdp/interior-drafts/volume-3/volume-3-full-6x9-interior-draft.pdf",
        "Spirit-filled visual lane with warmth, movement, oil/fire/rain restraint",
        "third devotional candidate",
    ),
    Product(
        "volume-3-journal",
        3,
        "Companion Journal",
        "Filled with the Holy Spirit Companion Journal",
        "A Companion Journal for Spirit-Filled Surrender",
        483,
        "downloads/production/kdp/companion-journal-drafts/volume-3/volume-3-companion-journal-6x9-draft.pdf",
        "paired journal cover, contemplative but bright enough for retail thumbnails",
        "paired support product",
    ),
]


SOURCE_FILES = [
    "downloads/production/master/volume-1-master-interior-manuscript.md",
    "downloads/production/master/volume-2-master-interior-manuscript.md",
    "downloads/production/master/volume-3-master-interior-manuscript.md",
    "downloads/production/master/volume-1-master-companion-journal.md",
    "downloads/production/master/volume-2-master-companion-journal.md",
    "downloads/production/master/volume-3-master-companion-journal.md",
    "downloads/production/kdp/interior-drafts/volume-1/volume-1-full-6x9-interior-draft.pdf",
    "downloads/production/kdp/interior-drafts/volume-2/volume-2-full-6x9-interior-draft.pdf",
    "downloads/production/kdp/interior-drafts/volume-3/volume-3-full-6x9-interior-draft.pdf",
    "downloads/production/kdp/companion-journal-drafts/volume-1/volume-1-companion-journal-6x9-draft.pdf",
    "downloads/production/kdp/companion-journal-drafts/volume-2/volume-2-companion-journal-6x9-draft.pdf",
    "downloads/production/kdp/companion-journal-drafts/volume-3/volume-3-companion-journal-6x9-draft.pdf",
    "downloads/production/kdp/full-wrap-drafts/volume-1-path-route-white-paper-full-wrap-draft.pdf",
    "downloads/production/kdp/full-wrap-drafts/volume-2-path-route-white-paper-full-wrap-draft.pdf",
    "downloads/production/kdp/full-wrap-drafts/volume-3-path-route-white-paper-full-wrap-draft.pdf",
    "downloads/production/kdp/companion-journal-full-wrap-drafts/volume-1-companion-journal-path-route-white-paper-full-wrap-draft.pdf",
    "downloads/production/kdp/companion-journal-full-wrap-drafts/volume-2-companion-journal-path-route-white-paper-full-wrap-draft.pdf",
    "downloads/production/kdp/companion-journal-full-wrap-drafts/volume-3-companion-journal-path-route-white-paper-full-wrap-draft.pdf",
]


KDP_SOURCES = [
    {
        "title": "KDP Proof and Author Copies",
        "url": "https://kdp.amazon.com/en_US/help/topic/G7BBN68RYX5UMDZF",
        "used_for": "proof versus author-copy distinction",
    },
    {
        "title": "KDP proof/author copy cost",
        "url": "https://kdp.amazon.com/en_US/help/topic/G2MYNEKHT443C2H2",
        "used_for": "author price equals print cost times quantity",
    },
    {
        "title": "KDP proof/author copy shipping",
        "url": "https://kdp.amazon.com/en_US/help/topic/GG6GRS7TKXVG6AGW",
        "used_for": "shipping is checked at checkout and is not free or Prime",
    },
    {
        "title": "KDP paperback printing cost",
        "url": "https://kdp.amazon.com/en_US/help/topic/G201834340",
        "used_for": "US black-ink regular-trim formula",
    },
    {
        "title": "KDP Printing Cost and Royalty Calculator",
        "url": "https://kdp.amazon.com/en_US/help/topic/GSQF43YAMUPFTMSP",
        "used_for": "final estimate must be verified in KDP calculator",
    },
]


DECISIONS = [
    {
        "area": "Title lock",
        "owner": "Lady D + Jon",
        "recommended_path": "Approve the current trilogy titles unless Lady D wants a stronger market-facing subtitle pass.",
        "decision_needed": "Final title and subtitle for each devotional and each companion journal.",
        "done_when": "Every KDP product title/subtitle is approved in writing and matches the cover/interior/title page.",
    },
    {
        "area": "Bible policy",
        "owner": "Lady D + Jon",
        "recommended_path": "Keep reference-only for the next proof, or approve one translation before adding full Scripture text.",
        "decision_needed": "References only, public-domain text, or licensed translation with required notice.",
        "done_when": "The copyright page carries the exact Bible translation and permission notice strategy.",
    },
    {
        "area": "July 6 content feedback",
        "owner": "IDC",
        "recommended_path": "Apply to samples first: fuller Scripture presence, warmer voice, simpler context lens, fused prompts.",
        "decision_needed": "Whether every day gets quoted Scripture text or a more compact reference-plus-reader-reflection format.",
        "done_when": "Sample pages are approved before trilogy-wide conversion.",
    },
    {
        "area": "Cover direction",
        "owner": "Lady D + Jon",
        "recommended_path": "Choose one coherent cover lane per volume, then regenerate final wraps only after page counts and paper type are locked.",
        "decision_needed": "Final cover concept, brightness, author name treatment, subtitle scale, paper type, and barcode plan.",
        "done_when": "The selected front cover and full-wrap proof pass thumbnail, spine, bleed, and barcode review.",
    },
    {
        "area": "Front/back matter",
        "owner": "Lady D",
        "recommended_path": "Approve dedication, author welcome, bio, acknowledgments, and series page before final interior export.",
        "decision_needed": "Final wording for author-sensitive pages.",
        "done_when": "No placeholders remain in copyright, dedication, bio, acknowledgments, or IDC imprint notes.",
    },
    {
        "area": "Launch order",
        "owner": "Lady D + Jon",
        "recommended_path": "Launch Volume 1 and its companion journal first, then Volumes 2 and 3 after proof confidence.",
        "decision_needed": "Single-volume launch, trilogy launch, or staggered release calendar.",
        "done_when": "KDP metadata, prices, categories, and review copies match the selected release sequence.",
    },
    {
        "area": "Adventist guardrail",
        "owner": "IDC theological audit",
        "recommended_path": "Keep Sabbath as seventh-day/Saturday and obedience as response to grace across copy, metadata, and prompts.",
        "decision_needed": "Confirm no Sunday-as-Sabbath language and no performance-based salvation phrasing.",
        "done_when": "Final proof audit shows zero Sunday-as-Sabbath drift and no earning-love framing.",
    },
    {
        "area": "KDP proof economics",
        "owner": "IDC + Jon",
        "recommended_path": "Use current KDP calculator and checkout for final proof/author-copy costs; use this sheet only for planning estimates.",
        "decision_needed": "How many proofs, who receives them, and what retail-price scenarios should be modeled.",
        "done_when": "Proof order quantity, shipping destination, retail prices, and author-copy path are documented.",
    },
]


def current_commit() -> str:
    try:
        return subprocess.check_output(["git", "rev-parse", "--short", "HEAD"], cwd=ROOT, text=True).strip()
    except Exception:
        return "unknown"


def write(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    cleaned = "\n".join(line.rstrip() for line in content.rstrip().splitlines())
    path.write_text(cleaned + "\n", encoding="utf-8")


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def word_count(path: Path) -> int | None:
    if path.suffix.lower() != ".md" or not path.exists():
        return None
    text = path.read_text(encoding="utf-8")
    return len(re.findall(r"\b[A-Za-z][A-Za-z']*\b", text))


def file_manifest(commit: str) -> dict[str, object]:
    rows = []
    for rel in SOURCE_FILES:
        path = ROOT / rel
        rows.append(
            {
                "path": rel,
                "exists": path.exists(),
                "bytes": path.stat().st_size if path.exists() else None,
                "sha256": sha256(path) if path.exists() else None,
                "word_count": word_count(path),
            }
        )
    return {
        "generated": GENERATED,
        "commit": commit,
        "status": "source_freeze_manifest_for_review_not_final_upload",
        "freeze_boundary": "These hashes describe the current review surfaces only. Final freeze happens after author decisions, Bible policy, final copyedit, cover lock, KDP Previewer, and physical proof.",
        "files": rows,
    }


def kdp_copy_rows() -> list[dict[str, object]]:
    rows = []
    for product in PRODUCTS:
        print_cost = round(1.00 + product.pages * 0.012, 2)
        rows.append(
            {
                "product": product.title,
                "kind": product.kind,
                "pages": product.pages,
                "amazon_com_black_ink_regular_trim_estimate": f"${print_cost:.2f}",
                "notes": "Planning estimate before shipping/tax; verify in KDP calculator and checkout.",
            }
        )
    return rows


def payload(commit: str, manifest: dict[str, object]) -> dict[str, object]:
    cost_rows = kdp_copy_rows()
    subtotal = sum(float(str(row["amazon_com_black_ink_regular_trim_estimate"]).replace("$", "")) for row in cost_rows)
    return {
        "generated": GENERATED,
        "commit": commit,
        "author": AUTHOR,
        "status": "author_decision_gate_ready_not_final_upload",
        "design_preset": "compact_reference_guide with customer_pack opening",
        "products": [product.__dict__ for product in PRODUCTS],
        "decisions": DECISIONS,
        "kdp_copy_economics": {
            "assumptions": [
                "US Amazon.com marketplace.",
                "Paperback, black ink, regular 6 x 9 trim, 110-828 page formula.",
                "Current review page counts, before final copyedit/page-count lock.",
                "Shipping, taxes, and marketplace differences are excluded.",
            ],
            "one_each_print_subtotal_before_shipping_tax": f"${subtotal:.2f}",
            "rows": cost_rows,
            "source_urls": KDP_SOURCES,
        },
        "source_manifest": manifest,
        "release_boundary": [
            "This sheet is a decision and source-control gate, not a final upload approval.",
            "Do not regenerate final wraps until trim, paper type, page counts, title/subtitle, and barcode strategy are locked.",
            "Do not add full Bible quotation text until the Bible policy is approved and the copyright/permission notice is inserted.",
            "Do not use public release language until KDP Previewer and physical proof review pass.",
            "Keep Juan Damon testimony/autobiography separate unless a new scope is approved.",
        ],
    }


def markdown(payload_data: dict[str, object]) -> str:
    product_rows = "\n".join(
        f"| {item['key']} | {item['title']} | {item['kind']} | {item['pages']} | {item['proof_role']} |"
        for item in payload_data["products"]
    )
    decision_rows = "\n".join(
        f"| {item['area']} | {item['owner']} | {item['decision_needed']} | {item['done_when']} |"
        for item in payload_data["decisions"]
    )
    kdp_rows = "\n".join(
        f"| {row['product']} | {row['pages']} | {row['amazon_com_black_ink_regular_trim_estimate']} | {row['notes']} |"
        for row in payload_data["kdp_copy_economics"]["rows"]
    )
    source_rows = "\n".join(
        f"| `{row['path']}` | {row['exists']} | {row['bytes'] or ''} | `{str(row['sha256'] or '')[:16]}` | {row['word_count'] or ''} |"
        for row in payload_data["source_manifest"]["files"]
    )
    source_links = "\n".join(
        f"- {item['title']}: {item['url']} ({item['used_for']})"
        for item in payload_data["kdp_copy_economics"]["source_urls"]
    )
    decisions_detail = "\n".join(
        f"### {idx}. {item['area']}\n\n- Owner: {item['owner']}\n- Recommended path: {item['recommended_path']}\n- Decision needed: {item['decision_needed']}\n- Done when: {item['done_when']}"
        for idx, item in enumerate(payload_data["decisions"], start=1)
    )
    boundary = "\n".join(f"- {item}" for item in payload_data["release_boundary"])
    assumptions = "\n".join(f"- {item}" for item in payload_data["kdp_copy_economics"]["assumptions"])
    return f"""# Lady D Author Decision Sheet

Generated: {payload_data['generated']}

Repo commit at generation: `{payload_data['commit']}`

Author: {payload_data['author']}

Status: Author decision gate ready. This is not final KDP upload approval.

Design preset: `{payload_data['design_preset']}`

## Purpose

This packet turns the July 8 plan into an approval surface. It lets Lady D and Jon lock the decisions that must be settled before final interiors, final wraps, KDP Previewer, and physical proof ordering.

## Product Approval Matrix

| Product Key | Title | Kind | Review Pages | Proof Role |
| --- | --- | --- | ---: | --- |
{product_rows}

## Immediate Decision Board

| Area | Owner | Decision Needed | Done When |
| --- | --- | --- | --- |
{decision_rows}

## Decision Detail

{decisions_detail}

## KDP Copy Economics Snapshot

Assumptions:

{assumptions}

Planning subtotal for one print copy of each of the six review products before shipping/tax: **{payload_data['kdp_copy_economics']['one_each_print_subtotal_before_shipping_tax']}**.

| Product | Pages | US Print Estimate | Notes |
| --- | ---: | ---: | --- |
{kdp_rows}

Official KDP references checked on {payload_data['generated']}:

{source_links}

## Source Freeze Manifest Snapshot

Freeze boundary: {payload_data['source_manifest']['freeze_boundary']}

| File | Exists | Bytes | SHA-256 Prefix | Words |
| --- | --- | ---: | --- | ---: |
{source_rows}

## Release Boundary

{boundary}
"""


def economics_markdown(payload_data: dict[str, object]) -> str:
    rows = "\n".join(
        f"| {row['product']} | {row['pages']} | {row['amazon_com_black_ink_regular_trim_estimate']} | {row['notes']} |"
        for row in payload_data["kdp_copy_economics"]["rows"]
    )
    assumptions = "\n".join(f"- {item}" for item in payload_data["kdp_copy_economics"]["assumptions"])
    sources = "\n".join(
        f"- {item['title']}: {item['url']} ({item['used_for']})"
        for item in payload_data["kdp_copy_economics"]["source_urls"]
    )
    return f"""# Lady D Amazon KDP Copy Economics Snapshot

Generated: {payload_data['generated']}

Status: Planning economics only. Verify inside KDP before ordering.

## Assumptions

{assumptions}

## Estimated Print Cost Before Shipping And Tax

| Product | Pages | US Print Estimate | Notes |
| --- | ---: | ---: | --- |
{rows}

Subtotal for one print copy of each product before shipping/tax: **{payload_data['kdp_copy_economics']['one_each_print_subtotal_before_shipping_tax']}**.

## Operating Notes

- Proof copies are for reviewing an unpublished upload before publication.
- Author copies apply after a book is live and use the live book file, not draft changes.
- Proof and author copy purchases do not create royalty payments.
- Shipping is not covered by free or Prime shipping for proof/author copy orders; confirm shipping at checkout before placing the order.
- Final retail prices should be modeled inside the KDP calculator after paper type and page counts are locked.

## Official KDP Sources Checked

{sources}
"""


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr[idx], LIGHT_FILL)
        run = hdr[idx].paragraphs[0].add_run(str(header))
        run.bold = True
        run.font.color.rgb = DARK_BLUE
        hdr[idx].paragraphs[0].paragraph_format.space_after = Pt(0)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row[: len(headers)]):
            cells[idx].text = str(value)
            cells[idx].paragraphs[0].paragraph_format.space_after = Pt(0)
    set_table_width(table, widths)


def add_para(doc: Document, text: str, *, bold: bool = False, color: RGBColor | None = None, size: float | None = None, after: int = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.name = "Calibri"
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    run.bold = bold


def build_docx(payload_data: dict[str, object], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "Lady D Author Decision Sheet"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = "IDC Publishing - Decision gate, not final upload approval"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(0)
    kicker_run = kicker.add_run("IDC Publishing KDP readiness")
    kicker_run.font.name = "Calibri"
    kicker_run.font.size = Pt(10)
    kicker_run.font.bold = True
    kicker_run.font.color.rgb = GOLD

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("Lady D Author Decision Sheet")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(30)
    title_run.font.bold = True
    title_run.font.color.rgb = INK

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    sub_run = subtitle.add_run("Title, Bible policy, cover, front/back matter, source freeze, and proof economics approval surface")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = MUTED

    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Author", payload_data["author"]],
            ["Generated", payload_data["generated"]],
            ["Status", "Decision gate ready; not final KDP upload approval"],
            ["Commit", payload_data["commit"]],
            ["Design preset", payload_data["design_preset"]],
        ],
        [1.55, 4.95],
    )

    doc.add_heading("Product Approval Matrix", level=1)
    product_rows = [
        [item["key"], item["title"], item["kind"], item["pages"], item["proof_role"]]
        for item in payload_data["products"]
    ]
    add_table(doc, ["Key", "Title", "Kind", "Pages", "Proof Role"], product_rows, [1.3, 2.0, 1.05, 0.6, 1.55])

    doc.add_heading("Immediate Decision Board", level=1)
    decision_rows = [
        [item["area"], item["owner"], item["decision_needed"], item["done_when"]]
        for item in payload_data["decisions"]
    ]
    add_table(doc, ["Area", "Owner", "Decision Needed", "Done When"], decision_rows, [1.15, 1.0, 2.15, 2.2])

    doc.add_page_break()
    doc.add_heading("KDP Copy Economics Snapshot", level=1)
    add_para(doc, f"Planning subtotal for one print copy of each of the six current review products before shipping/tax: {payload_data['kdp_copy_economics']['one_each_print_subtotal_before_shipping_tax']}", bold=True)
    economics_rows = [
        [row["product"], row["pages"], row["amazon_com_black_ink_regular_trim_estimate"], row["notes"]]
        for row in payload_data["kdp_copy_economics"]["rows"]
    ]
    add_table(doc, ["Product", "Pages", "Estimate", "Notes"], economics_rows, [2.6, 0.65, 0.8, 2.45])

    doc.add_heading("Source Freeze Snapshot", level=1)
    source_rows = [
        [Path(row["path"]).name, row["exists"], row["bytes"] or "", str(row["sha256"] or "")[:16]]
        for row in payload_data["source_manifest"]["files"][:18]
    ]
    add_table(doc, ["File", "Exists", "Bytes", "SHA-256 Prefix"], source_rows, [2.9, 0.75, 1.0, 1.85])

    doc.add_heading("Release Boundary", level=1)
    for item in payload_data["release_boundary"]:
        add_para(doc, item)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def pdf_paragraph(text: object, style: ParagraphStyle) -> Paragraph:
    return Paragraph(html.escape(str(text)).replace("\n", "<br/>"), style)


def build_pdf(payload_data: dict[str, object], path: Path) -> None:
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("DecisionH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=20, leading=23, textColor=colors.HexColor("#111827"), spaceAfter=10)
    h2 = ParagraphStyle("DecisionH2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=colors.HexColor("#2E74B5"), spaceBefore=12, spaceAfter=6)
    body = ParagraphStyle("DecisionBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.1, leading=11.4, alignment=TA_LEFT, spaceAfter=5)
    small = ParagraphStyle("DecisionSmall", parent=body, fontSize=7.7, leading=9.5)

    def make_table(headers: list[str], rows: list[list[object]], widths: list[float]) -> Table:
        data = [[pdf_paragraph(header, small) for header in headers]]
        for row in rows:
            data.append([pdf_paragraph(cell, small) for cell in row])
        table = Table(data, colWidths=[width * inch for width in widths], repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#AAB7C4")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )
        return table

    story = [
        pdf_paragraph("Lady D Author Decision Sheet", h1),
        pdf_paragraph("Title, Bible policy, cover, front/back matter, source freeze, and proof economics approval surface", body),
        pdf_paragraph("Status: decision gate ready; not final KDP upload approval.", body),
        Spacer(1, 8),
        pdf_paragraph("Product Approval Matrix", h2),
    ]
    story.append(
        make_table(
            ["Product", "Kind", "Pages", "Proof Role"],
            [[item["title"], item["kind"], item["pages"], item["proof_role"]] for item in payload_data["products"]],
            [2.65, 1.0, 0.6, 2.0],
        )
    )
    story.extend([Spacer(1, 8), pdf_paragraph("Immediate Decision Board", h2)])
    story.append(
        make_table(
            ["Area", "Owner", "Decision Needed", "Done When"],
            [[item["area"], item["owner"], item["decision_needed"], item["done_when"]] for item in payload_data["decisions"]],
            [1.05, 0.9, 2.15, 2.15],
        )
    )
    story.extend([Spacer(1, 8), pdf_paragraph("KDP Copy Economics", h2)])
    story.append(pdf_paragraph(f"Planning subtotal for one copy of each current review product before shipping/tax: {payload_data['kdp_copy_economics']['one_each_print_subtotal_before_shipping_tax']}", body))
    story.append(
        make_table(
            ["Product", "Pages", "Estimate"],
            [[row["product"], row["pages"], row["amazon_com_black_ink_regular_trim_estimate"]] for row in payload_data["kdp_copy_economics"]["rows"]],
            [4.25, 0.8, 1.0],
        )
    )
    story.extend([Spacer(1, 8), pdf_paragraph("Release Boundary", h2)])
    for item in payload_data["release_boundary"]:
        story.append(pdf_paragraph(item, body))

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.68 * inch, leftMargin=0.68 * inch, topMargin=0.68 * inch, bottomMargin=0.68 * inch)
    doc.build(story)


def html_page(payload_data: dict[str, object]) -> str:
    cards = [
        ("Decision Gate", "8", "Title, Bible, cover, front/back matter, launch order, guardrails, economics."),
        ("Products", "6", "Three devotionals plus three companion journals."),
        ("Proof Estimate", payload_data["kdp_copy_economics"]["one_each_print_subtotal_before_shipping_tax"], "One print copy each before shipping/tax, verify in KDP."),
        ("Source Files", str(len(payload_data["source_manifest"]["files"])), "Current review surfaces hashed for source-control review."),
    ]
    card_html = "".join(
        f'<article class="card"><span>{html.escape(label)}</span><strong>{html.escape(value)}</strong><p>{html.escape(desc)}</p></article>'
        for label, value, desc in cards
    )
    decision_rows = "".join(
        f"<tr><td>{html.escape(item['area'])}</td><td>{html.escape(item['owner'])}</td><td>{html.escape(item['decision_needed'])}</td><td>{html.escape(item['done_when'])}</td></tr>"
        for item in payload_data["decisions"]
    )
    product_rows = "".join(
        f"<tr><td>{html.escape(item['title'])}</td><td>{html.escape(item['kind'])}</td><td>{item['pages']}</td><td>{html.escape(item['current_cover_direction'])}</td></tr>"
        for item in payload_data["products"]
    )
    economics_rows = "".join(
        f"<tr><td>{html.escape(row['product'])}</td><td>{row['pages']}</td><td>{html.escape(row['amazon_com_black_ink_regular_trim_estimate'])}</td><td>{html.escape(row['notes'])}</td></tr>"
        for row in payload_data["kdp_copy_economics"]["rows"]
    )
    source_rows = "".join(
        f"<tr><td>{html.escape(Path(row['path']).name)}</td><td>{html.escape(str(row['exists']))}</td><td>{html.escape(str(row['bytes'] or ''))}</td><td><code>{html.escape(str(row['sha256'] or '')[:16])}</code></td></tr>"
        for row in payload_data["source_manifest"]["files"]
    )
    boundary_items = "".join(f"<li>{html.escape(item)}</li>" for item in payload_data["release_boundary"])
    source_links = "".join(
        f'<a href="{html.escape(item["url"])}">{html.escape(item["title"])}</a>'
        for item in payload_data["kdp_copy_economics"]["source_urls"]
    )
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Lady D Author Decision Sheet</title>
  <style>
    :root {{ --ink:#111827; --muted:#5b6474; --paper:#fffdf8; --mist:#f5f2eb; --indigo:#182646; --teal:#1d716f; --gold:#c99335; --coral:#b86464; --line:rgba(17,24,39,.14); --shadow:0 20px 58px rgba(24,38,70,.12); }}
    * {{ box-sizing:border-box; }}
    body {{ margin:0; color:var(--ink); font-family:Inter,ui-sans-serif,system-ui,-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif; background:linear-gradient(180deg,var(--paper),var(--mist)); line-height:1.5; }}
    nav {{ position:sticky; top:0; z-index:10; display:flex; flex-wrap:wrap; gap:8px; align-items:center; padding:10px clamp(16px,4vw,42px); background:rgba(24,38,70,.96); color:white; }}
    nav a, nav strong {{ color:white; text-decoration:none; font-weight:850; font-size:14px; }}
    header, main {{ max-width:1160px; margin:0 auto; padding:44px 22px; }}
    h1,h2,h3 {{ font-family:Georgia,"Times New Roman",serif; letter-spacing:0; line-height:1.05; margin:0 0 14px; }}
    h1 {{ font-size:clamp(42px,7vw,80px); max-width:980px; }}
    h2 {{ font-size:clamp(28px,4vw,46px); }}
    p {{ margin:0 0 14px; }}
    .lead {{ max-width:850px; font-size:clamp(18px,2vw,23px); color:#2e3746; }}
    .kicker {{ color:var(--teal); font-weight:950; letter-spacing:.14em; text-transform:uppercase; font-size:12px; margin-bottom:14px; }}
    .status {{ display:inline-block; background:var(--indigo); color:white; padding:7px 10px; border-radius:999px; font-size:12px; font-weight:900; margin:0 6px 8px 0; }}
    .actions {{ display:flex; flex-wrap:wrap; gap:10px; margin-top:18px; }}
    .actions a {{ display:inline-flex; align-items:center; min-height:38px; padding:8px 12px; border:1px solid var(--line); border-radius:999px; color:var(--teal); background:white; font-size:13px; font-weight:900; text-decoration:none; }}
    section {{ border-top:1px solid var(--line); padding:34px 0; }}
    .grid {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(230px,1fr)); gap:14px; }}
    .card, .panel {{ border:1px solid var(--line); background:white; border-radius:8px; padding:18px; box-shadow:var(--shadow); }}
    .card span {{ color:var(--gold); font-size:12px; font-weight:900; text-transform:uppercase; letter-spacing:.12em; }}
    .card strong {{ display:block; font-family:Georgia,"Times New Roman",serif; font-size:34px; line-height:1.05; margin:8px 0; color:#172247; }}
    table {{ width:100%; border-collapse:collapse; background:white; border:1px solid var(--line); border-radius:8px; overflow:hidden; }}
    th,td {{ padding:11px 12px; border-bottom:1px solid var(--line); vertical-align:top; text-align:left; }}
    th {{ background:#f4f6f9; color:var(--indigo); font-size:12px; text-transform:uppercase; letter-spacing:.08em; }}
    .table-wrap {{ overflow:auto; border-radius:8px; }}
    code {{ font-family:ui-monospace,SFMono-Regular,Menlo,Consolas,monospace; font-size:.9em; }}
    ul {{ margin:0; padding-left:22px; }}
    li {{ margin:8px 0; }}
    @media (max-width:720px) {{ nav {{ position:static; }} th,td {{ font-size:13px; }} }}
  </style>
</head>
<body>
  <nav>
    <strong>Lady D KDP Gate</strong>
    <a href="susan-damon-hub.html">Hub</a>
    <a href="production.html">Production Review</a>
    <a href="release-status.html">Release Dashboard</a>
    <a href="lady-d-plan-of-attack-2026-07-08.html">Plan</a>
    <a href="#decisions">Decisions</a>
    <a href="#economics">KDP Costs</a>
    <a href="#sources">Sources</a>
  </nav>
  <header>
    <div class="kicker">IDC Publishing author decision gate</div>
    <h1>Lady D Author Decision Sheet</h1>
    <p class="lead">A single approval surface for the decisions that must be settled before final interiors, final wraps, KDP Previewer, physical proofs, and public release language.</p>
    <p><span class="status">Generated {html.escape(payload_data['generated'])}</span><span class="status">Commit {html.escape(payload_data['commit'])}</span><span class="status">Not final upload approval</span></p>
    <div class="actions">
      <a href="downloads/production/kdp/author-decision-sheet/Lady-D-Author-Decision-Sheet-Pack.zip">Download ZIP</a>
      <a href="downloads/production/kdp/author-decision-sheet/lady-d-author-decision-sheet.pdf">PDF</a>
      <a href="downloads/production/kdp/author-decision-sheet/lady-d-author-decision-sheet.docx">DOCX</a>
      <a href="downloads/production/kdp/author-decision-sheet/lady-d-author-decision-sheet.md">Markdown</a>
      <a href="downloads/production/kdp/author-decision-sheet/lady-d-author-decision-sheet.json">JSON</a>
      <a href="downloads/production/kdp/author-decision-sheet/source-freeze-manifest.json">Source Manifest</a>
      <a href="downloads/production/kdp/author-decision-sheet/amazon-kdp-copy-economics.md">KDP Economics</a>
    </div>
  </header>
  <main>
    <section><div class="grid">{card_html}</div></section>
    <section>
      <h2>Product Approval Matrix</h2>
      <div class="table-wrap"><table><thead><tr><th>Product</th><th>Kind</th><th>Pages</th><th>Cover Direction</th></tr></thead><tbody>{product_rows}</tbody></table></div>
    </section>
    <section id="decisions">
      <h2>Immediate Decision Board</h2>
      <div class="table-wrap"><table><thead><tr><th>Area</th><th>Owner</th><th>Decision Needed</th><th>Done When</th></tr></thead><tbody>{decision_rows}</tbody></table></div>
    </section>
    <section id="economics">
      <h2>KDP Copy Economics</h2>
      <p class="lead">Planning subtotal for one print copy of each current review product before shipping/tax: <strong>{html.escape(payload_data['kdp_copy_economics']['one_each_print_subtotal_before_shipping_tax'])}</strong>. Verify the final numbers inside KDP before ordering.</p>
      <div class="table-wrap"><table><thead><tr><th>Product</th><th>Pages</th><th>Estimate</th><th>Notes</th></tr></thead><tbody>{economics_rows}</tbody></table></div>
    </section>
    <section id="sources">
      <h2>Source Freeze Snapshot</h2>
      <p class="lead">{html.escape(payload_data['source_manifest']['freeze_boundary'])}</p>
      <div class="table-wrap"><table><thead><tr><th>File</th><th>Exists</th><th>Bytes</th><th>SHA-256 Prefix</th></tr></thead><tbody>{source_rows}</tbody></table></div>
    </section>
    <section>
      <h2>Release Boundary</h2>
      <div class="panel"><ul>{boundary_items}</ul></div>
    </section>
    <section>
      <h2>Official KDP References</h2>
      <div class="actions">{source_links}</div>
    </section>
  </main>
</body>
</html>
"""


def make_zip(paths: list[Path]) -> Path:
    zip_path = OUT / "Lady-D-Author-Decision-Sheet-Pack.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in paths:
            archive.write(path, path.name)
    return zip_path


def copy_outputs(paths: list[Path]) -> None:
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)
    for path in paths:
        shutil.copy2(path, PUBLIC_OUT / path.name)
        shutil.copy2(path, LIBRARY_OUT / path.name)


def main() -> None:
    commit = current_commit()
    OUT.mkdir(parents=True, exist_ok=True)
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)

    manifest = file_manifest(commit)
    data = payload(commit, manifest)

    manifest_path = OUT / "source-freeze-manifest.json"
    payload_path = OUT / "lady-d-author-decision-sheet.json"
    markdown_path = OUT / "lady-d-author-decision-sheet.md"
    economics_path = OUT / "amazon-kdp-copy-economics.md"
    docx_path = OUT / "lady-d-author-decision-sheet.docx"
    pdf_path = OUT / "lady-d-author-decision-sheet.pdf"
    review_page = OUT / "lady-d-author-decision-sheet-review.html"

    write(manifest_path, json.dumps(manifest, indent=2))
    write(payload_path, json.dumps(data, indent=2))
    write(markdown_path, markdown(data))
    write(economics_path, economics_markdown(data))
    build_docx(data, docx_path)
    build_pdf(data, pdf_path)
    write(review_page, html_page(data))

    zip_path = make_zip([manifest_path, payload_path, markdown_path, economics_path, docx_path, pdf_path, review_page])
    all_paths = [manifest_path, payload_path, markdown_path, economics_path, docx_path, pdf_path, review_page, zip_path]
    copy_outputs(all_paths)

    shutil.copy2(review_page, SOURCE_PAGE)
    PUBLIC_PAGE.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(review_page, PUBLIC_PAGE)
    shutil.copy2(review_page, LIBRARY_OUT / SOURCE_PAGE.name)

    print(
        json.dumps(
            {
                "status": data["status"],
                "page": str(SOURCE_PAGE.relative_to(ROOT)),
                "zip": str(zip_path.relative_to(ROOT)),
                "print_subtotal_before_shipping_tax": data["kdp_copy_economics"]["one_each_print_subtotal_before_shipping_tax"],
                "source_files": len(manifest["files"]),
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
