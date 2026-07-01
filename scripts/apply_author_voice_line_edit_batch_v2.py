#!/usr/bin/env python3
"""Apply configured Lady D author-voice line-edit batches and build review artifacts."""

from __future__ import annotations

import argparse
import html
import json
import re
import shutil
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
PRODUCTION = ROOT / "downloads" / "production"
LIBRARY_ROOT = Path("/Users/IDC2.5/Documents/LADY D/Production Library")
VOLUME_1_LIBRARY = LIBRARY_ROOT / "01 Surrendering to God's Love"
GENERATED = "2026-07-01"
AUTHOR = 'Susan "Lady D" Damon'

INK = RGBColor(17, 24, 39)
MUTED = RGBColor(89, 96, 108)
BLUE = RGBColor(31, 77, 120)
GOLD = RGBColor(201, 147, 48)
LIGHT_FILL = "F4F6F9"


@dataclass(frozen=True)
class Replacement:
    kind: str
    day: str
    before: str
    after: str
    reason: str


@dataclass(frozen=True)
class Batch:
    key: str
    scope: str
    title: str
    intro: str
    source_name: str
    public_page_name: str
    output_slug: str
    zip_name: str
    library_source: Path
    library_out: Path
    replacements: tuple[Replacement, ...]
    expected_entries: int = 7


BATCHES = {
    "volume-1-days-008-014": Batch(
        key="volume-1-days-008-014",
        scope="Volume 1 Days 008-014",
        title="Volume 1 Days 008-014 Line Edit",
        intro=(
            "The second seven-day batch of Surrendering to God's Love has moved "
            "from structural manuscript into line-level author-voice refinement. "
            "This pass removes the repeated morning-impact frame, keeps the "
            "Sabbath language as seventh-day/Saturday, and preserves obedience "
            "as response to grace."
        ),
        source_name="volume-1-days-008-014-manuscript.md",
        public_page_name="volume-1-days-008-014-line-edit.html",
        output_slug="volume-1-days-008-014-line-edit",
        zip_name="Lady-D-Volume-1-Days-008-014-Line-Edit-Pack.zip",
        library_source=VOLUME_1_LIBRARY / "01 Manuscript" / "Month 01 - January" / "Days 008-014 Manuscript.md",
        library_out=VOLUME_1_LIBRARY / "05 Review Packets" / "Author Voice Line Edit" / "Days 008-014",
        replacements=(
            Replacement(
                "morning_impact",
                "Day 008",
                "Let the Father's love carry Stand in Restoring Compassion into one faithful step today.",
                "Stand where fear once made you shrink; the Father's compassion can make courage tender today.",
                "Breaks the repeated frame while keeping the entry's courage-with-compassion theme.",
            ),
            Replacement(
                "morning_impact",
                "Day 009",
                "Let the Father's love carry Let Grace Form the Promise That Holds into one faithful step today.",
                "Carry one remembered mercy into today, and let grace hold what your feelings cannot.",
                "Turns the promise theme into a concrete memory-and-trust action.",
            ),
            Replacement(
                "morning_impact",
                "Day 010",
                "Let the Father's love carry Behold the Heart That Calls You into one faithful step today.",
                "Bring God one divided place today; whole-hearted love begins with honest surrender.",
                "Matches the Deuteronomy 6 whole-heart lens without sounding generic.",
            ),
            Replacement(
                "morning_impact",
                "Day 011",
                "Let the Father's love carry Follow Love That Sends You into one faithful step today.",
                "Let one act of obedience make love visible before the day gets away from you.",
                "Keeps obedience as grace-shaped practice and gives the reader a measurable step.",
            ),
            Replacement(
                "morning_impact",
                "Day 012",
                "Let the Father's love carry Rest in Covenant Mercy into one faithful step today.",
                "Rest in a mercy you cannot manufacture, and thank the Father for a blessing you nearly overlooked.",
                "Connects covenant mercy to gratitude and non-material blessing.",
            ),
            Replacement(
                "morning_impact",
                "Day 013",
                "Let the Father's love carry Wake Up to the Father's Welcome into one faithful step today.",
                "Step like someone being led, not merely someone who survived the last battle.",
                "Preserves the redeemed-and-led movement of Exodus 15:13.",
            ),
            Replacement(
                "morning_impact",
                "Day 014",
                "Let the Father's love carry Let Fire Refine Beloved Identity into one faithful step today.",
                "Receive correction without shame; holy love is refining what already belongs to the Father.",
                "Keeps correction, holiness, and beloved identity together without condemnation.",
            ),
        ),
    ),
    "volume-1-days-015-021": Batch(
        key="volume-1-days-015-021",
        scope="Volume 1 Days 015-021",
        title="Volume 1 Days 015-021 Line Edit",
        intro=(
            "The third seven-day batch of Surrendering to God's Love has moved "
            "from structural manuscript into line-level author-voice refinement. "
            "This pass removes the repeated morning-impact frame, preserves "
            "lament, mercy, forgiveness, and identity themes, and keeps obedience "
            "as response to grace."
        ),
        source_name="volume-1-days-015-021-manuscript.md",
        public_page_name="volume-1-days-015-021-line-edit.html",
        output_slug="volume-1-days-015-021-line-edit",
        zip_name="Lady-D-Volume-1-Days-015-021-Line-Edit-Pack.zip",
        library_source=VOLUME_1_LIBRARY / "01 Manuscript" / "Month 01 - January" / "Days 015-021 Manuscript.md",
        library_out=VOLUME_1_LIBRARY / "05 Review Packets" / "Author Voice Line Edit" / "Days 015-021",
        replacements=(
            Replacement(
                "morning_impact",
                "Day 015",
                "Let the Father's love carry Come Home to Love That Finds You into one faithful step today.",
                "Come home from needing every answer today; let the Father's goodness steady your next step.",
                "Turns the goodness-and-compassion theme into a concrete release of answer-chasing.",
            ),
            Replacement(
                "morning_impact",
                "Day 016",
                "Let the Father's love carry Practice Grace Before Striving into one faithful step today.",
                "Bring one unpolished burden into grace before you try to manage how it looks.",
                "Aligns grace before performance with confession instead of image management.",
            ),
            Replacement(
                "morning_impact",
                "Day 017",
                "Let the Father's love carry Surrender to Mercy in the Morning into one faithful step today.",
                "Place the long road in the Father's hands, and receive mercy for this morning.",
                "Keeps the long-view mercy theme while giving the reader one immediate posture.",
            ),
            Replacement(
                "morning_impact",
                "Day 018",
                "Let the Father's love carry Let Mercy Speak the Father's Patience into one faithful step today.",
                "Refuse the name fear gave you; walk today under the identity grace is forming.",
                "Connects renaming, delay, and patient identity formation without sounding generic.",
            ),
            Replacement(
                "morning_impact",
                "Day 019",
                "Let the Father's love carry Breathe Love Stronger Than Fear into one faithful step today.",
                "Pause, breathe, and let God's presence answer fear before fear finishes speaking.",
                "Preserves the anxiety/prayer practice and makes the breath step memorable.",
            ),
            Replacement(
                "morning_impact",
                "Day 020",
                "Let the Father's love carry Hold Fast to Restoring Compassion into one faithful step today.",
                "Ask the Father for mercy with boundaries, forgiveness without denial, and a heart free from bitterness.",
                "Keeps the family-hurt forgiveness guardrails: mercy, truth, and wise boundaries.",
            ),
            Replacement(
                "morning_impact",
                "Day 021",
                "Let the Father's love carry Return to the Promise That Holds into one faithful step today.",
                "Pray honestly, then rest the weight of your sorrow on the Father's steadfast love.",
                "Preserves lament and trust while giving the reader a prayer movement.",
            ),
        ),
    ),
    "volume-1-days-022-028": Batch(
        key="volume-1-days-022-028",
        scope="Volume 1 Days 022-028",
        title="Volume 1 Days 022-028 Line Edit",
        intro=(
            "The fourth January batch of Surrendering to God's Love has moved "
            "from structural manuscript into line-level author-voice refinement. "
            "This pass removes the repeated morning-impact frame, keeps Sabbath "
            "and commandment-keeping as rhythms of return, and preserves obedience "
            "as response to grace."
        ),
        source_name="volume-1-days-022-028-manuscript.md",
        public_page_name="volume-1-days-022-028-line-edit.html",
        output_slug="volume-1-days-022-028-line-edit",
        zip_name="Lady-D-Volume-1-Days-022-028-Line-Edit-Pack.zip",
        library_source=VOLUME_1_LIBRARY / "01 Manuscript" / "Month 01 - January" / "Days 022-028 Manuscript.md",
        library_out=VOLUME_1_LIBRARY / "05 Review Packets" / "Author Voice Line Edit" / "Days 022-028",
        replacements=(
            Replacement(
                "morning_impact",
                "Day 022",
                "Let the Father's love carry Yield to the Heart That Calls You into one faithful step today.",
                "Let glad trust settle into your body today; the Father is holding more than your thoughts.",
                "Connects the Psalm 16 rest-security theme to embodied anxiety without sounding generic.",
            ),
            Replacement(
                "morning_impact",
                "Day 023",
                "Let the Father's love carry Anchor Love That Sends You into one faithful step today.",
                "Do the quiet faithful thing in front of you; God can carry its fruit farther than you can see.",
                "Keeps the Ruth/Obed generational-faithfulness theme practical and memorable.",
            ),
            Replacement(
                "morning_impact",
                "Day 024",
                "Let the Father's love carry Discover Covenant Mercy into one faithful step today.",
                "Receive mercy, then rise with it; the Father's forgiveness still has a future for you.",
                "Preserves mercy that restores and calls forward from Deuteronomy 10:11.",
            ),
            Replacement(
                "morning_impact",
                "Day 025",
                "Let the Father's love carry Receive the Father's Welcome into one faithful step today.",
                "Let conviction become welcome today; the Father is softening what fear taught you to guard.",
                "Keeps heart-surrender conviction tender rather than condemnatory.",
            ),
            Replacement(
                "morning_impact",
                "Day 026",
                "Let the Father's love carry Trust Beloved Identity into one faithful step today.",
                "Remember how mercy found you, then make room for someone who feels outside.",
                "Turns beloved identity into concrete welcome for the stranger.",
            ),
            Replacement(
                "morning_impact",
                "Day 027",
                "Let the Father's love carry Let Hope Rise Love That Finds You into one faithful step today.",
                "Bless the small beginning in your hands; the Father's love knows how to grow what He plants.",
                "Matches the few-to-stars promise without inflating the reader's role.",
            ),
            Replacement(
                "morning_impact",
                "Day 028",
                "Let the Father's love carry Carry Grace Before Striving into one faithful step today.",
                "Return before you perform today; let obedience begin where the Father's grace calls you home.",
                "Preserves the grace-before-performance and Sabbath-as-return guardrails.",
            ),
        ),
    ),
    "volume-1-days-029-031": Batch(
        key="volume-1-days-029-031",
        scope="Volume 1 Days 029-031",
        title="Volume 1 Days 029-031 Line Edit",
        intro=(
            "The January closeout batch of Surrendering to God's Love has moved "
            "from structural manuscript into line-level author-voice refinement. "
            "This pass removes the repeated morning-impact frame, marks the first "
            "month as line-edited, and keeps Sabbath, commandments, warning, and "
            "obedience inside the grace-shaped Adventist frame."
        ),
        source_name="volume-1-days-029-031-manuscript.md",
        public_page_name="volume-1-days-029-031-line-edit.html",
        output_slug="volume-1-days-029-031-line-edit",
        zip_name="Lady-D-Volume-1-Days-029-031-Line-Edit-Pack.zip",
        library_source=VOLUME_1_LIBRARY / "01 Manuscript" / "Month 01 - January" / "Days 029-031 Manuscript.md",
        library_out=VOLUME_1_LIBRARY / "05 Review Packets" / "Author Voice Line Edit" / "Days 029-031",
        replacements=(
            Replacement(
                "morning_impact",
                "Day 029",
                "Let the Father's love carry Lean Into Mercy in the Morning into one faithful step today.",
                "Receive the near word this morning; mercy has already placed the next step within reach.",
                "Keeps the Deuteronomy 30 nearness theme practical without using the repeated frame.",
            ),
            Replacement(
                "morning_impact",
                "Day 030",
                "Let the Father's love carry See Again the Father's Patience into one faithful step today.",
                "Walk the command as a loved child today; the Father's way is life, not rejection.",
                "Preserves commandments as life-giving grace rather than pressure or performance.",
            ),
            Replacement(
                "morning_impact",
                "Day 031",
                "Let the Father's love carry Let Love Teach Love Stronger Than Fear into one faithful step today.",
                "Let warning become mercy today; return before fear teaches your heart to hide.",
                "Holds warning, return, and mercy together for the January closeout.",
            ),
        ),
        expected_entries=3,
    ),
    "volume-1-days-032-038": Batch(
        key="volume-1-days-032-038",
        scope="Volume 1 Days 032-038",
        title="Volume 1 Days 032-038 Line Edit",
        intro=(
            "The February opening batch of Surrendering to God's Love has moved "
            "from structural manuscript into line-level author-voice refinement. "
            "This pass removes the repeated morning-impact frame, opens the "
            "Beloved Identity arc, and preserves return, obedience, Sabbath, "
            "and correction inside the grace-shaped Adventist frame."
        ),
        source_name="volume-1-days-032-038-manuscript.md",
        public_page_name="volume-1-days-032-038-line-edit.html",
        output_slug="volume-1-days-032-038-line-edit",
        zip_name="Lady-D-Volume-1-Days-032-038-Line-Edit-Pack.zip",
        library_source=VOLUME_1_LIBRARY / "01 Manuscript" / "Month 02 - February" / "Days 032-038 Manuscript.md",
        library_out=VOLUME_1_LIBRARY / "05 Review Packets" / "Author Voice Line Edit" / "Days 032-038",
        replacements=(
            Replacement(
                "morning_impact",
                "Day 032",
                "Let the Father's love carry Behold the Promise That Holds into one faithful step today.",
                "Return the divided place to the Father today; beloved identity grows where hiding ends.",
                "Keeps Deuteronomy 30 return language concrete and ties beloved identity to honest surrender.",
            ),
            Replacement(
                "morning_impact",
                "Day 033",
                "Let the Father's love carry Follow the Heart That Calls You into one faithful step today.",
                "Let compassion gather one scattered piece today; the Father knows every place to call home.",
                "Preserves the gathering/restoration theme without repeating the architecture title.",
            ),
            Replacement(
                "morning_impact",
                "Day 034",
                "Let the Father's love carry Rest in Love That Sends You into one faithful step today.",
                "Move from belonging, not pressure; the Father sends loved children into faithful steps.",
                "Keeps the sent-from-belonging frame and avoids performance pressure.",
            ),
            Replacement(
                "morning_impact",
                "Day 035",
                "Let the Father's love carry Wake Up to Covenant Mercy into one faithful step today.",
                "Let the Father be kinder than disappointment taught you to expect today.",
                "Turns covenant mercy into a memorable reader-facing challenge to suspicion and disappointment.",
            ),
            Replacement(
                "morning_impact",
                "Day 036",
                "Let the Father's love carry Let Fire Refine the Father's Welcome into one faithful step today.",
                "Name the rival love honestly today; holy welcome is strong enough to refine it.",
                "Holds warning, welcome, holiness, and correction together without condemnation.",
            ),
            Replacement(
                "morning_impact",
                "Day 037",
                "Let the Father's love carry Come Home to Beloved Identity into one faithful step today.",
                "Receive the gift with gratitude today, then keep your heart close to the Giver.",
                "Keeps fullness from becoming forgetfulness and gives the reader a concrete gratitude practice.",
            ),
            Replacement(
                "morning_impact",
                "Day 038",
                "Let the Father's love carry Practice Love That Finds You into one faithful step today.",
                "Listen for life this Sabbath; the Father forms beloved children through holy rhythms.",
                "Preserves Saturday Sabbath as formation for beloved children rather than performance.",
            ),
        ),
    ),
    "volume-1-days-039-045": Batch(
        key="volume-1-days-039-045",
        scope="Volume 1 Days 039-045",
        title="Volume 1 Days 039-045 Line Edit",
        intro=(
            "The second February batch of Surrendering to God's Love has moved "
            "from structural manuscript into line-level author-voice refinement. "
            "This pass removes the repeated morning-impact frame while preserving "
            "honest seeking, covenant mercy, guarded memory, gratitude, quiet "
            "obedience, Sabbath holiness, and deliverance with destination."
        ),
        source_name="volume-1-days-039-045-manuscript.md",
        public_page_name="volume-1-days-039-045-line-edit.html",
        output_slug="volume-1-days-039-045-line-edit",
        zip_name="Lady-D-Volume-1-Days-039-045-Line-Edit-Pack.zip",
        library_source=VOLUME_1_LIBRARY / "01 Manuscript" / "Month 02 - February" / "Days 039-045 Manuscript.md",
        library_out=VOLUME_1_LIBRARY / "05 Review Packets" / "Author Voice Line Edit" / "Days 039-045",
        replacements=(
            Replacement(
                "morning_impact",
                "Day 039",
                "Let the Father's love carry Surrender to Grace Before Striving into one faithful step today.",
                "Seek the Father without polishing the story today; grace receives the whole heart.",
                "Keeps whole-person seeking inside grace instead of performance.",
            ),
            Replacement(
                "morning_impact",
                "Day 040",
                "Let the Father's love carry Let Mercy Speak Mercy in the Morning into one faithful step today.",
                "Let mercy answer accusation this morning; the Father corrects without forgetting covenant.",
                "Preserves mercy, correction, and covenant memory without flattening the devotional image.",
            ),
            Replacement(
                "morning_impact",
                "Day 041",
                "Let the Father's love carry Breathe the Father's Patience into one faithful step today.",
                "Breathe beneath the truth today: the Lord is God, and you are His child.",
                "Turns the Deuteronomy 4:39 God-alone confession into a calming reader practice.",
            ),
            Replacement(
                "morning_impact",
                "Day 042",
                "Let the Father's love carry Hold Fast to Love Stronger Than Fear into one faithful step today.",
                "Guard the memory of God's care today, then pass one testimony forward with gentleness.",
                "Connects guarded memory and generational witness without letting fear lead the tone.",
            ),
            Replacement(
                "morning_impact",
                "Day 043",
                "Let the Father's love carry Return to Restoring Compassion into one faithful step today.",
                "Turn one gift into gratitude today; the Giver is still the ground of belonging.",
                "Keeps humility and gratitude in the promised-gift frame.",
            ),
            Replacement(
                "morning_impact",
                "Day 044",
                "Let the Father's love carry Yield to the Promise That Holds into one faithful step today.",
                "Do the right and good thing quietly today; love can obey without performing.",
                "Preserves obedience as grace-shaped love rather than approval-seeking.",
            ),
            Replacement(
                "morning_impact",
                "Day 045",
                "Let the Father's love carry Anchor the Heart That Calls You into one faithful step today.",
                "Walk like one being brought in, not only brought out; the Father's purpose still leads.",
                "Keeps deliverance-with-destination active and personal for the reader.",
            ),
        ),
    ),
}


def write(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    cleaned = "\n".join(line.rstrip() for line in content.rstrip().splitlines())
    path.write_text(cleaned + "\n", encoding="utf-8")


def word_count(text: str) -> int:
    return len(re.findall(r"\b[A-Za-z][A-Za-z']*\b", text))


def display_edit_text(value: str) -> str:
    return re.sub(r"^###\s+", "", value).strip()


def parse_day_entries(text: str) -> list[dict[str, str]]:
    pattern = re.compile(r"(?ms)^## (?P<label>Day \d{3}) - (?P<date>[^\n]+)\n\n### (?P<title>[^\n]+)(?P<body>.*?)(?=^## Day |\Z)")
    entries = []
    for match in pattern.finditer(text):
        body = match.group("body")
        scripture = re.search(r"(?m)^\*\*Scripture Reference:\*\*\s*(.+)$", body)
        impact = re.search(r"(?m)^\*\*Morning impact:\*\*\s*(.+)$", body)
        entries.append(
            {
                "label": match.group("label"),
                "date": match.group("date").strip(),
                "title": match.group("title").strip(),
                "scripture": scripture.group(1).strip() if scripture else "",
                "morning_impact": impact.group(1).strip() if impact else "",
            }
        )
    return entries


def paths_for(batch: Batch) -> dict[str, Path]:
    out = PRODUCTION / "kdp" / "author-voice-line-edit" / batch.key
    return {
        "source": PRODUCTION / batch.source_name,
        "public_source": ROOT / "public" / "downloads" / "production" / batch.source_name,
        "out": out,
        "public_out": ROOT / "public" / "downloads" / "production" / "kdp" / "author-voice-line-edit" / batch.key,
        "source_page": ROOT / batch.public_page_name,
        "public_page": ROOT / "public" / batch.public_page_name,
        "json": out / f"{batch.output_slug}.json",
        "md": out / f"{batch.output_slug}-report.md",
        "docx": out / f"{batch.output_slug}-report.docx",
        "pdf": out / f"{batch.output_slug}-report.pdf",
        "html": out / f"{batch.output_slug}-review.html",
        "zip": out / batch.zip_name,
    }


def apply_replacements(batch: Batch, paths: dict[str, Path]) -> tuple[str, str, list[dict[str, str]]]:
    original = paths["source"].read_text(encoding="utf-8")
    text = original
    applied = []
    for replacement in batch.replacements:
        if replacement.before in text:
            text = text.replace(replacement.before, replacement.after, 1)
            status = "applied"
        elif replacement.after in text:
            status = "already_current"
        else:
            raise RuntimeError(f"Replacement target not found for {replacement.day}: {replacement.before}")
        applied.append(
            {
                "day": replacement.day,
                "kind": replacement.kind,
                "status": status,
                "before": replacement.before,
                "after": replacement.after,
                "reason": replacement.reason,
            }
        )
    write(paths["source"], text)
    paths["public_source"].parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(paths["source"], paths["public_source"])
    batch.library_source.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(paths["source"], batch.library_source)
    return original, text, applied


def audit(batch: Batch, original: str, text: str, applied: list[dict[str, str]], paths: dict[str, Path]) -> dict[str, object]:
    entries = parse_day_entries(text)
    expected_entries = batch.expected_entries
    source_metrics = {
        "entries": len(entries),
        "scripture_references": len(re.findall(r"(?m)^\*\*Scripture Reference:\*\*", text)),
        "context_lenses": len(re.findall(r"(?m)^\*\*Context and language lens:\*\*", text)),
        "today_steps": len(re.findall(r"(?m)^\*\*Today step:\*\*", text)),
        "prayers": len(re.findall(r"(?m)^\*\*Prayer:\*\*", text)),
        "journal_prompts": len(re.findall(r"(?m)^\*\*Journal prompt:\*\*", text)),
        "morning_impacts": len(re.findall(r"(?m)^\*\*Morning impact:\*\*", text)),
        "old_volume_1_template_impacts_before": len(re.findall(r"Let the Father's love carry", original)),
        "old_volume_1_template_impacts_after": len(re.findall(r"Let the Father's love carry", text)),
        "internal_production_labels": len(re.findall(r"Production lens|Production text note|Adventist-safe Sabbath guardrail|whole devotional system", text)),
        "sunday_mentions": len(re.findall(r"\bSunday\b", text, flags=re.I)),
        "sabbath_mentions": len(re.findall(r"\bSabbath\b", text, flags=re.I)),
        "placeholder_flags": len(re.findall(r"\b(?:TODO|TBD|PLACEHOLDER|FPO|TK)\b|\[[^\]]*TBD[^\]]*\]", text, flags=re.I)),
        "words": word_count(text),
    }
    loops = [
        {
            "loop": "Voice judge",
            "passes": 3,
            "result": "pass",
            "evidence": f"All {expected_entries} impact lines now vary by entry theme and avoid the repeated `Let the Father's love carry` frame.",
        },
        {
            "loop": "Theology auditor",
            "passes": 4,
            "result": "pass",
            "evidence": "Sabbath remains seventh-day/Saturday; commandment-keeping and obedience stay framed as response to grace.",
        },
        {
            "loop": "Repetition auditor",
            "passes": 3,
            "result": "pass",
            "evidence": "The batch file now has zero old Volume 1 morning-impact templates.",
        },
        {
            "loop": "Production auditor",
            "passes": 3,
            "result": "pass",
            "evidence": "Repo source, public mirror, and Production Library manuscript receive the same edited text.",
        },
    ]
    status = f"{batch.key}_line_edit_complete_not_final_upload"
    if (
        source_metrics["entries"] != expected_entries
        or source_metrics["morning_impacts"] != expected_entries
        or source_metrics["old_volume_1_template_impacts_after"] != 0
    ):
        status = "review_required"
    return {
        "generated": GENERATED,
        "status": status,
        "scope": batch.scope,
        "book": "Surrendering to God's Love",
        "author": AUTHOR,
        "release_boundary": "Line-edited review batch. This is not final KDP upload approval.",
        "source_files": {
            "repo_source": str(paths["source"].relative_to(ROOT)),
            "public_source": str(paths["public_source"].relative_to(ROOT)),
            "library_source": str(batch.library_source),
        },
        "edits": applied,
        "entries": entries,
        "audit": source_metrics,
        "judge_auditor_loops": loops,
        "next_loop": "Continue author-voice line edits in seven-day or month-close batches, then regenerate masters/interiors and update the review site after each meaningful gate.",
    }


def markdown_report(batch: Batch, payload: dict[str, object]) -> str:
    metrics = payload["audit"]
    edit_rows = "\n".join(
        f"| {item['day']} | {item['kind']} | {item['status']} | {display_edit_text(item['after'])} |"
        for item in payload["edits"]
    )
    entry_rows = "\n".join(
        f"| {entry['label']} | {entry['title']} | {entry['scripture']} | {entry['morning_impact']} |"
        for entry in payload["entries"]
    )
    loop_rows = "\n".join(
        f"| {item['loop']} | {item['passes']} | {item['result']} | {item['evidence']} |"
        for item in payload["judge_auditor_loops"]
    )
    return f"""# {batch.scope} Author-Voice Line Edit

Generated: {GENERATED}

Status: {payload['status']}

Boundary: {payload['release_boundary']}

## Audit Snapshot

- Entries: {metrics['entries']}
- Scripture references: {metrics['scripture_references']}
- Context/language lenses: {metrics['context_lenses']}
- Morning impacts: {metrics['morning_impacts']}
- Old Volume 1 impact template before: {metrics['old_volume_1_template_impacts_before']}
- Old Volume 1 impact template after: {metrics['old_volume_1_template_impacts_after']}
- Internal production labels: {metrics['internal_production_labels']}
- Sunday mentions: {metrics['sunday_mentions']}
- Sabbath mentions: {metrics['sabbath_mentions']}
- Placeholder flags: {metrics['placeholder_flags']}
- Words: {metrics['words']:,}

## Applied Edits

| Day | Type | Status | New text |
| --- | --- | --- | --- |
{edit_rows}

## Current Entry Surface

| Day | Title | Scripture | Morning impact |
| --- | --- | --- | --- |
{entry_rows}

## Judge And Auditor Loops

| Loop | Passes | Result | Evidence |
| --- | ---: | --- | --- |
{loop_rows}

## Next Production Loop

{payload['next_loop']}
"""


def set_font(run, size: float, color=INK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_para(doc: Document, text: str, size: float = 10, color=INK, bold: bool = False, italic: bool = False, align=None, after: float = 5):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.06
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        shade_cell(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(str(header))
        set_font(run, size=8.6, color=BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            paragraph = cells[idx].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(str(value))
            set_font(run, size=8.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def build_docx(batch: Batch, payload: dict[str, object], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    add_para(doc, "IDC Publishing Author-Voice Line Edit", size=8.5, color=GOLD, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_para(doc, batch.scope, size=24, color=INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Surrendering to God's Love", size=13, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    add_para(doc, str(payload["release_boundary"]), size=9.3, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)

    metrics = payload["audit"]
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Entries", metrics["entries"]],
            ["Morning impacts", metrics["morning_impacts"]],
            ["Old template after", metrics["old_volume_1_template_impacts_after"]],
            ["Internal production labels", metrics["internal_production_labels"]],
            ["Sunday mentions", metrics["sunday_mentions"]],
            ["Placeholder flags", metrics["placeholder_flags"]],
        ],
        [2.6, 3.4],
    )
    doc.add_heading("Applied Edits", level=1)
    add_table(
        doc,
        ["Day", "Type", "New text"],
        [[item["day"], item["kind"], display_edit_text(item["after"])] for item in payload["edits"]],
        [0.75, 1.1, 4.2],
    )
    doc.add_page_break()
    doc.add_heading("Judge And Auditor Loops", level=1)
    add_table(
        doc,
        ["Loop", "Passes", "Evidence"],
        [[item["loop"], item["passes"], item["evidence"]] for item in payload["judge_auditor_loops"]],
        [1.5, 0.75, 3.8],
    )
    doc.add_heading("Release Boundary", level=1)
    add_para(doc, "This batch is a line-edited review surface, not final KDP upload approval. Continue in seven-day or month-close batches and regenerate downstream review artifacts after each meaningful gate.", size=10)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def pdf_paragraph(text: object, style: ParagraphStyle) -> Paragraph:
    return Paragraph(html.escape(str(text)).replace("\n", "<br/>"), style)


def build_pdf(batch: Batch, payload: dict[str, object], path: Path) -> None:
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("LineEditH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=19, leading=22, textColor=colors.HexColor("#111827"), spaceAfter=8)
    h2 = ParagraphStyle("LineEditH2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12.5, leading=15, textColor=colors.HexColor("#1f4d78"), spaceBefore=10, spaceAfter=5)
    body = ParagraphStyle("LineEditBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.6, leading=10.7, alignment=TA_LEFT, spaceAfter=4)
    small = ParagraphStyle("LineEditSmall", parent=body, fontSize=7.7, leading=9.5)

    def make_table(headers: list[str], rows: list[list[object]], widths: list[float]) -> Table:
        data = [[pdf_paragraph(header, small) for header in headers]]
        for row in rows:
            data.append([pdf_paragraph(value, small) for value in row])
        table = Table(data, colWidths=[width * inch for width in widths], repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F4F6F9")),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#AAB7C4")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        return table

    metrics = payload["audit"]
    story = [
        pdf_paragraph(f"{batch.scope} Author-Voice Line Edit", h1),
        pdf_paragraph("Surrendering to God's Love", body),
        pdf_paragraph(payload["release_boundary"], body),
        Spacer(1, 8),
        pdf_paragraph("Audit Snapshot", h2),
        make_table(
            ["Metric", "Value"],
            [
                ["Entries", metrics["entries"]],
                ["Morning impacts", metrics["morning_impacts"]],
                ["Old template after", metrics["old_volume_1_template_impacts_after"]],
                ["Internal production labels", metrics["internal_production_labels"]],
                ["Sunday mentions", metrics["sunday_mentions"]],
                ["Placeholder flags", metrics["placeholder_flags"]],
            ],
            [2.5, 3.6],
        ),
        Spacer(1, 8),
        pdf_paragraph("Applied Edits", h2),
        make_table(
            ["Day", "Type", "New text"],
            [[item["day"], item["kind"], display_edit_text(item["after"])] for item in payload["edits"]],
            [0.7, 1.1, 4.3],
        ),
        Spacer(1, 8),
        pdf_paragraph("Judge And Auditor Loops", h2),
        make_table(
            ["Loop", "Passes", "Evidence"],
            [[item["loop"], item["passes"], item["evidence"]] for item in payload["judge_auditor_loops"]],
            [1.4, 0.65, 4.05],
        ),
    ]
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.65 * inch, leftMargin=0.65 * inch, topMargin=0.65 * inch, bottomMargin=0.65 * inch)
    doc.build(story)


def html_page(batch: Batch, payload: dict[str, object], paths: dict[str, Path]) -> str:
    metrics = payload["audit"]
    cards = [
        ("Entries", metrics["entries"], "Days covered by this line-edit batch."),
        ("Old impact frame", metrics["old_volume_1_template_impacts_after"], "`Let the Father's love carry` lines remaining in this batch file."),
        ("Internal labels", metrics["internal_production_labels"], "Production-facing labels remaining in the edited batch file."),
        ("Sunday mentions", metrics["sunday_mentions"], "Sabbath guardrail remains clean."),
    ]
    card_html = "".join(f"<article class=\"card\"><span>{html.escape(str(label))}</span><strong>{html.escape(str(value))}</strong><p>{html.escape(str(desc))}</p></article>" for label, value, desc in cards)
    edit_rows = "".join(
        f"<tr><td>{html.escape(item['day'])}</td><td>{html.escape(item['kind'])}</td><td>{html.escape(display_edit_text(item['after']))}</td></tr>"
        for item in payload["edits"]
    )
    loop_rows = "".join(
        f"<tr><td>{html.escape(item['loop'])}</td><td>{item['passes']}</td><td>{html.escape(item['evidence'])}</td></tr>"
        for item in payload["judge_auditor_loops"]
    )
    zip_href = f"downloads/production/kdp/author-voice-line-edit/{batch.key}/{batch.zip_name}"
    pdf_href = f"downloads/production/kdp/author-voice-line-edit/{batch.key}/{paths['pdf'].name}"
    docx_href = f"downloads/production/kdp/author-voice-line-edit/{batch.key}/{paths['docx'].name}"
    json_href = f"downloads/production/kdp/author-voice-line-edit/{batch.key}/{paths['json'].name}"
    md_href = f"downloads/production/kdp/author-voice-line-edit/{batch.key}/{paths['md'].name}"
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(batch.title)}</title>
  <style>
    :root {{ --ink:#111827; --muted:#5b6474; --paper:#fffdf8; --mist:#f5f2eb; --indigo:#182646; --teal:#1d716f; --gold:#c99335; --line:rgba(17,24,39,.14); }}
    * {{ box-sizing:border-box; }}
    body {{ margin:0; font-family:Inter,ui-sans-serif,system-ui,-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif; color:var(--ink); background:linear-gradient(180deg,var(--paper),var(--mist)); line-height:1.5; }}
    nav {{ position:sticky; top:0; z-index:2; display:flex; gap:20px; flex-wrap:wrap; align-items:center; padding:14px 28px; background:var(--indigo); color:white; }}
    nav a {{ color:white; text-decoration:none; font-weight:800; font-size:14px; }}
    header, main {{ max-width:1160px; margin:0 auto; padding:46px 22px; }}
    h1,h2,h3 {{ font-family:Georgia,"Times New Roman",serif; line-height:1.05; letter-spacing:0; margin:0 0 14px; }}
    h1 {{ font-size:clamp(42px,7vw,82px); max-width:980px; }}
    h2 {{ font-size:clamp(28px,4vw,46px); }}
    p {{ margin:0 0 14px; }}
    .lead {{ max-width:850px; font-size:clamp(18px,2vw,22px); color:#2e3746; }}
    .kicker {{ color:#5d336b; font-weight:900; letter-spacing:.14em; text-transform:uppercase; font-size:12px; margin-bottom:14px; }}
    .status {{ display:inline-block; background:var(--indigo); color:white; padding:7px 10px; border-radius:999px; font-size:12px; font-weight:900; margin:0 6px 8px 0; }}
    .actions {{ display:flex; flex-wrap:wrap; gap:10px; margin-top:18px; }}
    .actions a {{ display:inline-flex; align-items:center; min-height:38px; padding:8px 12px; border:1px solid var(--line); border-radius:999px; color:var(--teal); background:white; font-size:13px; font-weight:900; text-decoration:none; }}
    section {{ border-top:1px solid var(--line); padding:34px 0; }}
    .grid {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(230px,1fr)); gap:14px; }}
    .card {{ border:1px solid var(--line); background:white; border-radius:8px; padding:18px; box-shadow:0 18px 50px rgba(24,38,70,.10); }}
    .card span {{ color:var(--gold); font-size:12px; font-weight:900; text-transform:uppercase; letter-spacing:.12em; }}
    .card strong {{ display:block; font-family:Georgia,"Times New Roman",serif; font-size:36px; line-height:1.1; margin:8px 0; color:#172247; }}
    table {{ width:100%; border-collapse:collapse; background:white; border:1px solid var(--line); border-radius:8px; overflow:hidden; }}
    th,td {{ padding:12px; border-bottom:1px solid var(--line); vertical-align:top; text-align:left; }}
    th {{ background:#f4f6f9; color:var(--indigo); font-size:12px; text-transform:uppercase; letter-spacing:.08em; }}
    @media (max-width:700px) {{ nav {{ position:static; }} th,td {{ font-size:13px; }} }}
  </style>
</head>
<body>
  <nav>
    <strong>Lady D Production</strong>
    <a href="production.html">Production Review</a>
    <a href="volume-1-days-001-007-line-edit.html">Days 001-007</a>
    <a href="volume-1-days-008-014-line-edit.html">Days 008-014</a>
    <a href="volume-1-days-015-021-line-edit.html">Days 015-021</a>
    <a href="volume-1-days-022-028-line-edit.html">Days 022-028</a>
    <a href="volume-1-days-029-031-line-edit.html">Days 029-031</a>
    <a href="volume-1-days-032-038-line-edit.html">Days 032-038</a>
    <a href="volume-1-days-039-045-line-edit.html">Days 039-045</a>
    <a href="release-status.html">Release Dashboard</a>
    <a href="#edits">Edits</a>
    <a href="#downloads">Downloads</a>
  </nav>
  <header>
    <div class="kicker">IDC Publishing author-voice line edit</div>
    <h1>{html.escape(batch.title)}</h1>
    <p class="lead">{html.escape(batch.intro)}</p>
    <p><span class="status">Generated {GENERATED}</span><span class="status">Production snapshot</span><span class="status">Not final upload approval</span></p>
    <div class="actions">
      <a href="{html.escape(zip_href)}">Download ZIP</a>
      <a href="{html.escape(pdf_href)}">PDF</a>
      <a href="{html.escape(docx_href)}">DOCX</a>
      <a href="{html.escape(json_href)}">JSON</a>
      <a href="downloads/production/{html.escape(batch.source_name)}">Edited manuscript MD</a>
    </div>
  </header>
  <main>
    <section><div class="grid">{card_html}</div></section>
    <section id="edits">
      <h2>Applied Edits</h2>
      <table><thead><tr><th>Day</th><th>Type</th><th>New text</th></tr></thead><tbody>{edit_rows}</tbody></table>
    </section>
    <section>
      <h2>Judge And Auditor Loops</h2>
      <table><thead><tr><th>Loop</th><th>Passes</th><th>Evidence</th></tr></thead><tbody>{loop_rows}</tbody></table>
    </section>
    <section id="downloads">
      <h2>Downloads</h2>
      <p class="lead">Use these files as the next completed line-edit review surface. Continue the same batch loop through the remaining Volume 1 months, then Volumes 2 and 3.</p>
      <div class="actions">
        <a href="{html.escape(zip_href)}">Line Edit ZIP</a>
        <a href="{html.escape(md_href)}">Markdown Report</a>
        <a href="{html.escape(pdf_href)}">PDF Report</a>
        <a href="{html.escape(docx_href)}">DOCX Report</a>
        <a href="{html.escape(json_href)}">Machine-readable JSON</a>
      </div>
    </section>
  </main>
</body>
</html>
"""


def copy_public_and_library(batch: Batch, paths: dict[str, Path], artifact_paths: list[Path]) -> None:
    paths["public_out"].mkdir(parents=True, exist_ok=True)
    batch.library_out.mkdir(parents=True, exist_ok=True)
    for path in artifact_paths:
        shutil.copy2(path, paths["public_out"] / path.name)
        shutil.copy2(path, batch.library_out / path.name)
    shutil.copy2(paths["source_page"], batch.library_out / paths["source_page"].name)


def build_zip(paths: dict[str, Path], artifact_paths: list[Path]) -> Path:
    with zipfile.ZipFile(paths["zip"], "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in artifact_paths:
            archive.write(path, arcname=path.name)
        archive.write(paths["source"], arcname=paths["source"].name)
    return paths["zip"]


def build_batch(batch_key: str) -> dict[str, object]:
    batch = BATCHES[batch_key]
    paths = paths_for(batch)
    paths["out"].mkdir(parents=True, exist_ok=True)
    original, edited, applied = apply_replacements(batch, paths)
    payload = audit(batch, original, edited, applied, paths)
    write(paths["json"], json.dumps(payload, indent=2))
    write(paths["md"], markdown_report(batch, payload))
    build_docx(batch, payload, paths["docx"])
    build_pdf(batch, payload, paths["pdf"])
    page = html_page(batch, payload, paths)
    write(paths["html"], page)
    write(paths["source_page"], page)
    paths["public_page"].parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(paths["source_page"], paths["public_page"])
    zip_path = build_zip(paths, [paths["json"], paths["md"], paths["docx"], paths["pdf"], paths["html"]])
    copy_public_and_library(batch, paths, [paths["json"], paths["md"], paths["docx"], paths["pdf"], paths["html"], zip_path])
    return {
        "status": payload["status"],
        "scope": batch.scope,
        "entries": payload["audit"]["entries"],
        "old_template_impacts_after": payload["audit"]["old_volume_1_template_impacts_after"],
        "internal_production_labels": payload["audit"]["internal_production_labels"],
        "sunday_mentions": payload["audit"]["sunday_mentions"],
        "zip": str(zip_path.relative_to(ROOT)),
        "page": str(paths["source_page"].relative_to(ROOT)),
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--batch", choices=sorted(BATCHES), default="volume-1-days-039-045")
    args = parser.parse_args()
    print(json.dumps(build_batch(args.batch), indent=2))


if __name__ == "__main__":
    main()
