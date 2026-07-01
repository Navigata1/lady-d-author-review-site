#!/usr/bin/env python3
"""Build the Lady D proof decision resolution pack."""

from __future__ import annotations

import html
import json
import shutil
import subprocess
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_release_upload_readiness_pack import AUTHOR, PRODUCTS, ROOT, current_commit


GENERATED = "2026-07-01"
AUDIT_PATH = ROOT / "downloads" / "production" / "kdp" / "trilogy-proof-audit" / "trilogy-proof-audit.json"
OUT = ROOT / "downloads" / "production" / "kdp" / "proof-decision-resolution"
PUBLIC_OUT = ROOT / "public" / "downloads" / "production" / "kdp" / "proof-decision-resolution"
SOURCE_PAGE = ROOT / "proof-decision-resolution.html"
PUBLIC_PAGE = ROOT / "public" / "proof-decision-resolution.html"
LIBRARY_OUT = Path("/Users/IDC2.5/Documents/LADY D/Production Library/_Shared/KDP Readiness/Proof Decision Resolution")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(89, 96, 108)
GOLD = RGBColor(122, 90, 0)
LIGHT_FILL = "E8EEF5"
SOFT_FILL = "F4F6F9"

MONTH_LENSES = {
    "January": "belonging, first trust, and beginning the year with received love",
    "February": "courage, confession, and love that keeps choosing the Lord",
    "March": "formation, mercy, and steadiness in ordinary obedience",
    "April": "renewal, surrender, and resurrection-shaped hope",
    "May": "patience, provision, and faithful daily practice",
    "June": "peace, release, and the grace to stop striving",
    "July": "hidden faithfulness, family witness, and quiet endurance",
    "August": "rested witness, Sabbath trust, and obedient availability",
    "September": "harvest, maturity, and fruit that remains",
    "October": "testing, correction, and holy steadiness",
    "November": "gratitude, remembrance, and covenant faithfulness",
    "December": "completion, hope, and carrying the year forward",
    "Bonus": "leap-day grace, holy interruption, and a bonus invitation to receive",
}


def write(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    cleaned = "\n".join(line.rstrip() for line in content.rstrip().splitlines())
    path.write_text(cleaned + "\n", encoding="utf-8")


def load_audit() -> dict[str, object]:
    return json.loads(AUDIT_PATH.read_text(encoding="utf-8"))


def product_title(key: str, audit: dict[str, object]) -> str:
    for product in audit["products"]:
        if product["key"] == key:
            return str(product["title"])
    return key


def product_kind(key: str, audit: dict[str, object]) -> str:
    for product in audit["products"]:
        if product["key"] == key:
            return str(product["kind"])
    return ""


def revision_direction(location: dict[str, object], repeated_value: str) -> str:
    month = str(location.get("month", ""))
    lens = MONTH_LENSES.get(month, "the scripture and monthly arc")
    return f"Retitle this occurrence toward {lens}; preserve the Scripture and devotional angle, but avoid repeating `{repeated_value}`."


def build_title_decisions(audit: dict[str, object]) -> list[dict[str, object]]:
    counters: Counter[int] = Counter()
    decisions: list[dict[str, object]] = []
    for row in audit["title_duplicates"]:
        volume = int(row["volume"])
        counters[volume] += 1
        locations = []
        for idx, location in enumerate(row["locations"]):
            direction = "Potential keeper if this is the strongest contextual use." if idx == 0 else revision_direction(location, str(row["value"]))
            locations.append({**location, "recommended_direction": direction})
        decisions.append(
            {
                "id": f"T-V{volume}-{counters[volume]:03d}",
                "issue_type": "duplicate_day_title",
                "volume": volume,
                "product": row["product"],
                "product_title": row["title"],
                "repeated_title": row["value"],
                "occurrence_count": row["count"],
                "recommended_action": "revise_secondary_occurrences_unless_author_accepts_refrain",
                "decision_prompt": "Should this title act as an intentional refrain, or should later occurrences be individualized for month and Scripture?",
                "recommended_resolution": "Keep at most one strongest occurrence as written, then retitle later occurrences for the local month lens before final upload.",
                "locations": locations,
            }
        )
    return decisions


def build_morning_impact_decisions(audit: dict[str, object]) -> list[dict[str, object]]:
    counters: Counter[int] = Counter()
    decisions: list[dict[str, object]] = []
    for row in audit["morning_impact_duplicates"]:
        volume = int(row["volume"])
        counters[volume] += 1
        month_counts = Counter(str(item.get("month", "")) for item in row["locations"])
        decisions.append(
            {
                "id": f"M-V{volume}-{counters[volume]:03d}",
                "issue_type": "repeated_morning_impact",
                "volume": volume,
                "product": row["product"],
                "product_title": row["title"],
                "repeated_line": row["value"],
                "occurrence_count": row["count"],
                "month_distribution": dict(sorted(month_counts.items())),
                "recommended_action": "choose_intentional_refrain_or_month_specific_revision",
                "decision_prompt": "Should the book keep this repeated morning cue as part of a deliberate devotional rhythm, or should entries receive more varied morning impact lines?",
                "route_a_accept": "Accept as a refrain system if Susan wants a familiar recurring morning cue across the year.",
                "route_b_revise": "Revise repeated uses into month-specific impact lines so the daily close feels more freshly tied to the entry.",
                "locations": row["locations"],
            }
        )
    return decisions


def build_theology_decisions(audit: dict[str, object]) -> list[dict[str, object]]:
    counters: Counter[str] = Counter()
    decisions: list[dict[str, object]] = []
    for item in audit["works_watch_contexts"]:
        if item["preliminary_read"] != "requires proof review":
            continue
        product = str(item["product"])
        counters[product] += 1
        decisions.append(
            {
                "id": f"G-{product}-{counters[product]:02d}",
                "issue_type": "grace_obedience_context_review",
                "volume": next((p.volume for p in PRODUCTS if p.key == product), None),
                "product": product,
                "product_title": product_title(product, audit),
                "kind": product_kind(product, audit),
                "line": item["line"],
                "term": item["term"],
                "snippet": item["snippet"],
                "recommended_action": "theology_review_required",
                "decision_prompt": "Does this context clearly keep obedience as response to grace, never as a way to earn God's love?",
                "recommended_resolution": "Approve only if the nearby language negates earning/performance or frames obedience as grace-shaped response; otherwise revise the sentence before final upload.",
            }
        )
    return decisions


def build_resolution(audit: dict[str, object], commit: str) -> dict[str, object]:
    title_decisions = build_title_decisions(audit)
    impact_decisions = build_morning_impact_decisions(audit)
    theology_decisions = build_theology_decisions(audit)
    all_decisions = title_decisions + impact_decisions + theology_decisions
    by_volume: dict[int, dict[str, int]] = {}
    for volume in (1, 2, 3):
        by_volume[volume] = {
            "title_decisions": sum(1 for item in title_decisions if item["volume"] == volume),
            "morning_impact_decisions": sum(1 for item in impact_decisions if item["volume"] == volume),
            "theology_decisions": sum(1 for item in theology_decisions if item["volume"] == volume),
            "total_decisions": sum(1 for item in all_decisions if item["volume"] == volume),
        }
    likely_aligned = sum(1 for item in audit["works_watch_contexts"] if item["preliminary_read"] == "likely grace-aligned")
    return {
        "generated": GENERATED,
        "commit": commit,
        "source_audit_commit": audit["commit"],
        "author": AUTHOR,
        "status": "decision_queue_ready_not_final_upload",
        "totals": {
            "decision_items": len(all_decisions),
            "title_decisions": len(title_decisions),
            "morning_impact_decisions": len(impact_decisions),
            "theology_decisions": len(theology_decisions),
            "likely_grace_aligned_watch_contexts": likely_aligned,
            "source_words_checked": audit["totals"]["words"],
            "sunday_mentions": audit["totals"]["sunday_mentions"],
            "placeholder_markers": audit["totals"]["placeholders"],
        },
        "volume_breakdown": by_volume,
        "recommended_loop": [
            "Resolve the 12 theology contexts first because they touch the Adventist grace/obedience guardrail.",
            "Decide whether repeated morning-impact lines are an intentional refrain system or should be individualized by month.",
            "Retitle duplicate day-title groups one volume at a time, keeping at most the strongest occurrence when a phrase repeats.",
            "Run a final copyedit pass after decisions are applied to the manuscripts.",
            "Regenerate final interiors and covers only after page counts are locked.",
            "Do not declare KDP final until Previewer and physical proof review pass.",
        ],
        "title_decisions": title_decisions,
        "morning_impact_decisions": impact_decisions,
        "theology_decisions": theology_decisions,
    }


def volume_name(volume: int, audit: dict[str, object]) -> str:
    for product in audit["products"]:
        if product["volume"] == volume and product["kind"] == "Devotional":
            return str(product["title"])
    return f"Volume {volume}"


def main_report_markdown(resolution: dict[str, object], audit: dict[str, object]) -> str:
    totals = resolution["totals"]
    rows = "\n".join(
        f"| Volume {volume} | {volume_name(volume, audit)} | {data['title_decisions']} | {data['morning_impact_decisions']} | {data['theology_decisions']} | {data['total_decisions']} |"
        for volume, data in resolution["volume_breakdown"].items()
    )
    loop = "\n".join(f"{idx}. {item}" for idx, item in enumerate(resolution["recommended_loop"], start=1))
    return f"""# Lady D Proof Decision Resolution Pack

Generated: {GENERATED}

Repo commit at generation: `{resolution['commit']}`

Source proof audit commit: `{resolution['source_audit_commit']}`

Author: {AUTHOR}

Status: Decision queue ready. This is not a final KDP upload declaration.

## Purpose

The trilogy proof audit identified repeated titles, repeated morning-impact lines, and 12 grace/obedience contexts requiring focused proof review. This pack turns those findings into an operator-ready decision queue so the next writing loop can apply changes deliberately instead of guessing.

## Decision Snapshot

- Total decision items: {totals['decision_items']}
- Duplicate-title decisions: {totals['title_decisions']}
- Repeated morning-impact decisions: {totals['morning_impact_decisions']}
- Theology context decisions: {totals['theology_decisions']}
- Likely grace-aligned watch contexts retained for audit trail: {totals['likely_grace_aligned_watch_contexts']}
- Source words checked by upstream audit: {totals['source_words_checked']:,}
- Sunday mentions in checked sources: {totals['sunday_mentions']}
- Placeholder markers in checked sources: {totals['placeholder_markers']}

## Volume Decision Breakdown

| Volume | Title | Title Decisions | Morning-Impact Decisions | Theology Decisions | Total |
| --- | --- | ---: | ---: | ---: | ---: |
{rows}

## Recommended Resolution Loop

{loop}

## Release Boundary

This pack is a bridge from audit evidence into manuscript revision. It does not approve any interior, cover, metadata sheet, or companion journal for final upload. Final KDP readiness still requires applied edits, final proof, Bible permissions decision, ISBN/barcode data, locked page counts, KDP Previewer, and physical proof review.
"""


def title_queue_markdown(decisions: list[dict[str, object]]) -> str:
    lines = [
        "# Title Revision Decision Queue",
        "",
        f"Generated: {GENERATED}",
        "",
        "Recommended rule: keep at most one strongest occurrence of a repeated title, then retitle later occurrences toward month, Scripture, and devotional angle.",
        "",
        "| ID | Volume | Repeated title | Count | Recommended decision | Locations |",
        "| --- | ---: | --- | ---: | --- | --- |",
    ]
    for decision in decisions:
        locations = "; ".join(
            f"{loc['heading']} line {loc['line']} - {loc['recommended_direction']}".replace("|", "/") for loc in decision["locations"]
        )
        lines.append(
            f"| {decision['id']} | {decision['volume']} | {decision['repeated_title']} | {decision['occurrence_count']} | {decision['recommended_action']} | {locations} |"
        )
    return "\n".join(lines)


def morning_impact_queue_markdown(decisions: list[dict[str, object]]) -> str:
    lines = [
        "# Morning-Impact Resolution Queue",
        "",
        f"Generated: {GENERATED}",
        "",
        "Recommended rule: decide whether each repeated line is an intentional refrain or should be made more entry-specific before final upload.",
        "",
        "| ID | Volume | Repeated line | Count | Decision route | First locations |",
        "| --- | ---: | --- | ---: | --- | --- |",
    ]
    for decision in decisions:
        locations = "; ".join(f"{loc['heading']} line {loc['line']}" for loc in decision["locations"][:10])
        if len(decision["locations"]) > 10:
            locations += f"; +{len(decision['locations']) - 10} more"
        lines.append(
            f"| {decision['id']} | {decision['volume']} | {decision['repeated_line']} | {decision['occurrence_count']} | {decision['recommended_action']} | {locations} |"
        )
    return "\n".join(lines)


def theology_queue_markdown(decisions: list[dict[str, object]]) -> str:
    lines = [
        "# Theological Context Resolution Queue",
        "",
        f"Generated: {GENERATED}",
        "",
        "These are the automated watchlist contexts that did not include a nearby grace-alignment marker. They require a human proof decision before final upload.",
        "",
        "| ID | Product | Line | Term | Recommended decision | Context |",
        "| --- | --- | ---: | --- | --- | --- |",
    ]
    for decision in decisions:
        snippet = str(decision["snippet"]).replace("|", "/")
        lines.append(
            f"| {decision['id']} | {decision['product_title']} {decision['kind']} | {decision['line']} | {decision['term']} | {decision['recommended_action']} | {snippet} |"
        )
    return "\n".join(lines)


def volume_sheet_markdown(resolution: dict[str, object], audit: dict[str, object], volume: int) -> str:
    title_decisions = [item for item in resolution["title_decisions"] if item["volume"] == volume]
    impact_decisions = [item for item in resolution["morning_impact_decisions"] if item["volume"] == volume]
    theology_decisions = [item for item in resolution["theology_decisions"] if item["volume"] == volume]
    data = resolution["volume_breakdown"][volume]
    title = volume_name(volume, audit)
    lines = [
        f"# Volume {volume} Proof Decision Sheet - {title}",
        "",
        f"Generated: {GENERATED}",
        "",
        "## Decision Counts",
        "",
        f"- Title decisions: {data['title_decisions']}",
        f"- Morning-impact decisions: {data['morning_impact_decisions']}",
        f"- Theology decisions: {data['theology_decisions']}",
        f"- Total decisions: {data['total_decisions']}",
        "",
        "## Resolution Order",
        "",
        "1. Resolve theology contexts first.",
        "2. Decide morning-impact refrain versus individualized daily lines.",
        "3. Retitle repeated titles using the month lens and Scripture angle.",
        "4. Re-run proof audit after manuscript updates.",
        "",
        "## Theology Contexts",
        "",
    ]
    if theology_decisions:
        for item in theology_decisions:
            lines.append(f"- {item['id']} line {item['line']} `{item['term']}`: {item['recommended_resolution']}")
    else:
        lines.append("- No priority theology contexts in this volume.")
    lines.extend(["", "## Title Decisions", ""])
    for item in title_decisions[:20]:
        lines.append(f"- {item['id']} `{item['repeated_title']}`: {item['recommended_resolution']}")
    if len(title_decisions) > 20:
        lines.append(f"- {len(title_decisions) - 20} additional title decisions are available in the full title queue.")
    lines.extend(["", "## Morning-Impact Decisions", ""])
    for item in impact_decisions:
        lines.append(f"- {item['id']} `{item['repeated_line']}`: choose Route A accept as refrain or Route B revise by month.")
    return "\n".join(lines)


def set_font(run, size: float | None = None, bold: bool = False, italic: bool = False, color=INK) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def display_action(value: object) -> str:
    labels = {
        "theology_review_required": "Theology review",
        "revise_secondary_occurrences_unless_author_accepts_refrain": "Revise or accept refrain",
        "choose_intentional_refrain_or_month_specific_revision": "Choose refrain or revise",
    }
    return labels.get(str(value), str(value).replace("_", " ").title())


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.text = "Lady D Proof Decision Resolution"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.runs[0], size=8.5, color=MUTED, italic=True)
    footer = section.footer.paragraphs[0]
    footer.text = "Decision queue - not final KDP upload approval"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.runs[0], size=8, color=MUTED)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK_BLUE)]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_para(doc: Document, text: str, *, size=11, color=INK, bold=False, italic=False, align=None, after=6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), size=10.5, color=INK)


def add_volume_table(doc: Document, resolution: dict[str, object], audit: dict[str, object]) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Volume", "Title", "Titles", "Impacts", "Theology", "Total"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font(run, size=8.5, color=DARK_BLUE, bold=True)
    repeat_table_header(table.rows[0])
    for volume, data in resolution["volume_breakdown"].items():
        row = table.add_row().cells
        values = [
            str(volume),
            volume_name(volume, audit),
            str(data["title_decisions"]),
            str(data["morning_impact_decisions"]),
            str(data["theology_decisions"]),
            str(data["total_decisions"]),
        ]
        for idx, value in enumerate(values):
            row[idx].text = value
            set_cell_margins(row[idx])
            for paragraph in row[idx].paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=8.5, color=INK)


def add_theology_table(doc: Document, resolution: dict[str, object]) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["ID", "Product", "Line", "Term", "Proof Decision"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, SOFT_FILL)
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font(run, size=8.5, color=DARK_BLUE, bold=True)
    repeat_table_header(table.rows[0])
    for item in resolution["theology_decisions"]:
        row = table.add_row().cells
        values = [item["id"], item["product_title"], str(item["line"]), item["term"], display_action(item["recommended_action"])]
        for idx, value in enumerate(values):
            row[idx].text = str(value)
            set_cell_margins(row[idx])
            for paragraph in row[idx].paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=8.5, color=INK)


def build_docx(resolution: dict[str, object], audit: dict[str, object]) -> Path:
    docx_path = OUT / "lady-d-proof-decision-resolution-pack.docx"
    doc = Document()
    setup_doc(doc)
    doc.core_properties.title = "Lady D Proof Decision Resolution Pack"
    doc.core_properties.author = "IDC Publishing"
    add_para(doc, "LADY D DEVOTIONAL LIBRARY", size=9, color=GOLD, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Proof Decision Resolution Pack", size=24, color=DARK_BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_para(doc, "Operator queue for title, morning-impact, and theology proof decisions", size=12, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_para(doc, f"Generated {GENERATED} from commit {resolution['commit']}", size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    doc.add_heading("Readiness Boundary", level=1)
    add_para(doc, "This pack converts proof findings into decisions. It does not make any file final upload-ready. Applied manuscript edits, final proof, Bible permissions, ISBN/barcode, KDP Previewer, and physical proof review still remain.", after=10)
    doc.add_heading("Decision Snapshot", level=1)
    totals = resolution["totals"]
    for item in [
        f"Total decision items: {totals['decision_items']}",
        f"Duplicate-title decisions: {totals['title_decisions']}",
        f"Repeated morning-impact decisions: {totals['morning_impact_decisions']}",
        f"Theology context decisions: {totals['theology_decisions']}",
        f"Source words checked by upstream audit: {totals['source_words_checked']:,}",
        f"Sunday mentions: {totals['sunday_mentions']}",
        f"Placeholder markers: {totals['placeholder_markers']}",
    ]:
        add_bullet(doc, item)
    doc.add_heading("Volume Decision Breakdown", level=1)
    add_volume_table(doc, resolution, audit)
    doc.add_heading("Priority Theology Contexts", level=1)
    add_para(doc, "Resolve these first because they touch the grace/obedience guardrail. Approve only if obedience remains response to grace, never a path to earning God's love.", after=6)
    add_theology_table(doc, resolution)
    doc.add_heading("Recommended Loop", level=1)
    for item in resolution["recommended_loop"]:
        add_bullet(doc, item)
    doc.save(docx_path)
    return docx_path


def convert_docx_to_pdf(docx_path: Path) -> Path:
    pdf_path = docx_path.with_suffix(".pdf")
    if pdf_path.exists():
        pdf_path.unlink()
    subprocess.check_call(["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUT), str(docx_path)])
    return pdf_path


def review_html(resolution: dict[str, object], audit: dict[str, object]) -> str:
    totals = resolution["totals"]
    volume_rows = "\n".join(
        f"<tr><td>Volume {volume}</td><td>{html.escape(volume_name(volume, audit))}</td><td>{data['title_decisions']}</td><td>{data['morning_impact_decisions']}</td><td>{data['theology_decisions']}</td><td>{data['total_decisions']}</td></tr>"
        for volume, data in resolution["volume_breakdown"].items()
    )
    theology_rows = "\n".join(
        f"<tr><td>{html.escape(item['id'])}</td><td>{html.escape(item['product_title'])}</td><td>{item['line']}</td><td>{html.escape(str(item['term']))}</td><td>{html.escape(display_action(item['recommended_action']))}</td></tr>"
        for item in resolution["theology_decisions"]
    )
    downloads = [
        ("Download proof decision resolution ZIP", "downloads/production/kdp/proof-decision-resolution/Lady-D-Proof-Decision-Resolution-Pack.zip"),
        ("Resolution report PDF", "downloads/production/kdp/proof-decision-resolution/lady-d-proof-decision-resolution-pack.pdf"),
        ("Resolution report DOCX", "downloads/production/kdp/proof-decision-resolution/lady-d-proof-decision-resolution-pack.docx"),
        ("Resolution register JSON", "downloads/production/kdp/proof-decision-resolution/proof-decision-resolution.json"),
        ("Title revision queue", "downloads/production/kdp/proof-decision-resolution/title-revision-decision-queue.md"),
        ("Morning-impact queue", "downloads/production/kdp/proof-decision-resolution/morning-impact-resolution-queue.md"),
        ("Theological context queue", "downloads/production/kdp/proof-decision-resolution/theological-context-resolution-queue.md"),
        ("Volume 1 decision sheet", "downloads/production/kdp/proof-decision-resolution/volume-1-proof-decision-sheet.md"),
        ("Volume 2 decision sheet", "downloads/production/kdp/proof-decision-resolution/volume-2-proof-decision-sheet.md"),
        ("Volume 3 decision sheet", "downloads/production/kdp/proof-decision-resolution/volume-3-proof-decision-sheet.md"),
    ]
    download_links = "\n".join(f'<a class="card" href="{href}">{html.escape(label)}</a>' for label, href in downloads)
    loop = "\n".join(f"<li>{html.escape(item)}</li>" for item in resolution["recommended_loop"])
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='12' fill='%23182646'/%3E%3Ctext x='50%25' y='55%25' dominant-baseline='middle' text-anchor='middle' font-family='Georgia,serif' font-size='26' fill='%23fffdf8'%3ELD%3C/text%3E%3C/svg%3E">
  <title>Lady D Proof Decision Resolution</title>
  <style>
    :root {{ --ink:#111827; --paper:#fffdf8; --mist:#f5f2eb; --indigo:#182646; --teal:#1d716f; --gold:#c99335; --plum:#56325f; --line:rgba(17,24,39,.14); }}
    * {{ box-sizing:border-box; }}
    body {{ margin:0; color:var(--ink); font-family:Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; background:linear-gradient(180deg,var(--paper),var(--mist)); line-height:1.5; }}
    header, main {{ max-width:1180px; margin:0 auto; padding:34px 22px; }}
    h1, h2, h3 {{ font-family:Georgia, "Times New Roman", serif; line-height:1.05; margin:0 0 14px; letter-spacing:0; overflow-wrap:anywhere; }}
    h1 {{ font-size:clamp(42px, 7vw, 82px); max-width:980px; }}
    h2 {{ font-size:clamp(28px, 4vw, 48px); }}
    p {{ margin:0 0 14px; }}
    .lead {{ font-size:clamp(18px, 2vw, 23px); max-width:860px; color:#2e3746; }}
    .kicker {{ color:var(--teal); font-weight:900; letter-spacing:.14em; text-transform:uppercase; font-size:12px; margin-bottom:14px; }}
    .badges {{ display:flex; flex-wrap:wrap; gap:8px; margin-top:18px; }}
    .badge {{ display:inline-flex; align-items:center; min-height:30px; padding:6px 10px; border-radius:999px; color:white; background:var(--indigo); font-size:12px; font-weight:900; }}
    section {{ border-top:1px solid var(--line); padding:34px 0; }}
    .grid {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(220px,1fr)); gap:14px; }}
    .card {{ display:block; border:1px solid var(--line); border-radius:8px; background:white; padding:18px; box-shadow:0 18px 50px rgba(24,38,70,.1); color:var(--teal); font-weight:900; text-decoration:none; }}
    .metric h3 {{ font-size:28px; color:var(--indigo); }}
    .metric p {{ color:#374151; }}
    table {{ width:100%; border-collapse:collapse; background:white; border:1px solid var(--line); border-radius:8px; overflow:hidden; }}
    th, td {{ padding:10px; border-bottom:1px solid var(--line); text-align:left; vertical-align:top; }}
    th {{ color:var(--indigo); background:#e8eef5; font-size:13px; }}
    td {{ font-size:14px; }}
    ul {{ margin:0; padding-left:22px; }}
    li {{ margin:8px 0; }}
    a {{ color:var(--teal); text-underline-offset:3px; }}
    @media (max-width:700px) {{ header, main {{ padding-inline:16px; }} table {{ display:block; overflow-x:auto; }} }}
  </style>
</head>
<body>
  <header>
    <div class="kicker">Proof decision resolution</div>
    <h1>Lady D proof decision resolution pack</h1>
    <p class="lead">This page turns the proof audit into the next operator queue: theology contexts first, then morning-impact rhythm decisions, then repeated-title revisions one volume at a time.</p>
    <div class="badges">
      <span class="badge">Generated {GENERATED}</span>
      <span class="badge">Commit {html.escape(str(resolution['commit']))}</span>
      <span class="badge">192 decision items</span>
      <span class="badge">Not final upload-ready</span>
    </div>
  </header>
  <main>
    <section>
      <h2>Decision Snapshot</h2>
      <div class="grid">
        <div class="card metric"><h3>{totals['decision_items']}</h3><p>Total proof decisions</p></div>
        <div class="card metric"><h3>{totals['title_decisions']}</h3><p>Duplicate-title decisions</p></div>
        <div class="card metric"><h3>{totals['morning_impact_decisions']}</h3><p>Morning-impact decisions</p></div>
        <div class="card metric"><h3>{totals['theology_decisions']}</h3><p>Theology contexts requiring review</p></div>
        <div class="card metric"><h3>{totals['likely_grace_aligned_watch_contexts']}</h3><p>Likely grace-aligned contexts retained</p></div>
        <div class="card metric"><h3>{totals['sunday_mentions']}</h3><p>Sunday mentions in checked sources</p></div>
      </div>
    </section>
    <section>
      <h2>Downloads</h2>
      <div class="grid">{download_links}</div>
    </section>
    <section>
      <h2>Volume Decision Breakdown</h2>
      <table><thead><tr><th>Volume</th><th>Title</th><th>Title decisions</th><th>Morning-impact decisions</th><th>Theology decisions</th><th>Total</th></tr></thead><tbody>{volume_rows}</tbody></table>
    </section>
    <section>
      <h2>Priority Theology Contexts</h2>
      <p class="lead">These 12 contexts should be resolved first because they determine whether language remains clear that obedience is response to grace, never a way to earn love.</p>
      <table><thead><tr><th>ID</th><th>Product</th><th>Line</th><th>Term</th><th>Recommended action</th></tr></thead><tbody>{theology_rows}</tbody></table>
    </section>
    <section>
      <h2>Recommended Loop</h2>
      <ul>{loop}</ul>
      <p><a href="production.html">Return to production review</a></p>
    </section>
  </main>
</body>
</html>
"""


def make_zip(paths: list[Path]) -> Path:
    zip_path = OUT / "Lady-D-Proof-Decision-Resolution-Pack.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in paths:
            zf.write(path, path.name)
    return zip_path


def sync(paths: list[Path]) -> None:
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)
    for path in paths:
        shutil.copy2(path, PUBLIC_OUT / path.name)
        shutil.copy2(path, LIBRARY_OUT / path.name)


def main() -> None:
    commit = current_commit()
    OUT.mkdir(parents=True, exist_ok=True)
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)
    audit = load_audit()
    resolution = build_resolution(audit, commit)

    generated: list[Path] = []
    resolution_json = OUT / "proof-decision-resolution.json"
    main_md = OUT / "lady-d-proof-decision-resolution-pack.md"
    title_md = OUT / "title-revision-decision-queue.md"
    title_json = OUT / "title-revision-decision-queue.json"
    impact_md = OUT / "morning-impact-resolution-queue.md"
    impact_json = OUT / "morning-impact-resolution-queue.json"
    theology_md = OUT / "theological-context-resolution-queue.md"
    theology_json = OUT / "theological-context-resolution-queue.json"
    review_page = OUT / "lady-d-proof-decision-resolution-review.html"

    write(resolution_json, json.dumps(resolution, indent=2))
    write(main_md, main_report_markdown(resolution, audit))
    write(title_md, title_queue_markdown(resolution["title_decisions"]))
    write(title_json, json.dumps(resolution["title_decisions"], indent=2))
    write(impact_md, morning_impact_queue_markdown(resolution["morning_impact_decisions"]))
    write(impact_json, json.dumps(resolution["morning_impact_decisions"], indent=2))
    write(theology_md, theology_queue_markdown(resolution["theology_decisions"]))
    write(theology_json, json.dumps(resolution["theology_decisions"], indent=2))
    write(review_page, review_html(resolution, audit))
    generated.extend([resolution_json, main_md, title_md, title_json, impact_md, impact_json, theology_md, theology_json, review_page])

    for volume in (1, 2, 3):
        sheet = OUT / f"volume-{volume}-proof-decision-sheet.md"
        write(sheet, volume_sheet_markdown(resolution, audit, volume))
        generated.append(sheet)

    docx_path = build_docx(resolution, audit)
    pdf_path = convert_docx_to_pdf(docx_path)
    generated.extend([docx_path, pdf_path])

    zip_path = make_zip(generated)
    all_paths = generated + [zip_path]
    sync(all_paths)
    shutil.copy2(review_page, SOURCE_PAGE)
    shutil.copy2(review_page, PUBLIC_PAGE)
    print(
        json.dumps(
            {
                "report": str(docx_path.relative_to(ROOT)),
                "pdf": str(pdf_path.relative_to(ROOT)),
                "zip": str(zip_path.relative_to(ROOT)),
                "page": str(SOURCE_PAGE.relative_to(ROOT)),
                "result": resolution["status"],
                "decision_items": resolution["totals"]["decision_items"],
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
