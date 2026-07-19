#!/usr/bin/env python3
"""Build paginated 6x9 review interiors from the transcript-directed reader edition."""

from __future__ import annotations

import json
import re
import shutil
import subprocess
import zipfile
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader, PdfWriter


ROOT = Path(__file__).resolve().parents[1]
SOURCE_ROOT = ROOT / "downloads" / "production" / "revised-reader-edition"
OUT = SOURCE_ROOT / "interiors"
PUBLIC_OUT = ROOT / "public" / "downloads" / "production" / "revised-reader-edition" / "interiors"
AUTHOR = 'Susan "Lady D" Damon'
GENERATED = date.today().isoformat()

INK = RGBColor(35, 33, 31)
MUTED = RGBColor(94, 89, 83)
PAPER = "FBFAF6"
SOFT = "F4F0EA"
WHITE = "FFFFFF"


@dataclass(frozen=True)
class Book:
    volume: int
    title: str
    subtitle: str
    lane: str
    accent: RGBColor
    accent_hex: str


BOOKS = (
    Book(1, "Surrendering to God's Love", "A 365-Day Devotional Journey into the Father's Heart", "God the Father", RGBColor(132, 79, 40), "844F28"),
    Book(2, "Walking with Jesus", "A 365-Day Devotional Journey with the Son", "Jesus the Son", RGBColor(36, 102, 106), "24666A"),
    Book(3, "Filled with the Holy Spirit", "A 365-Day Devotional Journey of Presence, Power, and Fruit", "The Holy Spirit", RGBColor(112, 63, 105), "703F69"),
)

BOOK_ORDER = {
    name: index
    for index, name in enumerate(
        (
            "Genesis", "Exodus", "Leviticus", "Numbers", "Deuteronomy", "Joshua", "Judges", "Ruth",
            "1 Samuel", "2 Samuel", "1 Kings", "2 Kings", "1 Chronicles", "2 Chronicles", "Ezra",
            "Nehemiah", "Esther", "Job", "Psalm", "Proverbs", "Ecclesiastes", "Song of Solomon",
            "Isaiah", "Jeremiah", "Lamentations", "Ezekiel", "Daniel", "Hosea", "Joel", "Amos",
            "Obadiah", "Jonah", "Micah", "Nahum", "Habakkuk", "Zephaniah", "Haggai", "Zechariah",
            "Malachi", "Matthew", "Mark", "Luke", "John", "Acts", "Romans", "1 Corinthians",
            "2 Corinthians", "Galatians", "Ephesians", "Philippians", "Colossians", "1 Thessalonians",
            "2 Thessalonians", "1 Timothy", "2 Timothy", "Titus", "Philemon", "Hebrews", "James",
            "1 Peter", "2 Peter", "1 John", "2 John", "3 John", "Jude", "Revelation",
        )
    )
}


def set_font(run, name: str = "Georgia", size: float | None = None, color: RGBColor = INK, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def paragraph_rule(paragraph, color: str, size: str = "6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_font(run, "Aptos", 8, MUTED)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (fld_char, instr, separate, text, end):
        run._r.append(node)


def setup_doc(doc: Document, book: Book) -> None:
    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4.2)
    normal.paragraph_format.line_spacing = 1.03

    for name, size, before, after in (("Title", 25, 0, 10), ("Heading 1", 19, 12, 8), ("Heading 2", 15, 8, 5), ("Heading 3", 11, 5, 3)):
        style = doc.styles[name]
        style.font.name = "Georgia"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
        style.font.size = Pt(size)
        style.font.color.rgb = book.accent if name != "Heading 3" else INK
        style.font.bold = name != "Heading 3"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(book.title.upper())
    set_font(run, "Aptos", 7.2, MUTED, bold=True)
    paragraph_rule(header, book.accent_hex, "3")
    page_number(section.footer.paragraphs[0])


def add_para(doc: Document, text: str, size: float = 9.5, color: RGBColor = INK, bold: bool = False, italic: bool = False, align=None, after: float = 4.2) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, "Georgia", size, color, bold, italic)


def add_title_page(doc: Document, book: Book, journal: bool = False) -> None:
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run(f"VOLUME {book.volume}  |  {book.lane.upper()}")
    set_font(run, "Aptos", 8.5, book.accent, bold=True)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(book.title)
    subtitle = f"{book.subtitle}\nCompanion Journal" if journal else book.subtitle
    add_para(doc, subtitle, 12, MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule_run = rule.add_run("                                       ")
    paragraph_rule(rule, book.accent_hex, "12")
    set_font(rule_run, "Aptos", 5, book.accent)
    add_para(doc, AUTHOR, 11, INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Reader Edition Review Proof", 8.5, MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def add_front_matter(doc: Document, book: Book, journal: bool = False) -> None:
    doc.add_heading("A Welcome from Lady D", level=1)
    add_para(doc, "Dear reader,", 10)
    add_para(doc, "Bring your real morning to these pages. Bring the questions that followed you into the night, the responsibilities already calling your name, and the places where your faith feels strong or tired. You do not have to perform here. Begin with Scripture. Let God speak before fear, pressure, or another person's opinion gets the first word.", 10)
    add_para(doc, "Read slowly enough to notice what meets your life. Pray honestly. Write what is true. Then carry one faithful response into the day. My prayer is that this journey helps you know God's love more deeply, walk with Jesus more closely, and make room for the Holy Spirit to form a life that is both tender and strong.", 10)
    add_para(doc, "With love,\nLady D", 10, book.accent, italic=True)
    doc.add_page_break()
    doc.add_heading("How to Use This Journal" if journal else "How to Use This Book", level=1)
    steps = (
        ("Return to the full KJV Scripture printed for the day.", "Tell the truth before God in the writing space.", "Complete one practical response without turning it into performance.", "Close with the written prayer or make it your own.")
        if journal
        else ("Read the full KJV Scripture printed at the beginning of the day.", "Receive the devotional as a conversation, not an assignment to perform.", "Use the context note only to clarify the verse; return to Scripture itself as the authority.", "Pray honestly, then complete the fused reflection and response in the companion journal.")
    )
    for number, step in enumerate(steps, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(7)
        n = p.add_run(f"{number}. ")
        set_font(n, "Aptos", 9, book.accent, bold=True)
        set_font(p.add_run(step), "Georgia", 9.7)
    doc.add_heading("Scripture Note", level=2)
    add_para(doc, "Scripture quotations are from the King James Version (KJV), using the standardized 1769 text distributed by eBible.org. Dates are shown without weekdays so the journey can be used in any year. The February 29 reading is bonus material for leap years or any day when an extra place to pause is needed.", 8.8, MUTED)
    doc.add_page_break()


def add_month_divider(doc: Document, book: Book, month: str, theme: dict) -> None:
    for _ in range(5):
        doc.add_paragraph()
    add_para(doc, month.upper(), 9, book.accent, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=7)
    heading = doc.add_paragraph(style="Title")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run(theme["name"])
    add_para(doc, theme["promise"], 11, MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    doc.add_page_break()


def add_label(doc: Document, text: str, book: Book, after: float = 2) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text.upper())
    set_font(run, "Aptos", 7.3, book.accent, bold=True)


def add_quote_box(doc: Document, reference: str, text: str, book: Book, compact: bool = False) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, 100 if compact else 120, 130, 100 if compact else 120, 130)
    shade_cell(cell, SOFT if compact else PAPER)
    ref = cell.paragraphs[0]
    ref.paragraph_format.space_after = Pt(3)
    set_font(ref.add_run(f"{reference}  |  KJV"), "Aptos", 7.6, book.accent, bold=True)
    quote = cell.add_paragraph()
    quote.paragraph_format.space_after = Pt(0)
    quote.paragraph_format.line_spacing = 1.0
    set_font(quote.add_run(text), "Georgia", 8.8 if compact else 9.2, INK, italic=True)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1.5)


def entry_marker(entry: dict) -> str:
    return "BONUS READING" if entry["day_number"] == 0 else f"DAY {entry['day_number']:03d}"


def entry_key(entry: dict) -> str:
    return "bonus" if entry["day_number"] == 0 else f"day-{entry['day_number']:03d}"


def add_devotional_entry(doc: Document, book: Book, entry: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(f"{entry_marker(entry)}  |  {entry['date'].upper()}"), "Aptos", 7.6, book.accent, bold=True)
    title = doc.add_paragraph(style="Heading 1")
    title.paragraph_format.keep_with_next = True
    title.add_run(entry["title"])
    add_quote_box(doc, entry["scripture_reference"], entry["scripture_text"], book)
    if entry["context"]:
        add_label(doc, entry["context_label"], book)
        add_para(doc, entry["context"], 8.3, MUTED, italic=True, after=4)
    if entry.get("scripture_connection_reference") and entry.get("scripture_connection_text"):
        add_quote_box(doc, entry["scripture_connection_reference"], entry["scripture_connection_text"], book, compact=True)
    for paragraph in entry["body"]:
        add_para(doc, paragraph, 9.25, INK, after=3.5)
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(2)
    closing.paragraph_format.space_after = Pt(5)
    paragraph_rule(closing, book.accent_hex, "5")
    set_font(closing.add_run(entry["closing"]), "Georgia", 10, book.accent, bold=True)
    add_label(doc, "Prayer", book)
    add_para(doc, entry["prayer"], 9.1, INK, italic=True, after=4)
    add_label(doc, "Reflect and respond", book)
    add_para(doc, entry["reflection_and_response"], 9.1, INK, bold=True, after=0)
    doc.add_page_break()


def scripture_sort(reference: str) -> tuple[int, int, int]:
    match = re.match(r"(.+?)\s+(\d+):(\d+)", reference)
    if not match:
        return (999, 0, 0)
    book, chapter, verse = match.groups()
    return (BOOK_ORDER.get(book, 999), int(chapter), int(verse))


def add_scripture_index(doc: Document, book: Book, entries: list[dict], page_map: dict[str, int]) -> None:
    doc.add_heading("Scripture Journey Index", level=1)
    add_para(doc, "Primary passages are listed in biblical order. Every listing includes the final reading title, date, day, and interior page.", 8.8, MUTED, after=8)
    for entry in sorted(entries, key=lambda item: scripture_sort(item["scripture_reference"])):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3.1)
        paragraph.paragraph_format.keep_together = True
        set_font(paragraph.add_run(entry["scripture_reference"]), "Aptos", 7.1, book.accent, bold=True)
        set_font(paragraph.add_run(" | "), "Aptos", 7.1, MUTED)
        set_font(paragraph.add_run(entry["title"]), "Aptos", 7.1, INK, bold=True)
        meta = paragraph.add_run()
        meta.add_break()
        day = "Bonus" if entry["day_number"] == 0 else f"Day {entry['day_number']}"
        page = page_map.get(entry_key(entry), "TBD")
        set_font(meta, "Aptos", 6.7, MUTED)
        meta.add_text(f"{entry['date']} | {day} | Page {page}")


def add_journal_entry(doc: Document, book: Book, entry: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(f"{entry_marker(entry)}  |  {entry['date'].upper()}"), "Aptos", 7.6, book.accent, bold=True)
    title = doc.add_paragraph(style="Heading 1")
    title.add_run(entry["title"])
    add_quote_box(doc, entry["scripture_reference"], entry["scripture_text"], book, compact=True)
    add_label(doc, "Observe", book)
    add_para(doc, entry["journal_observe"], 8.2, INK, bold=True, after=2)
    add_label(doc, "Reflect", book)
    add_para(doc, entry["journal_reflect"], 8.2, INK, bold=True, after=2)
    add_label(doc, "Act", book)
    add_para(doc, entry["journal_act"], 8.2, INK, bold=True, after=2)
    add_label(doc, "Prayer starter", book)
    add_para(doc, entry["prayer"], 8.0, MUTED, italic=True, after=2)
    add_label(doc, "Prayer record", book)
    add_para(doc, entry["journal_prayer_record"], 8.0, INK, bold=True, after=2)
    add_label(doc, "Follow-through", book)
    add_para(doc, entry["journal_follow_through"], 8.0, INK, bold=True, after=2)
    add_label(doc, "Write", book)
    for _ in range(9):
        line = doc.add_paragraph()
        line.paragraph_format.space_after = Pt(10)
        run = line.add_run("_" * 78)
        set_font(run, "Aptos", 7.4, MUTED)
    doc.add_page_break()


def add_month_review(doc: Document, book: Book, month: str, theme: dict) -> None:
    doc.add_heading(f"{month} Review", level=1)
    add_para(doc, theme["promise"], 10, book.accent, italic=True, after=10)
    for question in (
        f"Where did {theme['name'].lower()} become real in my life this month?",
        "What truth do I want to carry forward?",
        "What one response is God inviting me to continue?",
    ):
        add_para(doc, question, 9.7, INK, bold=True, after=4)
        for _ in range(4):
            line = doc.add_paragraph()
            line.paragraph_format.space_after = Pt(8)
            run = line.add_run("_" * 78)
            set_font(run, "Aptos", 7.4, MUTED)
    doc.add_page_break()


def load_book(book: Book) -> dict:
    path = SOURCE_ROOT / f"volume-{book.volume}" / f"volume-{book.volume}-reader-edition.json"
    return json.loads(path.read_text(encoding="utf-8"))


def build_devotional_docx(book: Book, payload: dict, page_map: dict[str, int] | None = None) -> Path:
    doc = Document()
    setup_doc(doc, book)
    add_title_page(doc, book)
    add_front_matter(doc, book)
    current_month = ""
    for entry in payload["entries"]:
        month = entry["date"].split()[0]
        if month != current_month:
            current_month = month
            theme = payload["themes"][["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"].index(month)]
            add_month_divider(doc, book, month, theme)
        add_devotional_entry(doc, book, entry)
    add_scripture_index(doc, book, payload["entries"], page_map or {})
    OUT.mkdir(parents=True, exist_ok=True)
    path = OUT / f"volume-{book.volume}-revised-reader-edition-6x9.docx"
    doc.save(path)
    return path


def build_journal_docx(book: Book, payload: dict) -> Path:
    doc = Document()
    setup_doc(doc, book)
    add_title_page(doc, book, journal=True)
    add_front_matter(doc, book, journal=True)
    current_month = ""
    current_theme: dict | None = None
    for entry in payload["entries"]:
        month = entry["date"].split()[0]
        if month != current_month:
            if current_month and current_theme:
                add_month_review(doc, book, current_month, current_theme)
            current_month = month
            current_theme = payload["themes"][["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"].index(month)]
            add_month_divider(doc, book, month, current_theme)
        add_journal_entry(doc, book, entry)
    if current_month and current_theme:
        add_month_review(doc, book, current_month, current_theme)
    OUT.mkdir(parents=True, exist_ok=True)
    path = OUT / f"volume-{book.volume}-revised-companion-journal-6x9.docx"
    doc.save(path)
    return path


def normalize_pdf(path: Path) -> None:
    reader = PdfReader(str(path))
    writer = PdfWriter()
    for page in reader.pages:
        page.mediabox.lower_left = (0, 0)
        page.mediabox.upper_right = (432, 648)
        page.cropbox.lower_left = (0, 0)
        page.cropbox.upper_right = (432, 648)
        page.trimbox.lower_left = (0, 0)
        page.trimbox.upper_right = (432, 648)
        writer.add_page(page)
    with path.open("wb") as handle:
        writer.write(handle)


def convert_pdf(docx: Path) -> Path:
    subprocess.check_call(["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(docx.parent), str(docx)])
    path = docx.with_suffix(".pdf")
    normalize_pdf(path)
    return path


def map_entry_pages(pdf: Path, entries: list[dict]) -> dict[str, int]:
    wanted = {entry_marker(entry): entry_key(entry) for entry in entries}
    found: dict[str, int] = {}
    for page_number, page in enumerate(PdfReader(str(pdf)).pages, start=1):
        text = page.extract_text() or ""
        for marker, key in wanted.items():
            if marker in text and key not in found:
                found[key] = page_number
    return found


def audit_pdf(path: Path) -> dict:
    reader = PdfReader(str(path))
    wrong_size = []
    blank_pages = []
    for number, page in enumerate(reader.pages, start=1):
        width = round(float(page.mediabox.width), 2)
        height = round(float(page.mediabox.height), 2)
        if (width, height) != (432.0, 648.0):
            wrong_size.append({"page": number, "size": [width, height]})
        if not (page.extract_text() or "").strip():
            blank_pages.append(number)
    return {"pages": len(reader.pages), "wrong_size_pages": wrong_size, "blank_pages": blank_pages}


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    report = {"generated": GENERATED, "status": "review-proof", "books": []}
    artifacts: list[Path] = []
    for book in BOOKS:
        payload = load_book(book)
        first_docx = build_devotional_docx(book, payload)
        first_pdf = convert_pdf(first_docx)
        page_map = map_entry_pages(first_pdf, payload["entries"])
        if len(page_map) != 366:
            raise RuntimeError(f"Volume {book.volume}: mapped {len(page_map)} of 366 entry pages")
        devotional_docx = build_devotional_docx(book, payload, page_map)
        devotional_pdf = convert_pdf(devotional_docx)
        journal_docx = build_journal_docx(book, payload)
        journal_pdf = convert_pdf(journal_docx)
        item = {
            "volume": book.volume,
            "title": book.title,
            "entries": len(payload["entries"]),
            "mapped_entry_pages": len(page_map),
            "devotional": audit_pdf(devotional_pdf),
            "journal": audit_pdf(journal_pdf),
            "page_map": page_map,
        }
        report["books"].append(item)
        artifacts.extend((devotional_docx, devotional_pdf, journal_docx, journal_pdf))

    audit_path = OUT / "lady-d-revised-interiors-build-audit.json"
    audit_path.write_text(json.dumps(report, indent=2) + "\n", encoding="utf-8")
    artifacts.append(audit_path)
    zip_path = OUT / "Lady-D-Revised-Reader-Edition-Interiors.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for artifact in artifacts:
            archive.write(artifact, artifact.name)
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    for artifact in (*artifacts, zip_path):
        shutil.copy2(artifact, PUBLIC_OUT / artifact.name)
    print(json.dumps(report, indent=2))


if __name__ == "__main__":
    main()
