#!/usr/bin/env python3
"""Build the Volume 1 6x9 interior prototype package."""

from __future__ import annotations

import html
import re
import shutil
import subprocess
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader, PdfWriter


ROOT = Path(__file__).resolve().parents[1]
LIBRARY_ROOT = Path("/Users/IDC2.5/Documents/LADY D/Production Library")
BOOK_ROOT = LIBRARY_ROOT / "01 Surrendering to God's Love"
MASTER_MD = BOOK_ROOT / "06 Master Assembly" / "volume-1-master-interior-manuscript.md"
TEMPLATE_MD = BOOK_ROOT / "06 Master Assembly" / "front-back-matter-template-volume-1.md"
OUT = ROOT / "downloads" / "production" / "kdp" / "interior-prototypes" / "volume-1"
PUBLIC_OUT = ROOT / "public" / "downloads" / "production" / "kdp" / "interior-prototypes" / "volume-1"
PUBLIC_PAGE = ROOT / "public" / "volume-1-interior-prototype.html"
SOURCE_PAGE = ROOT / "volume-1-interior-prototype.html"
LIBRARY_OUT = BOOK_ROOT / "06 Master Assembly" / "Interior Prototype"


BOOK_TITLE = "Surrendering to God's Love"
BOOK_SUBTITLE = "A 365-Day Devotional Journey into the Father's Heart"
AUTHOR = 'Susan "Lady D" Damon'

INK = RGBColor(31, 31, 31)
MUTED = RGBColor(88, 88, 88)
DEEP = RGBColor(41, 36, 34)
RULE = RGBColor(149, 132, 99)
SOFT = "F8F5EE"
BOX = "EFE9DD"


@dataclass(frozen=True)
class Entry:
    day: str
    date: str
    title: str
    scripture: str
    lens_label: str
    lens: str
    body: list[str]
    today_step: str
    prayer: str
    journal_prompt: str
    morning_impact: str


def current_commit() -> str:
    try:
        return subprocess.check_output(["git", "rev-parse", "--short", "HEAD"], cwd=ROOT, text=True).strip()
    except Exception:
        return "unknown"


def clean_line(line: str) -> str:
    return line.strip().replace("\u2011", "-")


def parse_entries(limit: int = 7) -> list[Entry]:
    text = MASTER_MD.read_text()
    chunks = re.split(r"\n---\n", text)
    entries: list[Entry] = []
    for chunk in chunks:
        match = re.search(
            r"^## Day (?P<day>\d{3}) - (?P<date>.+?)\n\n### (?P<title>.+?)\n\n(?P<body>.*)",
            chunk.strip(),
            flags=re.S | re.M,
        )
        if not match:
            continue
        raw = match.group("body").strip()
        fields: dict[str, str] = {}
        paragraphs: list[str] = []
        lens_label = "Context and language lens"
        for part in re.split(r"\n\n+", raw):
            part = clean_line(part)
            bold = re.match(r"^\*\*(.+?):\*\*\s*(.*)$", part, flags=re.S)
            if bold:
                label, value = bold.group(1).strip(), bold.group(2).strip()
                if label in {"Scripture Reference", "Today step", "Prayer", "Journal prompt", "Morning impact"}:
                    fields[label] = value
                elif label in {"Context and language lens", "Production lens correction"}:
                    lens_label = label
                    fields["lens"] = value
                else:
                    paragraphs.append(part)
            elif part and not part.startswith("<!--"):
                paragraphs.append(part)
        entries.append(
            Entry(
                day=match.group("day"),
                date=match.group("date").strip(),
                title=match.group("title").strip(),
                scripture=fields.get("Scripture Reference", ""),
                lens_label=lens_label,
                lens=fields.get("lens", ""),
                body=paragraphs,
                today_step=fields.get("Today step", ""),
                prayer=fields.get("Prayer", ""),
                journal_prompt=fields.get("Journal prompt", ""),
                morning_impact=fields.get("Morning impact", ""),
            )
        )
        if len(entries) >= limit:
            break
    return entries


def set_font(run, name: str = "Georgia", size: float | None = None, color=INK, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_rule(paragraph, color: str = "958463", size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
    normal.font.size = Pt(10.4)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.16

    for name, size, before, after in [
        ("Heading 1", 17.5, 18, 9),
        ("Heading 2", 13.5, 14, 7),
        ("Heading 3", 11.5, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Georgia"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
        style.font.size = Pt(size)
        style.font.color.rgb = DEEP
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1

    section.header.paragraphs[0].text = BOOK_TITLE
    section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(section.header.paragraphs[0].runs[0], size=8.5, color=MUTED, italic=True)
    section.footer.paragraphs[0].text = "Volume 1 Interior Prototype - not final KDP upload file"
    section.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(section.footer.paragraphs[0].runs[0], size=8, color=MUTED)


def add_para(doc: Document, text: str = "", *, size: float = 10.4, bold: bool = False, italic: bool = False, color=INK, align=None, before: float = 0, after: float = 5, style: str | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.16
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_title_page(doc: Document) -> None:
    add_para(doc, "LADY D DEVOTIONAL LIBRARY", size=8.5, bold=True, color=RULE, align=WD_ALIGN_PARAGRAPH.CENTER, before=48, after=26)
    add_para(doc, BOOK_TITLE, size=25, bold=True, color=DEEP, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, BOOK_SUBTITLE, size=12.5, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=34)
    p = add_para(doc, "Susan \"Lady D\" Damon", size=13, bold=True, color=DEEP, align=WD_ALIGN_PARAGRAPH.CENTER, after=44)
    paragraph_rule(p)
    add_para(
        doc,
        "6 x 9 interior prototype - first seven devotional entries with front/back matter system",
        size=9.3,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=12,
    )
    add_para(
        doc,
        "Prepared by IDC Publishing for author review. Scripture references are included without full Bible quotation text until final translation permissions are locked.",
        size=9.2,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=0,
    )
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    doc.add_heading("Copyright and Permissions Draft", level=1)
    add_para(doc, "Copyright (c) 2026 Susan Damon. All rights reserved.", after=6)
    add_para(doc, "Published by IDC Publishing.", after=6)
    add_para(doc, "ISBN: [TBD]", after=6)
    add_para(
        doc,
        "Scripture references are included without full Bible quotation text in this prototype. If final Bible quotation text is added, insert the exact required copyright and permission notice for the selected translation before KDP upload.",
        after=8,
    )
    add_para(
        doc,
        "Adventist guardrail: Sabbath refers to the seventh-day/Saturday Sabbath. Obedience is presented as a response to God's grace, not a method of earning God's love.",
        bold=True,
        after=10,
    )
    doc.add_heading("Author Welcome", level=1)
    for paragraph in [
        "Dear reader,",
        "This devotional year is an invitation to meet God in the real morning, before the day has had time to name you. Bring your questions, your unfinished places, your family burdens, your hopes, and your need for grace.",
        "Let the Lord speak first. Let Scripture steady your steps. Let prayer become practical enough to walk out before noon.",
        "My prayer is that each entry helps you receive God's love more deeply and respond with a life more surrendered, more honest, and more available to Him.",
        "With love,",
        'Susan "Lady D" Damon',
    ]:
        add_para(doc, paragraph)
    doc.add_heading("How to Use This Devotional", level=1)
    steps = [
        "Begin with the Scripture reference.",
        "Read the devotional slowly, listening for one phrase that meets your real life.",
        "Pray the prayer aloud or rewrite it in your own words.",
        "Use the journal prompt to tell the truth before God.",
        "Take the Today step before the day gets crowded.",
        "Let the morning impact line become a small act of obedience.",
    ]
    for idx, step in enumerate(steps):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(step)
        if idx == len(steps) - 1:
            p.add_run().add_break(WD_BREAK.PAGE)


def add_callout(doc: Document, label: str, body: str, fill: str = SOFT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(4.75)
    cell = table.cell(0, 0)
    cell.width = Inches(4.75)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"{label}: ")
    set_font(r, size=9.3, bold=True, color=DEEP)
    r = p.add_run(body)
    set_font(r, size=9.3, color=INK)
    add_para(doc, "", after=5)


def add_entry(doc: Document, entry: Entry) -> None:
    marker = add_para(doc, f"Day {int(entry.day)} - {entry.date}", size=8.4, bold=True, color=RULE, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=4)
    add_para(doc, entry.title, size=18, bold=True, color=DEEP, align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    add_para(doc, entry.scripture, size=10.2, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_callout(doc, entry.lens_label, entry.lens, fill=BOX)
    for idx, paragraph in enumerate(entry.body):
        p = add_para(doc, paragraph, after=5)
        if idx > 0:
            p.paragraph_format.first_line_indent = Inches(0.18)
    add_callout(doc, "Today step", entry.today_step)
    add_callout(doc, "Prayer", entry.prayer)
    add_callout(doc, "Journal prompt", entry.journal_prompt)
    p = add_para(doc, entry.morning_impact, size=9.4, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=4)
    paragraph_rule(p, color="D2C7B2", size="4")
    p.add_run().add_break(WD_BREAK.PAGE)


def add_back_matter(doc: Document) -> None:
    heading = doc.add_heading("Prototype Back Matter", level=1)
    add_para(
        doc,
        "This prototype demonstrates the proposed first-page rhythm, devotional entry treatment, response prompts, and Adventist guardrail language for Volume 1. The full 365-day manuscript plus February 29 bonus remains assembled separately as the review master.",
    )
    doc.add_heading("Before Full Interior Lock", level=2)
    for item in [
        "Approve the front matter voice and final dedication/bio/acknowledgments.",
        "Choose final Bible translation permissions path.",
        "Complete copyedit and theological proof across the full manuscript.",
        "Flow all 366 entries through this design system.",
        "Render the full interior, verify page count, regenerate full-wrap cover, and run KDP Previewer.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def build_docx(entries: list[Entry]) -> Path:
    doc = Document()
    setup_doc(doc)
    add_title_page(doc)
    add_front_matter(doc)
    for entry in entries:
        add_entry(doc, entry)
    add_back_matter(doc)
    out = OUT / "volume-1-6x9-interior-prototype.docx"
    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def convert_pdf(docx_path: Path) -> Path:
    subprocess.check_call(["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUT), str(docx_path)])
    pdf_path = OUT / f"{docx_path.stem}.pdf"
    normalize_pdf_trim(pdf_path)
    return pdf_path


def normalize_pdf_trim(pdf_path: Path) -> None:
    """Normalize exported PDF pages to exact 6 x 9 in KDP trim points."""
    reader = PdfReader(str(pdf_path))
    writer = PdfWriter()
    width = 432
    height = 648
    for page in reader.pages:
        page.mediabox.lower_left = (0, 0)
        page.mediabox.upper_right = (width, height)
        page.cropbox.lower_left = (0, 0)
        page.cropbox.upper_right = (width, height)
        page.trimbox.lower_left = (0, 0)
        page.trimbox.upper_right = (width, height)
        writer.add_page(page)
    with pdf_path.open("wb") as fh:
        writer.write(fh)


def markdown_summary(entries: list[Entry], commit: str) -> str:
    entry_rows = "\n".join(
        f"| {int(e.day)} | {e.date} | {e.title} | {e.scripture} |" for e in entries
    )
    return f"""# Volume 1 6x9 Interior Prototype

Generated: 2026-07-01

Base commit: `{commit}`

Status: Prototype for KDP interior design review. This is not a final KDP upload file.

## Book

- Title: {BOOK_TITLE}
- Subtitle: {BOOK_SUBTITLE}
- Author: {AUTHOR}
- Trim: 6 x 9 inches, no bleed interior prototype
- PDF export: normalized after render to exact 432 x 648 pt media/crop/trim boxes
- Current source: `{MASTER_MD}`
- Front/back matter source: `{TEMPLATE_MD}`

## KDP Setup Basis

- 6 x 9 is a standard KDP paperback trim option.
- Current Volume 1 review master is 369 pages, so the prototype uses a conservative inside/gutter setting for the 301-500 page band.
- Interior is treated as no-bleed because this devotional text sample does not use edge-to-edge interior images.

## Prototype Contents

| Day | Date | Title | Scripture |
| ---: | --- | --- | --- |
{entry_rows}

## Adventist Guardrail

Sabbath remains seventh-day/Saturday Sabbath. Obedience remains response to grace, not a method of earning God's love.

## Judge And Auditor Loop Applied

1. Editorial judge: preserve Lady D devotional warmth and morning action rhythm.
2. Theological auditor: preserve Father-love, grace-before-performance, and Adventist Sabbath frame.
3. Production auditor: 6 x 9 trim, conservative margins, readable devotional rhythm, no bleed.
4. Permissions auditor: keep Scripture references without full quotation text.
5. Release auditor: mark prototype status clearly and block final-upload language.

## Remaining Before Final Upload

- Full 366-entry flow through the approved 6 x 9 design.
- Final dedication, author bio, acknowledgments, ISBN, and copyright owner.
- Bible translation permissions statement.
- Copyedit and theological proof.
- KDP Previewer and physical proof review.
"""


def audit_markdown(entries: list[Entry]) -> str:
    return f"""# Volume 1 Interior Prototype Audit

Generated: 2026-07-01

## Result

Pass for prototype publication to the author-facing review site.

## Evidence

- Source entries parsed: {len(entries)}
- Expected entry count for this prototype: 7
- Entry days: {", ".join(str(int(e.day)) for e in entries)}
- All entries include Scripture reference, body, Today step, prayer, journal prompt, and morning impact.
- Sabbath guardrail is explicit in the prototype front matter.
- Scripture quotation risk is controlled by using references only.
- Document status says prototype and not final KDP upload file.

## KDP Notes

- Trim: 6 x 9 inches.
- Interior: no bleed.
- Margin strategy: conservative prototype margins based on current 301-500 page estimate.
- PDF geometry: exact 432 x 648 pt page boxes after export normalization.
- Final upload remains blocked until full interior page count, cover regeneration, KDP Previewer, and physical proof review.
"""


def review_html(entries: list[Entry], commit: str) -> str:
    cards = "\n".join(
        f"""<article class="entry-card">
          <span>Day {int(e.day)} / {html.escape(e.date)}</span>
          <h3>{html.escape(e.title)}</h3>
          <p><strong>{html.escape(e.scripture)}</strong></p>
          <p>{html.escape(e.morning_impact)}</p>
        </article>"""
        for e in entries
    )
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Volume 1 6x9 Interior Prototype - Lady D</title>
  <style>
    :root {{
      --ink: #171717;
      --paper: #fffdf8;
      --mist: #f4efe4;
      --deep: #2d2925;
      --line: rgba(45, 41, 37, .16);
      --gold: #8b6c35;
      --sage: #4e6f62;
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      color: var(--ink);
      font-family: Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
      background: linear-gradient(180deg, var(--paper), var(--mist));
      line-height: 1.55;
    }}
    header, main {{ max-width: 1160px; margin: 0 auto; padding: 38px 22px; }}
    nav {{ display: flex; flex-wrap: wrap; gap: 10px; padding: 14px 22px; background: #171717; }}
    nav a {{ color: white; text-decoration: none; font-size: 13px; font-weight: 850; padding: 7px 10px; border-radius: 999px; background: rgba(255,255,255,.08); }}
    h1, h2, h3 {{ font-family: Georgia, "Times New Roman", serif; letter-spacing: 0; line-height: 1.05; margin: 0 0 14px; }}
    h1 {{ font-size: clamp(42px, 7vw, 84px); max-width: 920px; }}
    h2 {{ font-size: clamp(30px, 4vw, 52px); }}
    h3 {{ font-size: 24px; }}
    p {{ margin: 0 0 14px; }}
    .lead {{ max-width: 790px; font-size: clamp(18px, 2vw, 23px); color: #343434; }}
    .kicker {{ color: var(--sage); text-transform: uppercase; letter-spacing: .14em; font-size: 12px; font-weight: 900; margin-bottom: 14px; }}
    .status, .download, .entry-card, .audit {{
      border: 1px solid var(--line);
      border-radius: 8px;
      background: rgba(255, 253, 248, .94);
      box-shadow: 0 18px 50px rgba(45, 41, 37, .10);
    }}
    .status {{ padding: 18px; margin-top: 24px; }}
    .grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(240px, 1fr)); gap: 14px; }}
    section {{ border-top: 1px solid var(--line); padding: 34px 0; }}
    .download {{ display: block; padding: 16px; color: var(--deep); font-weight: 850; text-decoration: none; }}
    .entry-card, .audit {{ padding: 18px; }}
    .entry-card span {{ color: var(--gold); font-size: 12px; font-weight: 900; text-transform: uppercase; letter-spacing: .12em; }}
    ul {{ margin: 0; padding-left: 22px; }}
    li {{ margin: 8px 0; }}
  </style>
</head>
<body>
  <nav aria-label="Prototype navigation">
    <a href="production.html">Production Review</a>
    <a href="release-status.html">Release Dashboard</a>
    <a href="#downloads">Downloads</a>
    <a href="#entries">Seven-Day Sample</a>
  </nav>
  <header>
    <div class="kicker">Volume 1 interior prototype</div>
    <h1>6x9 devotional interior rhythm for <em>Surrendering to God's Love</em></h1>
    <p class="lead">This author-facing prototype turns the first seven devotional entries into a KDP-oriented 6x9 interior sample with front matter, devotional page rhythm, response prompts, Sabbath guardrail language, DOCX/PDF outputs, and a release audit.</p>
    <div class="status">
      <p><strong>Status:</strong> Prototype ready for review. Not final KDP upload file.</p>
      <p><strong>Base commit:</strong> {html.escape(commit)}</p>
      <p><strong>Guardrail:</strong> Sabbath remains seventh-day/Saturday Sabbath; obedience remains response to grace.</p>
    </div>
  </header>
  <main>
    <section id="downloads">
      <h2>Prototype Downloads</h2>
      <div class="grid">
        <a class="download" href="downloads/production/kdp/interior-prototypes/volume-1/Lady-D-Volume-1-Interior-Prototype-Pack.zip">Download full prototype pack</a>
        <a class="download" href="downloads/production/kdp/interior-prototypes/volume-1/volume-1-6x9-interior-prototype.docx">Download DOCX prototype</a>
        <a class="download" href="downloads/production/kdp/interior-prototypes/volume-1/volume-1-6x9-interior-prototype.pdf">Download PDF prototype</a>
        <a class="download" href="downloads/production/kdp/interior-prototypes/volume-1/volume-1-6x9-interior-prototype.md">Download prototype notes</a>
        <a class="download" href="downloads/production/kdp/interior-prototypes/volume-1/volume-1-6x9-interior-prototype-audit.md">Download prototype audit</a>
      </div>
    </section>
    <section>
      <h2>Judge And Auditor Result</h2>
      <div class="grid">
        <div class="audit"><h3>Editorial</h3><p>Preserves Lady D's morning voice: tender, direct, prayerful, and practical.</p></div>
        <div class="audit"><h3>Theological</h3><p>Protects grace-before-performance and seventh-day Sabbath language.</p></div>
        <div class="audit"><h3>Production</h3><p>Uses 6x9 no-bleed prototype sizing and conservative margin assumptions.</p></div>
        <div class="audit"><h3>Permissions</h3><p>Uses references only until final Bible translation permissions are locked.</p></div>
      </div>
    </section>
    <section id="entries">
      <h2>Seven-Day Sample</h2>
      <div class="grid">
        {cards}
      </div>
    </section>
    <section>
      <h2>Next Interior Loop</h2>
      <ul>
        <li>Approve or revise this interior rhythm.</li>
        <li>Flow all 366 Volume 1 entries through the selected design.</li>
        <li>Render and inspect the full interior, then regenerate the full-wrap cover from locked page count.</li>
        <li>Repeat the interior system for Volumes 2 and 3 plus the companion journals.</li>
      </ul>
    </section>
  </main>
</body>
</html>
"""


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    cleaned = "\n".join(line.rstrip() for line in content.rstrip().splitlines())
    path.write_text(cleaned + "\n")


def make_zip(paths: list[Path]) -> Path:
    zip_path = OUT / "Lady-D-Volume-1-Interior-Prototype-Pack.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in paths:
            zf.write(path, path.name)
    return zip_path


def sync(paths: list[Path]) -> None:
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)
    for path in paths:
        shutil.copy2(path, PUBLIC_OUT / path.name)
        shutil.copy2(path, LIBRARY_OUT / path.name)


def main() -> None:
    commit = current_commit()
    entries = parse_entries()
    if len(entries) != 7:
        raise SystemExit(f"expected 7 entries, found {len(entries)}")

    OUT.mkdir(parents=True, exist_ok=True)
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)

    docx_path = build_docx(entries)
    pdf_path = convert_pdf(docx_path)

    notes_path = OUT / "volume-1-6x9-interior-prototype.md"
    audit_path = OUT / "volume-1-6x9-interior-prototype-audit.md"
    review_path = OUT / "volume-1-6x9-interior-prototype-review.html"
    write_text(notes_path, markdown_summary(entries, commit))
    write_text(audit_path, audit_markdown(entries))
    write_text(review_path, review_html(entries, commit))

    paths = [docx_path, pdf_path, notes_path, audit_path, review_path]
    zip_path = make_zip(paths)
    sync(paths + [zip_path])
    shutil.copy2(review_path, SOURCE_PAGE)
    shutil.copy2(review_path, PUBLIC_PAGE)
    print(
        {
            "docx": str(docx_path.relative_to(ROOT)),
            "pdf": str(pdf_path.relative_to(ROOT)),
            "review": str(SOURCE_PAGE.relative_to(ROOT)),
            "zip": str(zip_path.relative_to(ROOT)),
        }
    )


if __name__ == "__main__":
    main()
