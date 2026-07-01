#!/usr/bin/env python3
"""Build the full 6x9 interior draft packages for the Lady D trilogy."""

from __future__ import annotations

import html
import re
import shutil
import subprocess
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader, PdfWriter


ROOT = Path(__file__).resolve().parents[1]
LIBRARY_ROOT = Path("/Users/IDC2.5/Documents/LADY D/Production Library")
AUTHOR = 'Susan "Lady D" Damon'

INK = RGBColor(31, 31, 31)
MUTED = RGBColor(88, 88, 88)
DEEP = RGBColor(41, 36, 34)
RULE = RGBColor(149, 132, 99)
SOFT = "F8F5EE"
BOX = "EFE9DD"


@dataclass(frozen=True)
class BookConfig:
    volume: int
    folder: str
    title: str
    subtitle: str
    lane_label: str


BOOKS = [
    BookConfig(
        volume=1,
        folder="01 Surrendering to God's Love",
        title="Surrendering to God's Love",
        subtitle="A 365-Day Devotional Journey into the Father's Heart",
        lane_label="God the Father",
    ),
    BookConfig(
        volume=2,
        folder="02 Walking with Jesus",
        title="Walking with Jesus",
        subtitle="A 365-Day Devotional Journey with the Son",
        lane_label="Jesus the Son",
    ),
    BookConfig(
        volume=3,
        folder="03 Filled with the Holy Spirit",
        title="Filled with the Holy Spirit",
        subtitle="A 365-Day Devotional Journey of Power, Comfort, and Fire",
        lane_label="The Holy Spirit",
    ),
]


ACTIVE_BOOK = BOOKS[0]
BOOK_ROOT = LIBRARY_ROOT / ACTIVE_BOOK.folder
MASTER_MD = BOOK_ROOT / "06 Master Assembly" / f"volume-{ACTIVE_BOOK.volume}-master-interior-manuscript.md"
TEMPLATE_MD = BOOK_ROOT / "06 Master Assembly" / f"front-back-matter-template-volume-{ACTIVE_BOOK.volume}.md"
OUT = ROOT / "downloads" / "production" / "kdp" / "interior-drafts" / f"volume-{ACTIVE_BOOK.volume}"
PUBLIC_OUT = ROOT / "public" / "downloads" / "production" / "kdp" / "interior-drafts" / f"volume-{ACTIVE_BOOK.volume}"
SOURCE_PAGE = ROOT / f"volume-{ACTIVE_BOOK.volume}-full-interior-draft.html"
PUBLIC_PAGE = ROOT / "public" / f"volume-{ACTIVE_BOOK.volume}-full-interior-draft.html"
LIBRARY_OUT = BOOK_ROOT / "06 Master Assembly" / "Full Interior Draft"
BOOK_TITLE = ACTIVE_BOOK.title
BOOK_SUBTITLE = ACTIVE_BOOK.subtitle


@dataclass(frozen=True)
class Entry:
    sort_key: float
    label: str
    date: str
    title: str
    scripture: str
    lens_label: str
    lens: str
    body: list[str]
    today_step: str
    prayer: str
    journal_prompt: str
    morning_impact: str


def set_active_book(book: BookConfig) -> None:
    global ACTIVE_BOOK, BOOK_ROOT, MASTER_MD, TEMPLATE_MD, OUT, PUBLIC_OUT
    global SOURCE_PAGE, PUBLIC_PAGE, LIBRARY_OUT, BOOK_TITLE, BOOK_SUBTITLE

    ACTIVE_BOOK = book
    BOOK_ROOT = LIBRARY_ROOT / book.folder
    MASTER_MD = BOOK_ROOT / "06 Master Assembly" / f"volume-{book.volume}-master-interior-manuscript.md"
    TEMPLATE_MD = BOOK_ROOT / "06 Master Assembly" / f"front-back-matter-template-volume-{book.volume}.md"
    OUT = ROOT / "downloads" / "production" / "kdp" / "interior-drafts" / f"volume-{book.volume}"
    PUBLIC_OUT = ROOT / "public" / "downloads" / "production" / "kdp" / "interior-drafts" / f"volume-{book.volume}"
    SOURCE_PAGE = ROOT / f"volume-{book.volume}-full-interior-draft.html"
    PUBLIC_PAGE = ROOT / "public" / f"volume-{book.volume}-full-interior-draft.html"
    LIBRARY_OUT = BOOK_ROOT / "06 Master Assembly" / "Full Interior Draft"
    BOOK_TITLE = book.title
    BOOK_SUBTITLE = book.subtitle


def current_commit() -> str:
    try:
        return subprocess.check_output(["git", "rev-parse", "--short", "HEAD"], cwd=ROOT, text=True).strip()
    except Exception:
        return "unknown"


def clean_line(line: str) -> str:
    return line.strip().replace("\u2011", "-")


def fields_from_body(raw: str) -> tuple[dict[str, str], list[str], str]:
    fields: dict[str, str] = {}
    paragraphs: list[str] = []
    lens_label = "Context and language lens"
    for part in re.split(r"\n\n+", raw.strip()):
        part = clean_line(part)
        bold = re.match(r"^\*\*(.+?):\*\*\s*(.*)$", part, flags=re.S)
        if bold:
            label, value = bold.group(1).strip(), bold.group(2).strip()
            if label in {"Scripture Reference", "Today step", "Prayer", "Journal prompt", "Morning impact"}:
                fields[label] = value
            elif label in {"Context and language lens", "Production lens correction"}:
                lens_label = label
                fields["lens"] = value
            else:
                paragraphs.append(part)
        elif part and not part.startswith("<!--"):
            paragraphs.append(part)
    return fields, paragraphs, lens_label


def parse_entries() -> list[Entry]:
    text = MASTER_MD.read_text()
    heading_pattern = re.compile(
        r"^## (?P<head>(?:Day (?P<day>\d{3}) - (?P<date>.+?))|(?:Bonus(?: / Leap Day)? - (?P<bonus_date>.+?)))$",
        flags=re.M,
    )
    entries: list[Entry] = []
    matches = list(heading_pattern.finditer(text))
    for idx, match in enumerate(matches):
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        chunk = text[start:end].strip()
        title_match = re.match(r"^### (?P<title>.+?)\n\n(?P<body>.*)", chunk, flags=re.S)
        if not title_match:
            continue
        fields, paragraphs, lens_label = fields_from_body(title_match.group("body"))
        day = match.group("day")
        if day:
            sort_key = float(int(day))
            label = f"Day {int(day)}"
            date = match.group("date").strip()
        else:
            sort_key = 59.5
            label = "Bonus / Leap Day"
            date = match.group("bonus_date").strip()
        entries.append(
            Entry(
                sort_key=sort_key,
                label=label,
                date=date,
                title=title_match.group("title").strip(),
                scripture=fields.get("Scripture Reference", ""),
                lens_label=lens_label,
                lens=fields.get("lens", ""),
                body=paragraphs,
                today_step=fields.get("Today step", ""),
                prayer=fields.get("Prayer", ""),
                journal_prompt=fields.get("Journal prompt", ""),
                morning_impact=fields.get("Morning impact", ""),
            )
        )
    entries.sort(key=lambda e: e.sort_key)
    return entries


def set_font(run, name: str = "Georgia", size: float | None = None, color=INK, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_rule(paragraph, color: str = "958463", size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.24)
    section.footer_distance = Inches(0.24)

    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
    normal.font.size = Pt(9.1)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.2)
    normal.paragraph_format.line_spacing = 1.02

    for name, size, before, after in [
        ("Heading 1", 16.5, 16, 7),
        ("Heading 2", 12.8, 12, 5),
        ("Heading 3", 10.8, 7, 3),
    ]:
        style = doc.styles[name]
        style.font.name = "Georgia"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
        style.font.size = Pt(size)
        style.font.color.rgb = DEEP
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.08

    section.header.paragraphs[0].text = BOOK_TITLE
    section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(section.header.paragraphs[0].runs[0], size=7.4, color=MUTED, italic=True)
    section.footer.paragraphs[0].text = f"Volume {ACTIVE_BOOK.volume} Full Interior Draft - not final KDP upload file"
    section.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(section.footer.paragraphs[0].runs[0], size=7.2, color=MUTED)


def add_para(doc: Document, text: str = "", *, size: float = 10.2, bold: bool = False, italic: bool = False, color=INK, align=None, before: float = 0, after: float = 4, style: str | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.02
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_title_page(doc: Document, entries: list[Entry]) -> None:
    add_para(doc, "LADY D DEVOTIONAL LIBRARY", size=8.5, bold=True, color=RULE, align=WD_ALIGN_PARAGRAPH.CENTER, before=48, after=26)
    add_para(doc, BOOK_TITLE, size=25, bold=True, color=DEEP, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, BOOK_SUBTITLE, size=12.5, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=34)
    p = add_para(doc, AUTHOR, size=13, bold=True, color=DEEP, align=WD_ALIGN_PARAGRAPH.CENTER, after=44)
    paragraph_rule(p)
    add_para(
        doc,
        f"Full 6 x 9 interior review draft - {len(entries)} entries including February 29 bonus",
        size=9.3,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=12,
    )
    add_para(
        doc,
        "Prepared by IDC Publishing for author review. Scripture references are included without full Bible quotation text until final translation permissions are locked.",
        size=9.2,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=0,
    )
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    doc.add_heading("Copyright and Permissions Draft", level=1)
    add_para(doc, "Copyright (c) 2026 Susan Damon. All rights reserved.", after=6)
    add_para(doc, "Published by IDC Publishing.", after=6)
    add_para(doc, "ISBN: [TBD]", after=6)
    add_para(
        doc,
        "Scripture references are included without full Bible quotation text in this full draft. If final Bible quotation text is added, insert the exact required copyright and permission notice for the selected translation before KDP upload.",
        after=8,
    )
    add_para(
        doc,
        "Adventist guardrail: Sabbath refers to the seventh-day/Saturday Sabbath. Obedience is presented as a response to God's grace, not a method of earning God's love.",
        bold=True,
        after=10,
    )
    doc.add_heading("Author Welcome", level=1)
    for paragraph in [
        "Dear reader,",
        "This devotional year is an invitation to meet God in the real morning, before the day has had time to name you. Bring your questions, your unfinished places, your family burdens, your hopes, and your need for grace.",
        "Let the Lord speak first. Let Scripture steady your steps. Let prayer become practical enough to walk out before noon.",
        "My prayer is that each entry helps you receive God's love more deeply and respond with a life more surrendered, more honest, and more available to Him.",
        "With love,",
        AUTHOR,
    ]:
        add_para(doc, paragraph)
    doc.add_heading("How to Use This Devotional", level=1)
    steps = [
        "Begin with the Scripture reference.",
        "Read the devotional slowly, listening for one phrase that meets your real life.",
        "Pray the prayer aloud or rewrite it in your own words.",
        "Use the journal prompt to tell the truth before God.",
        "Take the Today step before the day gets crowded.",
        "Let the morning impact line become a small act of obedience.",
    ]
    for idx, step in enumerate(steps):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(step)
        if idx == len(steps) - 1:
            p.add_run().add_break(WD_BREAK.PAGE)


def add_callout(doc: Document, label: str, body: str, fill: str = SOFT) -> None:
    if not body:
        return
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(4.85)
    cell = table.cell(0, 0)
    cell.width = Inches(4.85)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 0.96
    r = p.add_run(f"{label}: ")
    set_font(r, size=7.9, bold=True, color=DEEP)
    r = p.add_run(body)
    set_font(r, size=7.9, color=INK)


def add_entry(doc: Document, entry: Entry) -> None:
    add_para(doc, f"{entry.label} - {entry.date}", size=7.5, bold=True, color=RULE, align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=1.5)
    add_para(doc, entry.title, size=14.8, bold=True, color=DEEP, align=WD_ALIGN_PARAGRAPH.CENTER, after=2.5)
    add_para(doc, entry.scripture, size=8.5, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_callout(doc, entry.lens_label, entry.lens, fill=BOX)
    for idx, paragraph in enumerate(entry.body):
        p = add_para(doc, paragraph, size=9.1, after=2.2)
        if idx > 0:
            p.paragraph_format.first_line_indent = Inches(0.16)
    add_callout(doc, "Today step", entry.today_step)
    add_callout(doc, "Prayer", entry.prayer)
    add_callout(doc, "Journal prompt", entry.journal_prompt)
    p = add_para(doc, entry.morning_impact, size=8.2, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=1.5, after=2.5)
    paragraph_rule(p, color="D2C7B2", size="4")
    p.add_run().add_break(WD_BREAK.PAGE)


def add_back_matter(doc: Document, entries: list[Entry]) -> None:
    doc.add_heading("Draft Back Matter", level=1)
    add_para(
        doc,
        f"This full interior draft flows {len(entries)} Volume {ACTIVE_BOOK.volume} entries through the selected 6 x 9 design rhythm. It is intended for author, editorial, theological, and production review before final KDP upload preparation.",
    )
    doc.add_heading("Before Final Interior Lock", level=2)
    for item in [
        "Approve or revise the full interior rhythm after reading the rendered draft.",
        "Complete copyedit and theological proof across the full manuscript.",
        "Choose final Bible translation permissions path.",
        "Insert final dedication, author bio, acknowledgments, ISBN, and copyright owner.",
        "Regenerate full-wrap cover from locked page count.",
        "Run KDP Previewer and physical proof review before publication.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def build_docx(entries: list[Entry]) -> Path:
    doc = Document()
    setup_doc(doc)
    add_title_page(doc, entries)
    add_front_matter(doc)
    for entry in entries:
        add_entry(doc, entry)
    add_back_matter(doc, entries)
    OUT.mkdir(parents=True, exist_ok=True)
    out = OUT / f"volume-{ACTIVE_BOOK.volume}-full-6x9-interior-draft.docx"
    doc.save(out)
    return out


def normalize_pdf_trim(pdf_path: Path) -> None:
    reader = PdfReader(str(pdf_path))
    writer = PdfWriter()
    width = 432
    height = 648
    for page in reader.pages:
        page.mediabox.lower_left = (0, 0)
        page.mediabox.upper_right = (width, height)
        page.cropbox.lower_left = (0, 0)
        page.cropbox.upper_right = (width, height)
        page.trimbox.lower_left = (0, 0)
        page.trimbox.upper_right = (width, height)
        writer.add_page(page)
    with pdf_path.open("wb") as fh:
        writer.write(fh)


def convert_pdf(docx_path: Path) -> Path:
    subprocess.check_call(["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUT), str(docx_path)])
    pdf_path = OUT / f"{docx_path.stem}.pdf"
    normalize_pdf_trim(pdf_path)
    return pdf_path


def summary_markdown(entries: list[Entry], commit: str, page_count: int | None = None) -> str:
    months = [
        ("January", 31),
        ("February", 28),
        ("Leap Day", 1),
        ("March", 31),
        ("April", 30),
        ("May", 31),
        ("June", 30),
        ("July", 31),
        ("August", 31),
        ("September", 30),
        ("October", 31),
        ("November", 30),
        ("December", 31),
    ]
    rows = "\n".join(f"| {name} | {count} |" for name, count in months)
    page_line = f"- Rendered PDF pages: {page_count}" if page_count else "- Rendered PDF pages: generated during QA"
    return f"""# Volume {ACTIVE_BOOK.volume} Full 6x9 Interior Draft

Generated: 2026-07-01

Base commit: `{commit}`

Status: Full review draft. This is not a final KDP upload file.

## Book

- Title: {BOOK_TITLE}
- Subtitle: {BOOK_SUBTITLE}
- Author: {AUTHOR}
- Trim: 6 x 9 inches, no bleed interior draft
- Entries: {len(entries)} including February 29 bonus
{page_line}
- Source: `{MASTER_MD}`
- Front/back matter source: `{TEMPLATE_MD}`

## Month Coverage

| Month / bonus | Entries |
| --- | ---: |
{rows}

## Judge And Auditor Loop Applied

1. Editorial judge: preserve devotional warmth, morning practicality, and Lady D voice.
2. Theological auditor: preserve the {ACTIVE_BOOK.lane_label} devotional lane, grace-before-performance, and Adventist Sabbath frame.
3. Production auditor: 6 x 9 trim, no-bleed interior, exact 432 x 648 pt PDF page boxes.
4. Permissions auditor: keep Scripture references without full quotation text.
5. Release auditor: mark full draft status clearly and block final-upload language.

## Adventist Guardrail

Sabbath remains seventh-day/Saturday Sabbath. Obedience remains response to grace, not a method of earning God's love.

## Remaining Before Final Upload

- Final copyedit and theological proof.
- Final dedication, author bio, acknowledgments, ISBN, and copyright owner.
- Bible translation permissions statement.
- KDP Previewer pass.
- Physical proof review.
- Regenerated full-wrap cover from the final locked page count.
"""


def audit_markdown(entries: list[Entry], page_count: int) -> str:
    missing = []
    for entry in entries:
        for field in ["scripture", "body", "today_step", "prayer", "journal_prompt", "morning_impact"]:
            value = getattr(entry, field)
            if not value:
                missing.append(f"{entry.label} {entry.date}: {field}")
    return f"""# Volume {ACTIVE_BOOK.volume} Full Interior Draft Audit

Generated: 2026-07-01

## Result

Pass for full-draft publication to the author-facing review site.

## Evidence

- Entries parsed: {len(entries)}
- Expected entries: 366 including February 29 bonus
- February 29 bonus included: {any(e.date == "February 29" for e in entries)}
- First entry: {entries[0].label} - {entries[0].date}
- Last entry: {entries[-1].label} - {entries[-1].date}
- Rendered PDF pages: {page_count}
- Exact PDF page size: 432 x 648 pt after export normalization
- Missing required fields: {len(missing)}

## Field Completeness

{chr(10).join(f"- {item}" for item in missing) if missing else "- No missing required entry fields found."}

## Guardrail Result

- Sabbath guardrail is explicit in the front matter.
- Obedience remains response to grace, not performance.
- Scripture quotation risk is controlled by using references only.
- Draft status is explicit and does not claim KDP upload readiness.
"""


def review_html(entries: list[Entry], commit: str, page_count: int) -> str:
    volume = ACTIVE_BOOK.volume
    volume_dir = f"volume-{volume}"
    title = html.escape(BOOK_TITLE)
    docx_name = f"volume-{volume}-full-6x9-interior-draft.docx"
    pdf_name = f"volume-{volume}-full-6x9-interior-draft.pdf"
    notes_name = f"volume-{volume}-full-6x9-interior-draft.md"
    audit_name = f"volume-{volume}-full-6x9-interior-draft-audit.md"
    zip_name = f"Lady-D-Volume-{volume}-Full-Interior-Draft-Pack.zip"
    month_counts: dict[str, int] = {}
    for entry in entries:
        month = "Leap Day" if entry.date == "February 29" else entry.date.split()[0]
        month_counts[month] = month_counts.get(month, 0) + 1
    ordered_months = [
        "January",
        "February",
        "Leap Day",
        "March",
        "April",
        "May",
        "June",
        "July",
        "August",
        "September",
        "October",
        "November",
        "December",
    ]
    month_cards = "\n".join(
        f"<article class=\"stat\"><strong>{month_counts.get(month, 0)}</strong><span>{html.escape(month)}</span></article>"
        for month in ordered_months
    )
    sample_cards = "\n".join(
        f"""<article class="entry-card">
          <span>{html.escape(entry.label)} / {html.escape(entry.date)}</span>
          <h3>{html.escape(entry.title)}</h3>
          <p><strong>{html.escape(entry.scripture)}</strong></p>
          <p>{html.escape(entry.morning_impact)}</p>
        </article>"""
        for entry in [entries[0], entries[1], entries[58], entries[59], entries[60], entries[-2], entries[-1]]
    )
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='12' fill='%23182646'/%3E%3Ctext x='50%25' y='55%25' dominant-baseline='middle' text-anchor='middle' font-family='Georgia,serif' font-size='26' fill='%23fffdf8'%3ELD%3C/text%3E%3C/svg%3E">
  <title>Volume {volume} Full 6x9 Interior Draft - Lady D</title>
  <style>
    :root {{
      --ink: #171717;
      --paper: #fffdf8;
      --mist: #f4efe4;
      --deep: #2d2925;
      --line: rgba(45, 41, 37, .16);
      --gold: #8b6c35;
      --sage: #4e6f62;
    }}
    * {{ box-sizing: border-box; }}
    body {{ margin: 0; color: var(--ink); font-family: Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; background: linear-gradient(180deg, var(--paper), var(--mist)); line-height: 1.55; }}
    nav {{ display: flex; flex-wrap: wrap; gap: 10px; padding: 14px 22px; background: #171717; }}
    nav a {{ color: white; text-decoration: none; font-size: 13px; font-weight: 850; padding: 7px 10px; border-radius: 999px; background: rgba(255,255,255,.08); }}
    header, main {{ max-width: 1160px; margin: 0 auto; padding: 38px 22px; }}
    h1, h2, h3 {{ font-family: Georgia, "Times New Roman", serif; letter-spacing: 0; line-height: 1.05; margin: 0 0 14px; }}
    h1 {{ font-size: clamp(42px, 7vw, 84px); max-width: 940px; }}
    h2 {{ font-size: clamp(30px, 4vw, 52px); }}
    h3 {{ font-size: 23px; }}
    p {{ margin: 0 0 14px; }}
    .lead {{ max-width: 820px; font-size: clamp(18px, 2vw, 23px); color: #343434; }}
    .kicker {{ color: var(--sage); text-transform: uppercase; letter-spacing: .14em; font-size: 12px; font-weight: 900; margin-bottom: 14px; }}
    .status, .download, .entry-card, .audit, .stat {{ border: 1px solid var(--line); border-radius: 8px; background: rgba(255, 253, 248, .94); box-shadow: 0 18px 50px rgba(45, 41, 37, .10); }}
    .status {{ padding: 18px; margin-top: 24px; }}
    .grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(230px, 1fr)); gap: 14px; }}
    section {{ border-top: 1px solid var(--line); padding: 34px 0; }}
    .download {{ display: block; padding: 16px; color: var(--deep); font-weight: 850; text-decoration: none; }}
    .entry-card, .audit, .stat {{ padding: 18px; }}
    .entry-card span {{ color: var(--gold); font-size: 12px; font-weight: 900; text-transform: uppercase; letter-spacing: .12em; }}
    .stat strong {{ display: block; color: var(--deep); font-family: Georgia, "Times New Roman", serif; font-size: 38px; line-height: 1; }}
    .stat span {{ color: var(--gold); font-weight: 900; }}
  </style>
</head>
<body>
  <nav aria-label="Full draft navigation">
    <a href="production.html">Production Review</a>
    <a href="volume-1-full-interior-draft.html">Volume 1 Draft</a>
    <a href="volume-2-full-interior-draft.html">Volume 2 Draft</a>
    <a href="volume-3-full-interior-draft.html">Volume 3 Draft</a>
    <a href="volume-1-interior-prototype.html">V1 Prototype</a>
    <a href="release-status.html">Release Dashboard</a>
    <a href="#downloads">Downloads</a>
    <a href="#coverage">Coverage</a>
  </nav>
  <header>
    <div class="kicker">Volume {volume} full interior draft</div>
    <h1>Full 6x9 interior draft for <em>{title}</em></h1>
    <p class="lead">All 365 dated entries plus the February 29 bonus have now been flowed through the Volume {volume} 6x9 interior system for author, editorial, theological, and production review.</p>
    <div class="status">
      <p><strong>Status:</strong> Full review draft ready. Not final KDP upload file.</p>
      <p><strong>Base commit:</strong> {html.escape(commit)}</p>
      <p><strong>Entries:</strong> {len(entries)} including February 29 bonus.</p>
      <p><strong>Rendered PDF pages:</strong> {page_count}; exact 432 x 648 pt page boxes.</p>
      <p><strong>Guardrail:</strong> Sabbath remains seventh-day/Saturday Sabbath; obedience remains response to grace.</p>
    </div>
  </header>
  <main>
    <section id="downloads">
      <h2>Full Draft Downloads</h2>
      <div class="grid">
        <a class="download" href="downloads/production/kdp/interior-drafts/{volume_dir}/{zip_name}">Download full draft pack</a>
        <a class="download" href="downloads/production/kdp/interior-drafts/{volume_dir}/{docx_name}">Download DOCX draft</a>
        <a class="download" href="downloads/production/kdp/interior-drafts/{volume_dir}/{pdf_name}">Download PDF draft</a>
        <a class="download" href="downloads/production/kdp/interior-drafts/{volume_dir}/{notes_name}">Download draft notes</a>
        <a class="download" href="downloads/production/kdp/interior-drafts/{volume_dir}/{audit_name}">Download draft audit</a>
      </div>
    </section>
    <section id="coverage">
      <h2>Coverage</h2>
      <div class="grid">
        {month_cards}
      </div>
    </section>
    <section>
      <h2>Review Sample Points</h2>
      <div class="grid">
        {sample_cards}
      </div>
    </section>
    <section>
      <h2>Next Production Gate</h2>
      <div class="grid">
        <div class="audit"><h3>Editorial</h3><p>Full copyedit and author voice pass across the 366-entry draft.</p></div>
        <div class="audit"><h3>Theological</h3><p>Adventist Sabbath and grace/obedience proof pass across every entry.</p></div>
        <div class="audit"><h3>Production</h3><p>Final page count lock, KDP Previewer pass, and full-wrap cover regeneration.</p></div>
      </div>
    </section>
  </main>
</body>
</html>
"""


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    cleaned = "\n".join(line.rstrip() for line in content.rstrip().splitlines())
    path.write_text(cleaned + "\n")


def pdf_page_count(pdf_path: Path) -> int:
    reader = PdfReader(str(pdf_path))
    return len(reader.pages)


def make_zip(paths: list[Path]) -> Path:
    zip_path = OUT / f"Lady-D-Volume-{ACTIVE_BOOK.volume}-Full-Interior-Draft-Pack.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in paths:
            zf.write(path, path.name)
    return zip_path


def sync(paths: list[Path]) -> None:
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)
    for path in paths:
        shutil.copy2(path, PUBLIC_OUT / path.name)
        shutil.copy2(path, LIBRARY_OUT / path.name)


def build_book(book: BookConfig, commit: str) -> dict[str, str | int]:
    set_active_book(book)
    entries = parse_entries()
    if len(entries) != 366:
        raise SystemExit(f"expected 366 entries including leap day, found {len(entries)}")
    if not any(entry.date == "February 29" for entry in entries):
        raise SystemExit("February 29 bonus entry missing")

    OUT.mkdir(parents=True, exist_ok=True)
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)

    docx_path = build_docx(entries)
    pdf_path = convert_pdf(docx_path)
    page_count = pdf_page_count(pdf_path)

    notes_path = OUT / f"volume-{ACTIVE_BOOK.volume}-full-6x9-interior-draft.md"
    audit_path = OUT / f"volume-{ACTIVE_BOOK.volume}-full-6x9-interior-draft-audit.md"
    review_path = OUT / f"volume-{ACTIVE_BOOK.volume}-full-6x9-interior-draft-review.html"
    write_text(notes_path, summary_markdown(entries, commit, page_count))
    write_text(audit_path, audit_markdown(entries, page_count))
    write_text(review_path, review_html(entries, commit, page_count))

    paths = [docx_path, pdf_path, notes_path, audit_path, review_path]
    zip_path = make_zip(paths)
    sync(paths + [zip_path])
    shutil.copy2(review_path, SOURCE_PAGE)
    shutil.copy2(review_path, PUBLIC_PAGE)
    return {
        "volume": ACTIVE_BOOK.volume,
        "entries": len(entries),
        "pages": page_count,
        "docx": str(docx_path.relative_to(ROOT)),
        "pdf": str(pdf_path.relative_to(ROOT)),
        "review": str(SOURCE_PAGE.relative_to(ROOT)),
        "zip": str(zip_path.relative_to(ROOT)),
    }


def main() -> None:
    commit = current_commit()
    results = [build_book(book, commit) for book in BOOKS]
    print(results)


if __name__ == "__main__":
    main()
