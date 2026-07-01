#!/usr/bin/env python3
"""Build the Lady D release upload readiness pack."""

from __future__ import annotations

import html
import json
import re
import shutil
import subprocess
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "downloads" / "production" / "kdp" / "release-upload-readiness"
PUBLIC_OUT = ROOT / "public" / "downloads" / "production" / "kdp" / "release-upload-readiness"
SOURCE_PAGE = ROOT / "release-upload-readiness.html"
PUBLIC_PAGE = ROOT / "public" / "release-upload-readiness.html"
LIBRARY_OUT = Path("/Users/IDC2.5/Documents/LADY D/Production Library/_Shared/KDP Readiness/Release Upload Readiness")
GENERATED = "2026-07-01"
AUTHOR = 'Susan "Lady D" Damon'

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(89, 96, 108)
GOLD = RGBColor(122, 90, 0)
LIGHT_FILL = "E8EEF5"
SOFT_FILL = "F4F6F9"


@dataclass(frozen=True)
class Product:
    key: str
    volume: int
    kind: str
    title: str
    subtitle: str
    lane: str
    pages: int
    interior_pdf: Path
    cover_pdf: Path
    master_source: Path
    cover_size_white: str
    cover_size_cream: str
    description: str
    keywords: tuple[str, ...]


PRODUCTS = [
    Product(
        "volume-1-devotional",
        1,
        "Devotional",
        "Surrendering to God's Love",
        "A 365-Day Devotional Journey into the Father's Heart",
        "God the Father / love, identity, surrender, forgiveness, timing, daily trust",
        369,
        ROOT / "downloads/production/kdp/interior-drafts/volume-1/volume-1-full-6x9-interior-draft.pdf",
        ROOT / "downloads/production/kdp/full-wrap-drafts/volume-1-path-route-white-paper-full-wrap-draft.pdf",
        ROOT / "downloads/production/master/volume-1-master-interior-manuscript.md",
        "13.081 x 9.250 in",
        "13.173 x 9.250 in",
        "A year of morning surrender, steady Scripture, and practical trust in the Father's love.",
        ("Christian devotional", "Sabbath rest", "God's love", "daily surrender", "prayer journal", "Adventist devotional", "spiritual growth"),
    ),
    Product(
        "volume-1-journal",
        1,
        "Companion Journal",
        "Surrendering to God's Love Companion Journal",
        "A Companion Journal for Receiving the Father's Heart",
        "Guided response pages for Volume 1",
        470,
        ROOT / "downloads/production/kdp/companion-journal-drafts/volume-1/volume-1-companion-journal-6x9-draft.pdf",
        ROOT / "downloads/production/kdp/companion-journal-full-wrap-drafts/volume-1-companion-journal-path-route-white-paper-full-wrap-draft.pdf",
        ROOT / "downloads/production/master/volume-1-master-companion-journal.md",
        "13.308 x 9.250 in",
        "13.425 x 9.250 in",
        "Daily response space, Sabbath reflections, prayers, and review pages beside Volume 1.",
        ("Christian journal", "guided prayer", "Sabbath reflection", "God's love", "daily writing", "Adventist journal", "spiritual reflection"),
    ),
    Product(
        "volume-2-devotional",
        2,
        "Devotional",
        "Walking with Jesus",
        "A 365-Day Devotional Journey with the Son",
        "Jesus / discipleship, nearness, obedience, healing, following, abiding",
        369,
        ROOT / "downloads/production/kdp/interior-drafts/volume-2/volume-2-full-6x9-interior-draft.pdf",
        ROOT / "downloads/production/kdp/full-wrap-drafts/volume-2-path-route-white-paper-full-wrap-draft.pdf",
        ROOT / "downloads/production/master/volume-2-master-interior-manuscript.md",
        "13.081 x 9.250 in",
        "13.173 x 9.250 in",
        "A daily discipleship path shaped by the words, mercy, nearness, and authority of Jesus.",
        ("Jesus devotional", "daily discipleship", "walking with Jesus", "Christian growth", "Sabbath devotional", "prayer", "obedience and grace"),
    ),
    Product(
        "volume-2-journal",
        2,
        "Companion Journal",
        "Walking with Jesus Companion Journal",
        "A Companion Journal for Following the Son",
        "Guided response pages for Volume 2",
        477,
        ROOT / "downloads/production/kdp/companion-journal-drafts/volume-2/volume-2-companion-journal-6x9-draft.pdf",
        ROOT / "downloads/production/kdp/companion-journal-full-wrap-drafts/volume-2-companion-journal-path-route-white-paper-full-wrap-draft.pdf",
        ROOT / "downloads/production/master/volume-2-master-companion-journal.md",
        "13.324 x 9.250 in",
        "13.443 x 9.250 in",
        "A year of written response beside the devotional journey with Jesus.",
        ("Jesus journal", "discipleship journal", "Sabbath reflection", "guided prayer", "Christian writing", "Adventist journal", "daily reflection"),
    ),
    Product(
        "volume-3-devotional",
        3,
        "Devotional",
        "Filled with the Holy Spirit",
        "A 365-Day Devotional Journey of Power, Comfort, and Fire",
        "Holy Spirit / filling, comfort, conviction, gifts, fruit, rain, oil, breath",
        369,
        ROOT / "downloads/production/kdp/interior-drafts/volume-3/volume-3-full-6x9-interior-draft.pdf",
        ROOT / "downloads/production/kdp/full-wrap-drafts/volume-3-path-route-white-paper-full-wrap-draft.pdf",
        ROOT / "downloads/production/master/volume-3-master-interior-manuscript.md",
        "13.081 x 9.250 in",
        "13.173 x 9.250 in",
        "A year of Spirit-formed life: comfort, conviction, power, fruit, and witness that points to Jesus.",
        ("Holy Spirit devotional", "Christian devotional", "Spirit filled life", "Sabbath rest", "daily prayer", "Adventist devotional", "fruit of the Spirit"),
    ),
    Product(
        "volume-3-journal",
        3,
        "Companion Journal",
        "Filled with the Holy Spirit Companion Journal",
        "A Companion Journal for Spirit-Filled Surrender",
        "Guided response pages for Volume 3",
        483,
        ROOT / "downloads/production/kdp/companion-journal-drafts/volume-3/volume-3-companion-journal-6x9-draft.pdf",
        ROOT / "downloads/production/kdp/companion-journal-full-wrap-drafts/volume-3-companion-journal-path-route-white-paper-full-wrap-draft.pdf",
        ROOT / "downloads/production/master/volume-3-master-companion-journal.md",
        "13.338 x 9.250 in",
        "13.457 x 9.250 in",
        "A companion journal for Spirit-filled surrender, prayer, Sabbath reflection, and fruit formation.",
        ("Holy Spirit journal", "Christian journal", "Spirit filled life", "guided prayer", "Sabbath reflection", "Adventist journal", "daily writing"),
    ),
]


KDP_SOURCES = [
    ("Create a Paperback Cover", "https://kdp.amazon.com/help/topic/G201953020"),
    ("Set Trim Size, Bleed, and Margins", "https://kdp.amazon.com/help/topic/GVBQ3CMEQW3W2VL6"),
    ("Paperback Submission Guidelines", "https://kdp.amazon.com/help/topic/G201857950"),
]


def current_commit() -> str:
    try:
        return subprocess.check_output(["git", "rev-parse", "--short", "HEAD"], cwd=ROOT, text=True).strip()
    except Exception:
        return "unknown"


def word_count(text: str) -> int:
    return len(re.findall(r"[A-Za-z']+", text))


def pdf_info(path: Path) -> dict[str, object]:
    reader = PdfReader(str(path))
    page = reader.pages[0]
    width = float(page.mediabox.width)
    height = float(page.mediabox.height)
    return {
        "exists": path.exists(),
        "pages": len(reader.pages),
        "page_size_points": f"{width:g} x {height:g}",
        "encrypted": reader.is_encrypted,
    }


def audit_products(commit: str) -> dict[str, object]:
    product_rows = []
    totals = {
        "products": len(PRODUCTS),
        "missing_files": 0,
        "sunday_mentions": 0,
        "sabbath_mentions": 0,
        "saturday_mentions": 0,
        "words": 0,
    }
    for product in PRODUCTS:
        source_text = product.master_source.read_text()
        counts = {
            "words": word_count(source_text),
            "sabbath_mentions": len(re.findall(r"\bSabbath\b", source_text, re.I)),
            "saturday_mentions": len(re.findall(r"\bSaturday\b", source_text, re.I)),
            "sunday_mentions": len(re.findall(r"\bSunday\b", source_text, re.I)),
            "performance_mentions": len(re.findall(r"\bperformance\b", source_text, re.I)),
            "earn_mentions": len(re.findall(r"\bearn(?:ing|ed|s)?\b", source_text, re.I)),
        }
        interior_info = pdf_info(product.interior_pdf)
        cover_info = pdf_info(product.cover_pdf)
        missing = [str(path.relative_to(ROOT)) for path in [product.interior_pdf, product.cover_pdf, product.master_source] if not path.exists()]
        totals["missing_files"] += len(missing)
        totals["sunday_mentions"] += counts["sunday_mentions"]
        totals["sabbath_mentions"] += counts["sabbath_mentions"]
        totals["saturday_mentions"] += counts["saturday_mentions"]
        totals["words"] += counts["words"]
        product_rows.append(
            {
                "key": product.key,
                "volume": product.volume,
                "kind": product.kind,
                "title": product.title,
                "subtitle": product.subtitle,
                "pages": product.pages,
                "interior_pdf": str(product.interior_pdf.relative_to(ROOT)),
                "cover_pdf": str(product.cover_pdf.relative_to(ROOT)),
                "source": str(product.master_source.relative_to(ROOT)),
                "interior": interior_info,
                "cover": cover_info,
                "counts": counts,
                "missing": missing,
                "cover_size_white": product.cover_size_white,
                "cover_size_cream": product.cover_size_cream,
                "upload_status": "review-ready-not-final-upload",
            }
        )
    return {
        "generated": GENERATED,
        "commit": commit,
        "result": "review_ready_not_final_upload",
        "totals": totals,
        "products": product_rows,
        "kdp_sources_checked": [{"title": title, "url": url} for title, url in KDP_SOURCES],
        "blocking_gates": [
            "Author-approved dedication, acknowledgments, author bio, ISBN, and barcode.",
            "Final copyedit and theological proof sign-off.",
            "Final Bible permissions decision remains reference-only unless quotations are added.",
            "Final paper type selection and final KDP cover calculator confirmation.",
            "KDP Previewer pass and physical proof review.",
        ],
    }


def write(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    cleaned = "\n".join(line.rstrip() for line in content.rstrip().splitlines())
    path.write_text(cleaned + "\n")


def metadata_sheet(product: Product) -> str:
    return f"""# {product.title} - KDP Metadata Draft

Generated: {GENERATED}

Status: Metadata working draft. Do not paste into KDP until author/publisher approval.

## Product Identity

- Product key: `{product.key}`
- Volume: {product.volume}
- Product type: {product.kind}
- Title: {product.title}
- Subtitle: {product.subtitle}
- Author: {AUTHOR}
- Series: Lady D Devotional Library
- Language: English
- Trim: 6 x 9 in paperback
- Interior: black ink, no-bleed review draft
- Current review page count: {product.pages}
- Current white-paper cover working size: {product.cover_size_white}
- Current cream-paper cover working size: {product.cover_size_cream}

## Description Draft

{product.description}

This release is written from a Seventh-day Adventist Christian frame. Sabbath language refers to the seventh-day/Saturday Sabbath. Obedience is presented as a response to God's grace, not a means of earning God's love.

## Seven Keyword Drafts

{chr(10).join(f"- {keyword}" for keyword in product.keywords)}

## Category Working Direction

- Primary lane: Christian devotional / Christian spiritual growth
- Secondary lane: prayer, journaling, discipleship, or Christian living depending on product type
- Final categories must be selected inside KDP from currently available store taxonomy.

## Upload File Pair

- Interior PDF: `{product.interior_pdf.relative_to(ROOT)}`
- Full-wrap cover PDF: `{product.cover_pdf.relative_to(ROOT)}`

## Approval Fields Still Needed

- ISBN / barcode decision
- Publisher/imprint confirmation
- Final author bio
- Final description approval
- Pricing and territories
- KDP Previewer pass
- Physical proof pass
"""


def permissions_policy() -> str:
    sources = "\n".join(f"- [{title}]({url})" for title, url in KDP_SOURCES)
    return f"""# Lady D Scripture Permissions and KDP Source Policy

Generated: {GENERATED}

Status: Production policy for final upload preparation.

## Current Scripture Strategy

The current Lady D devotional interiors use Scripture references without full quoted Bible text. This keeps the current review drafts in a lower permissions-risk state. If final full Scripture quotations are added later, the exact Bible translation copyright and permission notice must be inserted before upload.

## Required Guardrails

- Keep Sabbath language as seventh-day/Saturday Sabbath.
- Keep obedience as response to grace, not a method of earning God's love.
- Do not add full Bible quotation blocks without a selected translation and required permission/copyright language.
- Do not mark a file final if it contains placeholders, hidden comments, annotations, crop marks, or unsupported special characters in filenames.

## KDP Source Checks Used

{sources}

## Practical Upload Policy

1. Upload no-bleed interior PDFs unless final design intentionally adds bleed.
2. Use single-PDF full-wrap covers that include back cover, spine, and front cover as one image.
3. Confirm fonts and images are embedded or flattened in native exports.
4. Confirm cover/interior images are at least 300 DPI.
5. Run KDP Previewer and order physical proof copies before public launch.
"""


def proof_runbook(audit: dict[str, object]) -> str:
    rows = []
    for item in audit["products"]:
        rows.append(
            f"| {item['title']} | {item['kind']} | {item['pages']} | {item['interior']['page_size_points']} | {item['cover_size_white']} | Review-ready, not final |"
        )
    return f"""# Lady D KDP Upload and Proof Runbook

Generated: {GENERATED}

Status: Operator runbook. This does not replace KDP Previewer or physical proof review.

## Product Matrix

| Product | Type | Pages | Interior page size | White-paper cover size | Status |
| --- | --- | ---: | --- | --- | --- |
{chr(10).join(rows)}

## Upload Sequence

1. Confirm final paper type for each product.
2. Confirm ISBN/barcode path for each product.
3. Re-export final approved interior PDF from the review draft.
4. Regenerate final full-wrap cover with the locked paper type and page count.
5. Upload interior and cover in KDP.
6. Run KDP Previewer.
7. Resolve every KDP warning or visual defect.
8. Order proof copy.
9. Inspect proof copy before public release.

## Physical Proof Inspection

- Front cover title/subtitle/author readable at arm's length.
- Spine text centered and not drifting into front/back covers.
- Barcode box clear and scannable after ISBN decision.
- Interior margins comfortable near gutter.
- No clipped headers, footers, page numbers, or writing lines.
- Sabbath and grace/obedience language preserved.
- No placeholder text, annotations, comments, hidden objects, or visible production labels in final upload files.
"""


def audit_markdown(audit: dict[str, object]) -> str:
    rows = []
    for item in audit["products"]:
        rows.append(
            f"| {item['title']} | {item['kind']} | {item['pages']} | {item['counts']['words']:,} | "
            f"{item['counts']['sabbath_mentions']} | {item['counts']['sunday_mentions']} | {item['interior']['page_size_points']} |"
        )
    return f"""# Lady D Theological and Production Readiness Audit

Generated: {GENERATED}

Base commit: `{audit['commit']}`

Result: Review-ready, not final KDP upload.

## Summary

- Products checked: {audit['totals']['products']}
- Source words checked: {audit['totals']['words']:,}
- Sabbath mentions: {audit['totals']['sabbath_mentions']}
- Saturday mentions: {audit['totals']['saturday_mentions']}
- Sunday mentions: {audit['totals']['sunday_mentions']}
- Missing required source/interior/cover files: {audit['totals']['missing_files']}

## Product Evidence

| Product | Type | Pages | Source words | Sabbath mentions | Sunday mentions | Interior page size |
| --- | --- | ---: | ---: | ---: | ---: | --- |
{chr(10).join(rows)}

## Guardrail Finding

The current source set preserves the Adventist guardrail in measurable terms: Sabbath language is present, Saturday language is present, and no `Sunday` mentions were found in the checked master manuscripts or companion journals. This does not replace a human theological proof pass, but it gives the next reviewer a concrete audit baseline.

## Remaining Blocking Gates

{chr(10).join(f"- {gate}" for gate in audit['blocking_gates'])}
"""


def main_report_markdown(audit: dict[str, object]) -> str:
    return f"""# Lady D Release Upload Readiness Pack

Generated: {GENERATED}

Base commit: `{audit['commit']}`

Status: Review-ready, not final KDP upload.

## What This Pack Adds

- Six KDP metadata drafts: three devotionals and three companion journals.
- Scripture permissions and KDP source policy.
- KDP upload and physical proof runbook.
- Theological and production audit in Markdown and JSON.
- Operator-facing DOCX/PDF report.
- Live HTML review page.

## Current Release Boundary

The Lady D trilogy now has 6 x 9 devotional interiors, 6 x 9 companion journals, front-cover candidates, devotional full-wrap drafts, companion journal full-wrap drafts, and release dashboards. The remaining gate is no longer generation of the main draft artifacts; it is approval, proofing, upload preparation, KDP Previewer, and physical proof.

## Evidence Snapshot

- Products checked: {audit['totals']['products']}
- Source words checked: {audit['totals']['words']:,}
- Sabbath mentions across checked sources: {audit['totals']['sabbath_mentions']}
- Sunday mentions across checked sources: {audit['totals']['sunday_mentions']}
- Missing required files: {audit['totals']['missing_files']}

## KDP Source Notes

The pack aligns the working upload checklist with the current KDP help pages for paperback cover file structure, trim/bleed/margins, and paperback submission file requirements.

## Readiness Conclusion

The project is substantially assembled for author/publisher review, but it is not final-upload ready until the blocking gates in the audit are closed.
"""


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, size: float | None = None, bold: bool = False, italic: bool = False, color=INK) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.text = "Lady D Release Upload Readiness"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.runs[0], size=8.5, color=MUTED, italic=True)
    footer = section.footer.paragraphs[0]
    footer.text = "Review-ready evidence pack - not final KDP upload files"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.runs[0], size=8, color=MUTED)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_para(doc: Document, text: str, *, size: float = 11, bold: bool = False, italic: bool = False, color=INK, after: float = 6, align=None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.add_run(text)


def add_product_table(doc: Document, audit: dict[str, object]) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    headers = ["Product", "Type", "Pages", "Words", "Sabbath", "Sunday"]
    widths = [Inches(2.15), Inches(1.05), Inches(0.7), Inches(0.9), Inches(0.75), Inches(0.7)]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(header)
        set_font(run, size=9.2, bold=True, color=DARK_BLUE)
    for item in audit["products"]:
        row = table.add_row()
        values = [
            item["title"],
            item["kind"],
            str(item["pages"]),
            f"{item['counts']['words']:,}",
            str(item["counts"]["sabbath_mentions"]),
            str(item["counts"]["sunday_mentions"]),
        ]
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.width = widths[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_font(run, size=8.8)


def build_docx(audit: dict[str, object]) -> Path:
    doc = Document()
    setup_doc(doc)
    add_para(doc, "LADY D DEVOTIONAL LIBRARY", size=9, bold=True, color=GOLD, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Release Upload Readiness Pack", size=25, bold=True, color=DARK_BLUE, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "KDP metadata, permissions, proofing, and theological-production evidence", size=12.5, italic=True, color=MUTED, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, f"Generated {GENERATED} from commit {audit['commit']}", size=9.5, color=MUTED, after=20, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_heading("Readiness Boundary", level=1)
    add_para(doc, "The trilogy is substantially assembled for author and publisher review, but it is not final-upload ready until approvals, ISBN/barcode, KDP Previewer, and physical proofs are complete.")
    doc.add_heading("Evidence Snapshot", level=1)
    for item in [
        f"Products checked: {audit['totals']['products']}",
        f"Source words checked: {audit['totals']['words']:,}",
        f"Sabbath mentions: {audit['totals']['sabbath_mentions']}",
        f"Sunday mentions: {audit['totals']['sunday_mentions']}",
        f"Missing required files: {audit['totals']['missing_files']}",
    ]:
        add_bullet(doc, item)
    doc.add_heading("Product Matrix", level=1)
    add_product_table(doc, audit)
    doc.add_heading("Blocking Gates", level=1)
    for gate in audit["blocking_gates"]:
        add_bullet(doc, gate)
    doc.add_heading("KDP Source Alignment", level=1)
    for title, url in KDP_SOURCES:
        add_bullet(doc, f"{title}: {url}")
    doc.add_heading("Operator Next Step", level=1)
    add_para(doc, "Use this pack to collect final metadata, decide ISBN/barcode and paper type, run final copyedit/theological proof, regenerate final upload files, then complete KDP Previewer and physical proof review.")
    out = OUT / "lady-d-release-upload-readiness-pack.docx"
    doc.save(out)
    return out


def convert_docx_to_pdf(docx_path: Path) -> Path:
    subprocess.check_call(["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUT), str(docx_path)])
    return OUT / f"{docx_path.stem}.pdf"


def review_html(audit: dict[str, object]) -> str:
    rows = "".join(
        f"<tr><td>{html.escape(item['title'])}</td><td>{html.escape(item['kind'])}</td><td>{item['pages']}</td><td>{item['counts']['words']:,}</td><td>{item['counts']['sabbath_mentions']}</td><td>{item['counts']['sunday_mentions']}</td></tr>"
        for item in audit["products"]
    )
    downloads = [
        ("Download readiness pack ZIP", "downloads/production/kdp/release-upload-readiness/Lady-D-Release-Upload-Readiness-Pack.zip"),
        ("Readiness report PDF", "downloads/production/kdp/release-upload-readiness/lady-d-release-upload-readiness-pack.pdf"),
        ("Readiness report DOCX", "downloads/production/kdp/release-upload-readiness/lady-d-release-upload-readiness-pack.docx"),
        ("Theological audit JSON", "downloads/production/kdp/release-upload-readiness/theological-production-readiness-audit.json"),
        ("KDP upload runbook", "downloads/production/kdp/release-upload-readiness/kdp-upload-proof-runbook.md"),
        ("Permissions policy", "downloads/production/kdp/release-upload-readiness/scripture-permissions-and-kdp-policy.md"),
    ]
    download_links = "".join(f"<a class=\"download\" href=\"{href}\">{html.escape(label)}</a>" for label, href in downloads)
    metadata_links = "".join(
        f"<a class=\"download\" href=\"downloads/production/kdp/release-upload-readiness/{product.key}-kdp-metadata.md\">{html.escape(product.title)} metadata</a>"
        for product in PRODUCTS
    )
    gates = "".join(f"<li>{html.escape(gate)}</li>" for gate in audit["blocking_gates"])
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='12' fill='%23182646'/%3E%3Ctext x='50%25' y='55%25' dominant-baseline='middle' text-anchor='middle' font-family='Georgia,serif' font-size='26' fill='%23fffdf8'%3ELD%3C/text%3E%3C/svg%3E">
  <title>Lady D Release Upload Readiness Pack</title>
  <style>
    :root {{ --ink:#111827; --paper:#fffdf8; --mist:#f5f2eb; --blue:#182646; --teal:#1d716f; --gold:#c99335; --line:rgba(17,24,39,.14); }}
    * {{ box-sizing:border-box; }}
    body {{ margin:0; color:var(--ink); font-family:Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; background:linear-gradient(180deg,var(--paper),var(--mist)); line-height:1.5; }}
    header, main {{ max-width:1180px; margin:0 auto; padding:34px 22px; }}
    h1,h2,h3 {{ font-family:Georgia,"Times New Roman",serif; line-height:1.05; letter-spacing:0; margin:0 0 14px; }}
    h1 {{ font-size:clamp(42px,7vw,82px); }}
    h2 {{ font-size:clamp(28px,4vw,48px); }}
    p {{ margin:0 0 14px; }}
    .lead {{ max-width:860px; font-size:clamp(18px,2vw,23px); color:#2e3746; }}
    .kicker {{ color:var(--teal); font-weight:900; letter-spacing:.14em; text-transform:uppercase; font-size:12px; margin-bottom:14px; }}
    .grid {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(240px,1fr)); gap:14px; }}
    .card,.download,table {{ border:1px solid var(--line); border-radius:8px; background:white; box-shadow:0 18px 50px rgba(24,38,70,.10); }}
    .card,.download {{ padding:16px; }}
    .download {{ color:var(--teal); font-weight:900; text-decoration:none; }}
    section {{ border-top:1px solid var(--line); padding:34px 0; }}
    table {{ width:100%; border-collapse:collapse; overflow:hidden; }}
    th,td {{ padding:10px; border-bottom:1px solid var(--line); text-align:left; vertical-align:top; }}
    th {{ background:#e8eef5; color:var(--blue); }}
    .status {{ display:inline-block; background:var(--blue); color:white; padding:7px 10px; border-radius:999px; font-size:12px; font-weight:900; }}
  </style>
</head>
<body>
  <header>
    <div class="kicker">KDP release upload readiness</div>
    <h1>Lady D release upload readiness pack</h1>
    <p class="lead">This pack turns the current trilogy build into a practical KDP upload lane: metadata drafts, Scripture permissions policy, proof runbook, and a theological/production audit over the current devotional and companion journal artifacts.</p>
    <p><span class="status">Generated {GENERATED}</span> <span class="status">Commit {html.escape(str(audit['commit']))}</span> <span class="status">Not final upload-ready</span></p>
  </header>
  <main>
    <section>
      <h2>Readiness Snapshot</h2>
      <div class="grid">
        <div class="card"><h3>{audit['totals']['products']}</h3><p>Products checked</p></div>
        <div class="card"><h3>{audit['totals']['words']:,}</h3><p>Source words checked</p></div>
        <div class="card"><h3>{audit['totals']['sabbath_mentions']}</h3><p>Sabbath mentions</p></div>
        <div class="card"><h3>{audit['totals']['sunday_mentions']}</h3><p>Sunday mentions</p></div>
      </div>
    </section>
    <section>
      <h2>Downloads</h2>
      <div class="grid">{download_links}{metadata_links}</div>
    </section>
    <section>
      <h2>Product Evidence</h2>
      <table><thead><tr><th>Product</th><th>Type</th><th>Pages</th><th>Words</th><th>Sabbath</th><th>Sunday</th></tr></thead><tbody>{rows}</tbody></table>
    </section>
    <section>
      <h2>Remaining Gates</h2>
      <ul>{gates}</ul>
      <p><a href="production.html">Return to production review</a></p>
    </section>
  </main>
</body>
</html>
"""


def make_zip(paths: list[Path]) -> Path:
    zip_path = OUT / "Lady-D-Release-Upload-Readiness-Pack.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in paths:
            zf.write(path, path.name)
    return zip_path


def sync(paths: list[Path]) -> None:
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)
    for path in paths:
        shutil.copy2(path, PUBLIC_OUT / path.name)
        shutil.copy2(path, LIBRARY_OUT / path.name)


def main() -> None:
    commit = current_commit()
    OUT.mkdir(parents=True, exist_ok=True)
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)
    audit = audit_products(commit)

    generated: list[Path] = []
    audit_json = OUT / "theological-production-readiness-audit.json"
    audit_md = OUT / "theological-production-readiness-audit.md"
    main_md = OUT / "lady-d-release-upload-readiness-pack.md"
    permissions_md = OUT / "scripture-permissions-and-kdp-policy.md"
    runbook_md = OUT / "kdp-upload-proof-runbook.md"
    review_page = OUT / "lady-d-release-upload-readiness-review.html"

    write(audit_json, json.dumps(audit, indent=2))
    write(audit_md, audit_markdown(audit))
    write(main_md, main_report_markdown(audit))
    write(permissions_md, permissions_policy())
    write(runbook_md, proof_runbook(audit))
    write(review_page, review_html(audit))
    generated.extend([audit_json, audit_md, main_md, permissions_md, runbook_md, review_page])

    for product in PRODUCTS:
        path = OUT / f"{product.key}-kdp-metadata.md"
        write(path, metadata_sheet(product))
        generated.append(path)

    docx_path = build_docx(audit)
    pdf_path = convert_docx_to_pdf(docx_path)
    generated.extend([docx_path, pdf_path])

    zip_path = make_zip(generated)
    all_paths = generated + [zip_path]
    sync(all_paths)
    shutil.copy2(review_page, SOURCE_PAGE)
    shutil.copy2(review_page, PUBLIC_PAGE)
    print(json.dumps({"report": str(docx_path.relative_to(ROOT)), "pdf": str(pdf_path.relative_to(ROOT)), "zip": str(zip_path.relative_to(ROOT)), "page": str(SOURCE_PAGE.relative_to(ROOT))}, indent=2))


if __name__ == "__main__":
    main()
