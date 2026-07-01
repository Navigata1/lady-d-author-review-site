#!/usr/bin/env python3
"""Build KDP interior finalization artifacts for the Lady D trilogy."""

from __future__ import annotations

import html
import re
import shutil
import subprocess
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "downloads" / "production" / "kdp" / "interior-finalization"
PUBLIC_OUT = ROOT / "public" / "downloads" / "production" / "kdp" / "interior-finalization"
PUBLIC_PAGE = ROOT / "public" / "release-status.html"
SOURCE_PAGE = ROOT / "release-status.html"
LIBRARY_ROOT = Path("/Users/IDC2.5/Documents/LADY D/Production Library")
SHARED_LIBRARY = LIBRARY_ROOT / "_Shared" / "KDP Readiness" / "Interior Finalization"


@dataclass(frozen=True)
class Book:
    volume: int
    title: str
    subtitle: str
    lane: str
    manuscript_pages: int
    journal_pages: int
    manuscript_words: int
    journal_words: int
    sabbath_mentions: int
    folder: str
    cover_width_white: str
    cover_width_cream: str
    journal_cover_width_white: str
    journal_cover_width_cream: str


BOOKS = [
    Book(
        1,
        "Surrendering to God's Love",
        "A 365-Day Devotional Journey into the Father's Heart",
        "God the Father / love, identity, surrender, forgiveness, timing, daily trust",
        369,
        470,
        130601,
        21065,
        279,
        "01 Surrendering to God's Love",
        "13.081 x 9.250 in",
        "13.173 x 9.250 in",
        "13.308 x 9.250 in",
        "13.425 x 9.250 in",
    ),
    Book(
        2,
        "Walking with Jesus",
        "A 365-Day Devotional Journey with the Son",
        "Jesus / discipleship, nearness, obedience, healing, following, abiding",
        369,
        477,
        97711,
        18940,
        300,
        "02 Walking with Jesus",
        "13.081 x 9.250 in",
        "13.173 x 9.250 in",
        "13.324 x 9.250 in",
        "13.443 x 9.250 in",
    ),
    Book(
        3,
        "Filled with the Holy Spirit",
        "A 365-Day Devotional Journey of Power, Comfort, and Fire",
        "Holy Spirit / filling, comfort, conviction, gifts, fruit, rain, oil, breath",
        369,
        483,
        94175,
        19003,
        194,
        "03 Filled with the Holy Spirit",
        "13.081 x 9.250 in",
        "13.173 x 9.250 in",
        "13.338 x 9.250 in",
        "13.457 x 9.250 in",
    ),
]


def current_commit() -> str:
    try:
        return subprocess.check_output(["git", "rev-parse", "--short", "HEAD"], cwd=ROOT, text=True).strip()
    except Exception:
        return "unknown"


def write(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    cleaned = "\n".join(line.rstrip() for line in content.rstrip().splitlines())
    path.write_text(cleaned + "\n")


def front_back_template(book: Book) -> str:
    return f"""# {book.title} - KDP Front and Back Matter Template

Generated: 2026-07-01

Status: Draft publishing template. This is not legal advice and is not the final KDP upload file.

## Book Metadata

- Title: {book.title}
- Subtitle: {book.subtitle}
- Author line: Susan "Lady D" Damon
- Series: Lady D Devotional Library
- Volume: {book.volume}
- Spiritual lane: {book.lane}
- Current manuscript review pages: {book.manuscript_pages}
- Current companion journal 6 x 9 review pages: {book.journal_pages}

## Recommended Front Matter Order

1. Half-title page: `{book.title}`
2. Series page: `Lady D Devotional Library`
3. Title page: `{book.title}` / `{book.subtitle}` / `Susan "Lady D" Damon`
4. Copyright page with final copyright owner, ISBN, imprint, permissions, and edition data.
5. Dedication page: `[Author-provided dedication goes here.]`
6. Author welcome: a short note from Susan "Lady D" Damon in her own voice.
7. How to use this devotional.
8. Scripture and permissions note.
9. Sabbath and grace framing note.
10. Reader orientation.
11. January 1 devotional entry.

## Draft Copyright Page Text

Copyright (c) 2026 Susan Damon. All rights reserved.

Published by IDC Publishing.

ISBN: [TBD]

No part of this publication may be reproduced, stored, or transmitted in any form without written permission from the publisher, except for brief quotations in reviews or devotional discussion.

Scripture references are included in the current review manuscript without full Bible text. If final Bible quotation text is added, insert the exact required copyright and permission notice for the selected Bible translation before KDP upload.

This devotional is written from a Seventh-day Adventist Christian frame. Sabbath language refers to the seventh-day/Saturday Sabbath. Obedience is presented as a response to God's grace, not a means of earning God's love.

## Draft Author Welcome

Dear reader,

This devotional year is an invitation to meet God in the real morning, before the day has had time to name you. Bring your questions, your unfinished places, your family burdens, your hopes, and your need for grace. Let the Lord speak first. Let Scripture steady your steps. Let prayer become practical enough to walk out before noon.

My prayer is that each entry helps you receive God's love more deeply and respond with a life more surrendered, more honest, and more available to Him.

With love,

Susan "Lady D" Damon

## Draft How to Use This Devotional

1. Begin with the Scripture reference.
2. Read the devotional slowly, listening for one phrase that meets your real life.
3. Pray the prayer aloud or rewrite it in your own words.
4. Use the journal prompt to tell the truth before God.
5. Take the Today step before the day gets crowded.
6. Let the morning impact line become a small act of obedience.

## Draft Sabbath and Grace Framing Note

Throughout this devotional, Sabbath refers to the seventh-day/Saturday Sabbath. The goal is not religious performance. Sabbath is treated as a gift of rest, worship, trust, and holy rhythm. Obedience is framed as the fruit of receiving God's grace, not a way to earn His love.

## Recommended Back Matter Order

1. Closing year-end reflection.
2. Invitation to use the companion journal.
3. Prayer of surrender for the next season.
4. About Susan "Lady D" Damon.
5. About IDC Publishing.
6. Other volumes in the Lady D Devotional Library.
7. Acknowledgments.

## Draft Closing Reflection

You have walked through a year of Scripture, prayer, surrender, and practical obedience. Do not rush past what God has formed. Return to the entries that still speak. Carry the prayers into the next season. Let the Lord continue what He has begun.

## Required Before Final Upload

- Replace all placeholders.
- Lock final Bible translation and permissions statement.
- Confirm copyright owner and ISBN.
- Confirm author dedication, acknowledgments, and author bio.
- Run final copyedit and theological review.
- Re-export final interior PDF from the designed 6 x 9 template.
- Run KDP Previewer and proof-order review.
"""


def kit_markdown(commit: str) -> str:
    rows = "\n".join(
        f"| Volume {b.volume} | {b.title} | {b.manuscript_pages} | {b.journal_pages} | "
        f"{b.cover_width_white} / {b.cover_width_cream} | {b.journal_cover_width_white} / {b.journal_cover_width_cream} |" for b in BOOKS
    )
    template_links = "\n".join(
        f"- Volume {b.volume}: `front-back-matter-template-volume-{b.volume}.md`" for b in BOOKS
    )
    return f"""# Lady D KDP Interior Finalization Kit

Generated: 2026-07-01

Repo commit at generation: `{commit}`

Status: Production-planning kit. This is not a final KDP upload package.

## Purpose

The three dated devotional manuscripts are complete at 365 dated entries plus the February 29 bonus for each volume. This kit turns the next publishing phase into a concrete checklist: front matter, back matter, permissions, trim readiness, final interior design, and KDP upload gates.

## Current Book Matrix

| Volume | Title | Devotional review pages | Companion journal 6 x 9 pages | Devotional white / cream cover size | Journal white / cream cover size |
| --- | --- | ---: | ---: | --- | --- |
{rows}

## Completed Production Evidence

- 1,098 day cards across the three volumes.
- Three master interior manuscripts in Markdown, DOCX, and PDF.
- Three master companion journals in Markdown, DOCX, and PDF.
- Master assembly audit result: pass.
- Sabbath guardrail preserved: Sunday mentions remain zero in master manuscripts and journals.
- Nine 6 x 9 front-cover art candidates.
- Path-route full-wrap draft pack for the three main manuscripts.
- Path-route full-wrap draft pack for the three companion journals.
- KDP trim and cover readiness worksheet.
- Volume 1 6 x 9 interior prototype package for the first seven devotional entries.
- Full 6 x 9 interior review drafts for Volumes 1, 2, and 3 with 366 entries each.
- Full 6 x 9 companion journal review drafts for Volumes 1, 2, and 3 with 365 daily pages plus one February 29 bonus page each.
- Release upload readiness pack with six KDP metadata sheets, permissions policy, proof runbook, and theological-production audit.
- Trilogy proof and copyedit audit pack with repetition ledgers, theological watchlist contexts, and volume proof checklists.
- Proof decision resolution pack regenerated with zero current decision items.
- Proof decision application pack generated; the prior 192-item title, morning-impact, and theology queue is preserved as evidence and now audits clear.
- Reader-facing lens application and author-voice copyedit gate generated; the source and public manuscript mirrors now audit at 1,098 context/language lenses and zero internal production labels.
- Volume 1 Days 001-028 author-voice line edits completed across four seven-day batches; the edited batches now have 28 varied morning-impact lines, zero old Volume 1 impact templates, zero internal production labels, and zero Sunday mentions.
- Public Vercel review page for author-facing review.

## Interior Finalization Deliverables In This Kit

{template_links}

- `lady-d-kdp-interior-finalization-kit.md`
- `lady-d-kdp-interior-finalization-kit.docx`
- `lady-d-kdp-interior-finalization-kit.pdf`
- `lady-d-release-readiness-dashboard.html`
- `Lady-D-KDP-Interior-Finalization-Kit.zip`

## Standard Front Matter System

Use this order unless the author or publisher makes a deliberate change:

1. Half-title page.
2. Series page.
3. Title page.
4. Copyright page.
5. Dedication.
6. Author welcome.
7. How to use this devotional.
8. Scripture and permissions note.
9. Sabbath and grace framing note.
10. Reader orientation.
11. January 1 devotional entry.

## Standard Back Matter System

1. Closing year-end reflection.
2. Companion journal invitation.
3. Prayer of surrender for the next season.
4. About Susan "Lady D" Damon.
5. About IDC Publishing.
6. Other volumes in the Lady D Devotional Library.
7. Acknowledgments.

## Permissions Gate

The current manuscripts include Scripture references, not full Bible quotation text. Before KDP upload, choose one of these paths:

1. Keep references only and avoid adding Bible quotation text.
2. Add full Scripture text only after selecting a translation and inserting its required copyright and permission notice.
3. If using public-domain text, still disclose the translation/source clearly and consistently.

## Adventist Guardrail

Sabbath means seventh-day/Saturday Sabbath. Obedience must remain response to grace, not a way to earn God's love. Any final copyedit, typesetting, metadata, or marketing copy must preserve that theological frame.

## Judge And Auditor Loop

Run three to seven passes depending on risk:

1. Editorial judge: reader clarity, voice, flow, and author-heart fit.
2. Theological auditor: Sabbath frame, grace/obedience frame, Jesus-centered Spirit language.
3. Permissions auditor: Bible references/quotations, copyright notice, ISBN/imprint data.
4. Production auditor: 6 x 9 margins, page count, running heads, page numbers, image resolution.
5. Retail judge: cover thumbnail readability, Amazon metadata, description, categories, and series cohesion.
6. Proof auditor: KDP Previewer, printed proof, spine alignment, trim, bleed, and barcode area.
7. Final release judge: all blockers closed and no placeholder text remains.

## Remaining Gates Before True KDP Upload

- Final paper type and trim choices.
- Final front matter and back matter approval.
- Final author bio, dedication, acknowledgments, and ISBN.
- Final Bible permissions statement.
- Final copyedit of all three manuscripts and journals.
- Final approved 6 x 9 upload interiors with locked page counts.
- Regenerated full-wrap covers from final page counts.
- Final KDP metadata approval for all three devotionals and all three companion journals.
- Continue author-voice line edits beyond the completed Volume 1 Days 001-028 batches.
- Author approval of companion journal rhythm and front/back matter.
- KDP Previewer pass for each upload file.
- Physical proof review before public release.

## Recommended Next Production Step

Use the devotional and companion journal 6 x 9 review drafts plus the release upload readiness pack, trilogy proof audit pack, proof decision resolution pack, proof decision application pack, author-voice copyedit gate, and completed Volume 1 Days 001-028 line-edit batches as the next trilogy-wide copyedit and theological proof surface. The proof decision queue is clear; continue seven-day author-voice line edits, metadata approval, and final cover regeneration from locked page counts and paper type.
"""


def dashboard_html(commit: str) -> str:
    cards = "".join(
        f"""
        <article class="card">
          <span>Volume {b.volume}</span>
          <h3>{html.escape(b.title)}</h3>
          <p>{html.escape(b.subtitle)}</p>
          <p><strong>Lane:</strong> {html.escape(b.lane)}</p>
          <dl>
            <div><dt>Manuscript</dt><dd>{b.manuscript_pages} pages / {b.manuscript_words:,} words</dd></div>
            <div><dt>Journal</dt><dd>{b.journal_pages} pages / {b.journal_words:,} source words</dd></div>
            <div><dt>Status</dt><dd>Dated manuscript, devotional 6 x 9 draft, and companion journal 6 x 9 draft are built. Final copyedited upload files are still required.</dd></div>
            <div><dt>Devotional white / cream cover</dt><dd>{b.cover_width_white} / {b.cover_width_cream}</dd></div>
            <div><dt>Journal white / cream cover</dt><dd>{b.journal_cover_width_white} / {b.journal_cover_width_cream}</dd></div>
          </dl>
        </article>
        """
        for b in BOOKS
    )
    completed = [
        "1,098 devotional day cards generated and live.",
        "Three master manuscripts assembled in Markdown, DOCX, and PDF.",
        "Three companion journals assembled in Markdown, DOCX, and PDF.",
        "Nine 6 x 9 front-cover art candidates generated.",
        "Path-route full-wrap draft covers generated for the three main manuscripts.",
        "Path-route full-wrap draft covers generated for the three companion journals.",
        "KDP trim and cover worksheet generated from official KDP guidance.",
        "Volume 1 6 x 9 interior prototype package generated for the first seven devotional entries.",
        "Full 6 x 9 interior review drafts generated for Volumes 1, 2, and 3 with 366 entries each.",
        "Full 6 x 9 companion journal review drafts generated for Volumes 1, 2, and 3.",
        "Interior finalization front/back matter templates created for all three volumes.",
        "Release upload readiness pack generated with KDP metadata drafts, permissions policy, proof runbook, and theological-production audit.",
        "Trilogy proof and copyedit audit pack generated with repetition ledgers, theological watchlist contexts, and volume proof checklists.",
        "Proof decision resolution pack regenerated with zero current title, morning-impact, or theology decisions.",
        "Proof decision application pack generated; the prior 192-item proof queue is preserved and now audits clear.",
        "Author-voice copyedit gate generated with 1,098 source context/language lenses and zero internal production labels in the master manuscripts.",
        "Volume 1 Days 001-028 author-voice line edits completed across four seven-day batches with 28 varied morning-impact lines, zero old Volume 1 impact templates, zero internal production labels, and zero Sunday mentions.",
    ]
    remaining = [
        "Final paper type and ISBN/barcode data.",
        "Author-approved dedication, acknowledgments, and bio.",
        "Final Bible translation permissions statement.",
        "Remaining author-voice line edits after the completed Volume 1 Days 001-028 batches.",
        "Author approval of companion journal rhythm and front/back matter.",
        "Final approved 6 x 9 devotional and journal upload interiors with locked page counts.",
        "Regenerated full-wrap covers from locked page counts.",
        "KDP Previewer and physical proof review.",
    ]
    completed_html = "".join(f"<li>{html.escape(item)}</li>" for item in completed)
    remaining_html = "".join(f"<li>{html.escape(item)}</li>" for item in remaining)
    guardrails = [
        "Sabbath remains seventh-day/Saturday Sabbath throughout the release lane.",
        "Obedience remains a response to God's grace, not a method of earning God's love.",
        "Final copyedit and metadata must preserve the Adventist theological frame.",
        "Scripture references may stay reference-only unless a final translation permission statement is added.",
    ]
    guardrails_html = "".join(f"<li>{html.escape(item)}</li>" for item in guardrails)
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='12' fill='%23182646'/%3E%3Ctext x='50%25' y='55%25' dominant-baseline='middle' text-anchor='middle' font-family='Georgia,serif' font-size='26' fill='%23fffdf8'%3ELD%3C/text%3E%3C/svg%3E">
  <title>Lady D Release Readiness Dashboard</title>
  <style>
    :root {{
      --ink: #111827;
      --paper: #fffdf8;
      --mist: #f5f2eb;
      --indigo: #182646;
      --teal: #1d716f;
      --gold: #c99335;
      --line: rgba(17, 24, 39, .14);
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      color: var(--ink);
      font-family: Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
      background: linear-gradient(180deg, var(--paper), var(--mist));
      line-height: 1.5;
    }}
    header, main {{ max-width: 1180px; margin: 0 auto; padding: 34px 22px; }}
    h1, h2, h3 {{ font-family: Georgia, "Times New Roman", serif; line-height: 1.05; margin: 0 0 14px; letter-spacing: 0; }}
    h1 {{ font-size: clamp(42px, 7vw, 84px); max-width: 960px; }}
    h2 {{ font-size: clamp(28px, 4vw, 48px); }}
    p {{ margin: 0 0 14px; }}
    .lead {{ font-size: clamp(18px, 2vw, 23px); max-width: 820px; color: #2e3746; }}
    .kicker {{ color: var(--teal); font-weight: 900; letter-spacing: .14em; text-transform: uppercase; font-size: 12px; margin-bottom: 14px; }}
    .grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(260px, 1fr)); gap: 16px; }}
    .card, .panel {{
      border: 1px solid var(--line);
      background: white;
      border-radius: 8px;
      padding: 18px;
      box-shadow: 0 18px 50px rgba(24, 38, 70, .10);
    }}
    .card span {{ color: var(--gold); font-size: 12px; font-weight: 900; text-transform: uppercase; letter-spacing: .12em; }}
    .actions {{ display: flex; flex-wrap: wrap; gap: 10px; margin-top: 18px; }}
    .actions a {{
      display: inline-flex;
      align-items: center;
      min-height: 38px;
      padding: 8px 11px;
      border: 1px solid var(--line);
      border-radius: 999px;
      color: var(--teal);
      background: white;
      font-size: 13px;
      font-weight: 900;
      text-decoration: none;
    }}
    dl {{ margin: 16px 0 0; display: grid; gap: 10px; }}
    dt {{ font-weight: 900; color: var(--indigo); }}
    dd {{ margin: 0; color: #374151; }}
    section {{ border-top: 1px solid var(--line); padding: 34px 0; }}
    ul {{ margin: 0; padding-left: 22px; }}
    li {{ margin: 8px 0; }}
    .status {{ display: inline-block; background: var(--indigo); color: white; padding: 7px 10px; border-radius: 999px; font-size: 12px; font-weight: 900; }}
    a {{ color: var(--teal); text-underline-offset: 3px; }}
  </style>
</head>
<body>
  <header>
    <div class="kicker">IDC Publishing release dashboard</div>
    <h1>Lady D Devotional Library release readiness</h1>
    <p class="lead">The three devotional manuscripts are complete at the dated-entry level and now have master assemblies, companion journals, cover candidates, KDP trim math, devotional and companion journal full-wrap drafts, interior finalization templates, the first Volume 1 6 x 9 interior prototype, full 6 x 9 devotional drafts, full 6 x 9 companion journal drafts, a release-upload readiness pack, a trilogy proof/copyedit audit pack, a proof decision resolution pack, proof decision application evidence showing the prior 192-item queue now cleared, an author-voice copyedit gate, and the first four completed Volume 1 seven-day author-voice line-edit batches covering Days 001-028. This dashboard separates what is complete from what still gates true KDP upload readiness.</p>
    <p><span class="status">Generated 2026-07-01</span> <span class="status">Base commit: {html.escape(commit)}</span></p>
    <div class="actions">
      <a href="production.html">Production Review</a>
      <a href="volume-1-interior-prototype.html">Volume 1 Interior Prototype</a>
      <a href="volume-1-full-interior-draft.html">Volume 1 Full Draft</a>
      <a href="volume-2-full-interior-draft.html">Volume 2 Full Draft</a>
      <a href="volume-3-full-interior-draft.html">Volume 3 Full Draft</a>
      <a href="volume-1-companion-journal-draft.html">Volume 1 Journal Draft</a>
      <a href="volume-2-companion-journal-draft.html">Volume 2 Journal Draft</a>
      <a href="volume-3-companion-journal-draft.html">Volume 3 Journal Draft</a>
      <a href="downloads/production/kdp/interior-finalization/Lady-D-KDP-Interior-Finalization-Kit.zip">Interior Kit ZIP</a>
      <a href="downloads/production/kdp/interior-finalization/lady-d-kdp-interior-finalization-kit.pdf">Interior Kit PDF</a>
      <a href="downloads/production/kdp/companion-journal-full-wrap-drafts/Lady-D-Companion-Journal-Full-Wrap-Draft-Pack.zip">Journal Full-Wrap ZIP</a>
      <a href="release-upload-readiness.html">Upload Readiness Page</a>
      <a href="downloads/production/kdp/release-upload-readiness/Lady-D-Release-Upload-Readiness-Pack.zip">Upload Readiness ZIP</a>
      <a href="downloads/production/kdp/release-upload-readiness/lady-d-release-upload-readiness-pack.pdf">Upload Readiness PDF</a>
      <a href="trilogy-proof-audit.html">Proof Audit Page</a>
      <a href="downloads/production/kdp/trilogy-proof-audit/Lady-D-Trilogy-Proof-Audit-Pack.zip">Proof Audit ZIP</a>
      <a href="downloads/production/kdp/trilogy-proof-audit/lady-d-trilogy-proof-audit-pack.pdf">Proof Audit PDF</a>
      <a href="proof-decision-resolution.html">Proof Decision Page</a>
      <a href="downloads/production/kdp/proof-decision-resolution/Lady-D-Proof-Decision-Resolution-Pack.zip">Proof Decision ZIP</a>
      <a href="downloads/production/kdp/proof-decision-resolution/lady-d-proof-decision-resolution-pack.pdf">Proof Decision PDF</a>
      <a href="proof-decision-application.html">Proof Applied Page</a>
      <a href="downloads/production/kdp/proof-decision-application/Lady-D-Proof-Decision-Application-Pack.zip">Proof Applied ZIP</a>
      <a href="downloads/production/kdp/proof-decision-application/lady-d-proof-decision-application-pack.pdf">Proof Applied PDF</a>
      <a href="author-voice-copyedit.html">Author Voice Gate</a>
      <a href="downloads/production/kdp/author-voice-copyedit/Lady-D-Author-Voice-Copyedit-Pack.zip">Author Voice ZIP</a>
      <a href="downloads/production/kdp/author-voice-copyedit/lady-d-author-voice-copyedit-pack.pdf">Author Voice PDF</a>
      <a href="volume-1-days-001-007-line-edit.html">V1 Days 001-007 Line Edit</a>
      <a href="volume-1-days-008-014-line-edit.html">V1 Days 008-014 Line Edit</a>
      <a href="volume-1-days-015-021-line-edit.html">V1 Days 015-021 Line Edit</a>
      <a href="volume-1-days-022-028-line-edit.html">V1 Days 022-028 Line Edit</a>
      <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-001-007/Lady-D-Volume-1-Days-001-007-Line-Edit-Pack.zip">Days 001-007 ZIP</a>
      <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-008-014/Lady-D-Volume-1-Days-008-014-Line-Edit-Pack.zip">Days 008-014 ZIP</a>
      <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-015-021/Lady-D-Volume-1-Days-015-021-Line-Edit-Pack.zip">Days 015-021 ZIP</a>
      <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-022-028/Lady-D-Volume-1-Days-022-028-Line-Edit-Pack.zip">Days 022-028 ZIP</a>
    </div>
  </header>
  <main>
    <section>
      <h2>Book Status</h2>
      <div class="grid">{cards}</div>
    </section>
    <section class="grid">
      <div class="panel">
        <h2>Finalized So Far</h2>
        <ul>{completed_html}</ul>
      </div>
      <div class="panel">
        <h2>Remaining Gates Before KDP Upload</h2>
        <ul>{remaining_html}</ul>
      </div>
      <div class="panel">
        <h2>Guardrail Snapshot</h2>
        <ul>{guardrails_html}</ul>
      </div>
    </section>
    <section>
      <h2>Release Upload Readiness Pack</h2>
      <p class="lead">The newest pack gathers KDP-facing metadata drafts, Scripture permissions policy, proof-run instructions, and a machine-readable theological-production audit for all three devotionals and all three companion journals.</p>
      <div class="grid">
        <article class="card">
          <span>Coverage</span>
          <h3>Six KDP products</h3>
          <p>Three devotional books and three companion journals now have dedicated metadata sheets with descriptions, keyword candidates, upload file pairs, and approval fields.</p>
        </article>
        <article class="card">
          <span>Evidence</span>
          <h3>377,171 source words</h3>
          <p>The regenerated proof audit reports zero Sunday mentions, zero placeholders, zero duplicate-title groups, zero repeated morning-impact groups, and zero priority theology contexts across the current review set.</p>
        </article>
        <article class="card">
          <span>Boundary</span>
          <h3>Review ready, not upload final</h3>
          <p>Final ISBN/barcode, paper type, Bible permissions, final proof, KDP Previewer, and physical proof review still gate public release.</p>
        </article>
      </div>
      <div class="actions">
        <a href="release-upload-readiness.html">Open Upload Readiness Page</a>
        <a href="downloads/production/kdp/release-upload-readiness/Lady-D-Release-Upload-Readiness-Pack.zip">Download Upload Readiness ZIP</a>
        <a href="downloads/production/kdp/release-upload-readiness/theological-production-readiness-audit.json">Audit JSON</a>
        <a href="downloads/production/kdp/release-upload-readiness/kdp-upload-proof-runbook.md">Proof Runbook</a>
        <a href="downloads/production/kdp/release-upload-readiness/scripture-permissions-and-kdp-policy.md">Permissions Policy</a>
      </div>
    </section>
    <section>
      <h2>Trilogy Proof And Copyedit Audit Pack</h2>
      <p class="lead">The newest proof pack documents the regenerated clean gate after the proof decisions were applied: structure checks, placeholder scan, Adventist Sabbath guardrail confirmation, repeated-title ledger, repeated morning-impact ledger, and grace/performance context review.</p>
      <div class="grid">
        <article class="card">
          <span>Structure</span>
          <h3>Three devotional masters pass</h3>
          <p>Each devotional master contains 365 dated entries, one February 29 bonus, 366 Scripture references, 366 prayers, and 366 morning-impact lines.</p>
        </article>
        <article class="card">
          <span>Guardrail</span>
          <h3>Zero Sunday mentions</h3>
          <p>The audit checks 377,171 source words and reports zero placeholders, zero Sunday mentions, and zero priority theology contexts requiring resolution.</p>
        </article>
        <article class="card">
          <span>Editorial result</span>
          <h3>Repetition ledgers clear</h3>
          <p>The regenerated ledger reports zero duplicate title groups and zero repeated morning-impact groups across the six review products.</p>
        </article>
      </div>
      <div class="actions">
        <a href="trilogy-proof-audit.html">Open Proof Audit Page</a>
        <a href="downloads/production/kdp/trilogy-proof-audit/Lady-D-Trilogy-Proof-Audit-Pack.zip">Download Proof Audit ZIP</a>
        <a href="downloads/production/kdp/trilogy-proof-audit/trilogy-proof-audit.json">Audit JSON</a>
        <a href="downloads/production/kdp/trilogy-proof-audit/title-and-impact-repetition-ledger.md">Repetition Ledger</a>
        <a href="downloads/production/kdp/trilogy-proof-audit/theological-watchlist-context-ledger.md">Theological Watchlist</a>
      </div>
    </section>
    <section>
      <h2>Proof Decision Resolution Pack</h2>
      <p class="lead">The newest resolution pack has been regenerated after source application. It records a clear decision queue while preserving the machine-readable register for audit traceability.</p>
      <div class="grid">
        <article class="card">
          <span>Decision queue</span>
          <h3>0 proof decisions</h3>
          <p>The current register includes no duplicate-title decisions, no morning-impact rhythm decisions, and no priority theology contexts.</p>
        </article>
        <article class="card">
          <span>Priority</span>
          <h3>Grace preserved</h3>
          <p>The cleaned audit keeps obedience framed as response to grace and never as a way to earn God's love.</p>
        </article>
        <article class="card">
          <span>Operator surface</span>
          <h3>Trace remains</h3>
          <p>The prior queue and source application ledger are preserved in the proof decision application pack.</p>
        </article>
      </div>
      <div class="actions">
        <a href="proof-decision-resolution.html">Open Proof Decision Page</a>
        <a href="downloads/production/kdp/proof-decision-resolution/Lady-D-Proof-Decision-Resolution-Pack.zip">Download Proof Decision ZIP</a>
        <a href="downloads/production/kdp/proof-decision-resolution/proof-decision-resolution.json">Decision Register JSON</a>
        <a href="downloads/production/kdp/proof-decision-resolution/title-revision-decision-queue.md">Title Queue</a>
        <a href="downloads/production/kdp/proof-decision-resolution/morning-impact-resolution-queue.md">Morning-Impact Queue</a>
        <a href="downloads/production/kdp/proof-decision-resolution/theological-context-resolution-queue.md">Theology Queue</a>
      </div>
    </section>
    <section>
      <h2>Proof Decision Application Pack</h2>
      <p class="lead">The application pack proves the prior 192-item proof decision queue was applied into the source manuscripts and regenerated through the proof lane. It is evidence of a cleared proof-decision queue, not final KDP upload approval.</p>
      <div class="grid">
        <article class="card">
          <span>Prior queue</span>
          <h3>192 applied</h3>
          <p>The preserved snapshot includes 150 title decisions, 30 morning-impact decisions, and 12 theology decisions.</p>
        </article>
        <article class="card">
          <span>Current audit</span>
          <h3>0 current blockers</h3>
          <p>The regenerated audit reports zero duplicate title groups, zero repeated impact groups, and zero priority theology contexts.</p>
        </article>
        <article class="card">
          <span>Boundary</span>
          <h3>Still not upload final</h3>
          <p>Final copyedit, author approvals, ISBN/barcode, locked covers, KDP Previewer, and physical proofs still gate public release.</p>
        </article>
      </div>
      <div class="actions">
        <a href="proof-decision-application.html">Open Proof Applied Page</a>
        <a href="downloads/production/kdp/proof-decision-application/Lady-D-Proof-Decision-Application-Pack.zip">Download Proof Applied ZIP</a>
        <a href="downloads/production/kdp/proof-decision-application/proof-decision-application.json">Application JSON</a>
        <a href="downloads/production/kdp/proof-decision-application/proof-decision-source-application.json">Source Application Ledger</a>
        <a href="downloads/production/kdp/proof-decision-application/proof-decision-prior-queue-snapshot.json">Prior Queue Snapshot</a>
      </div>
    </section>
    <section>
      <h2>Author Voice Copyedit Gate</h2>
      <p class="lead">The author-voice gate proves the internal production-lens language has been removed from the source and public manuscript surfaces, then names the remaining voice-rhythm work honestly before any final upload claim.</p>
      <div class="grid">
        <article class="card">
          <span>Reader-facing lens cleanup</span>
          <h3>1,098 context lenses</h3>
          <p>Every devotional entry across the three volumes now has a reader-facing context/language lens in the source manuscript set and public mirror.</p>
        </article>
        <article class="card">
          <span>Internal label audit</span>
          <h3>0 remaining</h3>
          <p>The regenerated master manuscripts report zero internal production-lens labels and zero Sunday mentions.</p>
        </article>
        <article class="card">
          <span>Next editorial surface</span>
          <h3>22 queue items</h3>
          <p>The pack identifies repeated sentences, morning-impact patterns, and voice-rhythm issues for the next Lady D line edit.</p>
        </article>
      </div>
      <div class="actions">
        <a href="author-voice-copyedit.html">Open Author Voice Gate</a>
        <a href="downloads/production/kdp/author-voice-copyedit/Lady-D-Author-Voice-Copyedit-Pack.zip">Download Author Voice ZIP</a>
        <a href="downloads/production/kdp/author-voice-copyedit/lady-d-author-voice-copyedit-pack.pdf">Author Voice PDF</a>
        <a href="downloads/production/kdp/author-voice-copyedit/lady-d-author-voice-copyedit-pack.docx">Author Voice DOCX</a>
        <a href="downloads/production/kdp/author-voice-copyedit/author-voice-copyedit-pack.json">Author Voice JSON</a>
        <a href="downloads/production/kdp/author-voice-copyedit/reader-facing-lens-application.md">Lens Application Report</a>
      </div>
    </section>
    <section>
      <h2>Author-Voice Line Edit Progress</h2>
      <p class="lead">Volume 1 Days 001-028 has moved from structural manuscript into line-edited review surface across four seven-day batches. The edited batches have 28 varied morning-impact lines, zero old Volume 1 impact templates, zero internal production labels, and zero Sunday mentions. This is continuing review progress, not final KDP upload approval.</p>
      <div class="grid">
        <article class="card">
          <span>Completed batch</span>
          <h3>Volume 1 Days 001-007</h3>
          <p>The edited pilot belongs to <em>Surrendering to God's Love</em> and is mirrored into public downloads and the Production Library.</p>
        </article>
        <article class="card">
          <span>Completed batch</span>
          <h3>Volume 1 Days 008-014</h3>
          <p>The second January batch is now mirrored into public downloads and the Production Library with its own PDF, DOCX, JSON, and ZIP pack.</p>
        </article>
        <article class="card">
          <span>Completed batch</span>
          <h3>Volume 1 Days 015-021</h3>
          <p>The third January batch is now mirrored into public downloads and the Production Library with its own PDF, DOCX, JSON, and ZIP pack.</p>
        </article>
        <article class="card">
          <span>Completed batch</span>
          <h3>Volume 1 Days 022-028</h3>
          <p>The fourth January batch is now mirrored into public downloads and the Production Library with its own PDF, DOCX, JSON, and ZIP pack.</p>
        </article>
        <article class="card">
          <span>Combined result</span>
          <h3>28 varied impacts</h3>
          <p>The repeated `Let the Father's love carry` impact frame has been removed from all four completed batches.</p>
        </article>
        <article class="card">
          <span>Guardrail</span>
          <h3>0 Sunday mentions</h3>
          <p>The Adventist Sabbath frame remains seventh-day/Saturday and obedience remains response to grace in all four edited batches.</p>
        </article>
      </div>
      <div class="actions">
        <a href="volume-1-days-001-007-line-edit.html">Open Days 001-007 Page</a>
        <a href="volume-1-days-008-014-line-edit.html">Open Days 008-014 Page</a>
        <a href="volume-1-days-015-021-line-edit.html">Open Days 015-021 Page</a>
        <a href="volume-1-days-022-028-line-edit.html">Open Days 022-028 Page</a>
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-001-007/Lady-D-Volume-1-Days-001-007-Line-Edit-Pack.zip">Days 001-007 ZIP</a>
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-008-014/Lady-D-Volume-1-Days-008-014-Line-Edit-Pack.zip">Days 008-014 ZIP</a>
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-015-021/Lady-D-Volume-1-Days-015-021-Line-Edit-Pack.zip">Days 015-021 ZIP</a>
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-022-028/Lady-D-Volume-1-Days-022-028-Line-Edit-Pack.zip">Days 022-028 ZIP</a>
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-008-014/volume-1-days-008-014-line-edit-report.pdf">Days 008-014 PDF</a>
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-008-014/volume-1-days-008-014-line-edit-report.docx">Days 008-014 DOCX</a>
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-008-014/volume-1-days-008-014-line-edit.json">Days 008-014 JSON</a>
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-015-021/volume-1-days-015-021-line-edit-report.pdf">Days 015-021 PDF</a>
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-015-021/volume-1-days-015-021-line-edit-report.docx">Days 015-021 DOCX</a>
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-015-021/volume-1-days-015-021-line-edit.json">Days 015-021 JSON</a>
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-022-028/volume-1-days-022-028-line-edit-report.pdf">Days 022-028 PDF</a>
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-022-028/volume-1-days-022-028-line-edit-report.docx">Days 022-028 DOCX</a>
        <a href="downloads/production/kdp/author-voice-line-edit/volume-1-days-022-028/volume-1-days-022-028-line-edit.json">Days 022-028 JSON</a>
      </div>
    </section>
    <section>
      <h2>Active Recommendation</h2>
      <p class="lead">Use the three devotional drafts, three companion journal drafts, release-upload readiness pack, trilogy proof audit pack, proof decision resolution pack, proof decision application pack, author-voice copyedit gate, and Volume 1 Days 001-028 line-edit batches as the next copyedit and theological proof surface. Do not mark any file as final upload-ready until the remaining author-voice line edits, ISBN/barcode, paper type, Bible permissions, KDP Previewer, and physical proof review pass.</p>
      <p><a href="production.html">Return to production review page</a></p>
    </section>
  </main>
</body>
</html>
"""


def make_zip(paths: list[Path]) -> Path:
    zip_path = OUT / "Lady-D-KDP-Interior-Finalization-Kit.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in paths:
            zf.write(path, path.name)
    shutil.copy2(zip_path, PUBLIC_OUT / zip_path.name)
    shutil.copy2(zip_path, SHARED_LIBRARY / zip_path.name)
    return zip_path


def set_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for name, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[name]
        style.font.name = "Georgia"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(24, 38, 70)


def clean_markdown_text(value: str) -> str:
    return value.replace("`", "").replace("**", "").replace("*", "")


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        cells = [clean_markdown_text(cell.strip()) for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for idx, cell in enumerate(rows[0]):
        run = table.rows[0].cells[idx].paragraphs[0].add_run(cell)
        run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row[: len(cells)]):
            cells[idx].text = value


def build_docx(markdown: str) -> Path:
    docx_path = OUT / "lady-d-kdp-interior-finalization-kit.docx"
    doc = Document()
    set_doc_styles(doc)
    doc.core_properties.title = "Lady D KDP Interior Finalization Kit"
    doc.core_properties.author = "IDC Publishing"

    table_lines: list[str] = []
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if line.startswith("|"):
            table_lines.append(line)
            continue
        if table_lines:
            add_markdown_table(doc, table_lines)
            table_lines = []
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(clean_markdown_text(line[2:]), level=0)
        elif line.startswith("## "):
            doc.add_heading(clean_markdown_text(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(clean_markdown_text(line[4:]), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(clean_markdown_text(line[2:]), style="List Bullet")
        elif re.match(r"^\d+\\. ", line):
            doc.add_paragraph(clean_markdown_text(re.sub(r"^\d+\\. ", "", line)), style="List Number")
        else:
            doc.add_paragraph(clean_markdown_text(line))
    if table_lines:
        add_markdown_table(doc, table_lines)

    doc.save(docx_path)
    shutil.copy2(docx_path, PUBLIC_OUT / docx_path.name)
    shutil.copy2(docx_path, SHARED_LIBRARY / docx_path.name)
    return docx_path


def convert_docx_to_pdf(docx_path: Path) -> Path:
    pdf_path = docx_path.with_suffix(".pdf")
    subprocess.check_call(["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUT), str(docx_path)])
    shutil.copy2(pdf_path, PUBLIC_OUT / pdf_path.name)
    shutil.copy2(pdf_path, SHARED_LIBRARY / pdf_path.name)
    return pdf_path


def main() -> None:
    commit = current_commit()
    OUT.mkdir(parents=True, exist_ok=True)
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    SHARED_LIBRARY.mkdir(parents=True, exist_ok=True)

    generated: list[Path] = []

    for book in BOOKS:
        path = OUT / f"front-back-matter-template-volume-{book.volume}.md"
        write(path, front_back_template(book))
        generated.append(path)
        shutil.copy2(path, PUBLIC_OUT / path.name)
        book_dir = LIBRARY_ROOT / book.folder / "06 Master Assembly"
        shutil.copy2(path, book_dir / path.name)

    kit = OUT / "lady-d-kdp-interior-finalization-kit.md"
    write(kit, kit_markdown(commit))
    generated.append(kit)
    shutil.copy2(kit, PUBLIC_OUT / kit.name)
    shutil.copy2(kit, SHARED_LIBRARY / kit.name)
    docx_path = build_docx(kit.read_text())
    pdf_path = convert_docx_to_pdf(docx_path)
    generated.extend([docx_path, pdf_path])

    dashboard = OUT / "lady-d-release-readiness-dashboard.html"
    write(dashboard, dashboard_html(commit))
    generated.append(dashboard)
    shutil.copy2(dashboard, PUBLIC_OUT / dashboard.name)
    shutil.copy2(dashboard, SOURCE_PAGE)
    shutil.copy2(dashboard, PUBLIC_PAGE)
    shutil.copy2(dashboard, SHARED_LIBRARY / dashboard.name)

    zip_path = make_zip(generated)
    print({"kit": str(kit.relative_to(ROOT)), "dashboard": str(dashboard.relative_to(ROOT)), "zip": str(zip_path.relative_to(ROOT))})


if __name__ == "__main__":
    main()
