#!/usr/bin/env python3
"""Build full 6x9 companion journal draft packages for the Lady D trilogy."""

from __future__ import annotations

import html
import json
import re
import shutil
import subprocess
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader, PdfWriter


ROOT = Path(__file__).resolve().parents[1]
LIBRARY_ROOT = Path("/Users/IDC2.5/Documents/LADY D/Production Library")
AUTHOR = 'Susan "Lady D" Damon'
GENERATED = "2026-07-01"

INK = RGBColor(31, 31, 31)
MUTED = RGBColor(84, 84, 84)
DEEP = RGBColor(42, 37, 33)
RULE = RGBColor(143, 120, 79)
SAGE = RGBColor(74, 101, 82)
SOFT_FILL = "F8F5EE"
BOX_FILL = "EFE9DD"


@dataclass(frozen=True)
class BookConfig:
    volume: int
    folder: str
    title: str
    devotional_subtitle: str
    journal_subtitle: str
    lane_label: str


@dataclass(frozen=True)
class JournalItem:
    kind: str
    sort_key: float
    week_title: str
    heading: str
    focus: str
    write: str
    practice: str
    body: tuple[str, ...]


BOOKS = [
    BookConfig(
        volume=1,
        folder="01 Surrendering to God's Love",
        title="Surrendering to God's Love",
        devotional_subtitle="A 365-Day Devotional Journey into the Father's Heart",
        journal_subtitle="A Companion Journal for Receiving the Father's Heart",
        lane_label="God the Father",
    ),
    BookConfig(
        volume=2,
        folder="02 Walking with Jesus",
        title="Walking with Jesus",
        devotional_subtitle="A 365-Day Devotional Journey with the Son",
        journal_subtitle="A Companion Journal for Following the Son",
        lane_label="Jesus the Son",
    ),
    BookConfig(
        volume=3,
        folder="03 Filled with the Holy Spirit",
        title="Filled with the Holy Spirit",
        devotional_subtitle="A 365-Day Devotional Journey of Power, Comfort, and Fire",
        journal_subtitle="A Companion Journal for Spirit-Filled Surrender",
        lane_label="The Holy Spirit",
    ),
]


ACTIVE_BOOK = BOOKS[0]
BOOK_ROOT = LIBRARY_ROOT / ACTIVE_BOOK.folder
MASTER_MD = ROOT / "downloads" / "production" / "master" / f"volume-{ACTIVE_BOOK.volume}-master-companion-journal.md"
OUT = ROOT / "downloads" / "production" / "kdp" / "companion-journal-drafts" / f"volume-{ACTIVE_BOOK.volume}"
PUBLIC_OUT = ROOT / "public" / "downloads" / "production" / "kdp" / "companion-journal-drafts" / f"volume-{ACTIVE_BOOK.volume}"
SOURCE_PAGE = ROOT / f"volume-{ACTIVE_BOOK.volume}-companion-journal-draft.html"
PUBLIC_PAGE = ROOT / "public" / f"volume-{ACTIVE_BOOK.volume}-companion-journal-draft.html"
LIBRARY_OUT = BOOK_ROOT / "06 Master Assembly" / "Companion Journal Full Interior Draft"


def set_active_book(book: BookConfig) -> None:
    global ACTIVE_BOOK, BOOK_ROOT, MASTER_MD, OUT, PUBLIC_OUT, SOURCE_PAGE, PUBLIC_PAGE, LIBRARY_OUT

    ACTIVE_BOOK = book
    BOOK_ROOT = LIBRARY_ROOT / book.folder
    MASTER_MD = ROOT / "downloads" / "production" / "master" / f"volume-{book.volume}-master-companion-journal.md"
    OUT = ROOT / "downloads" / "production" / "kdp" / "companion-journal-drafts" / f"volume-{book.volume}"
    PUBLIC_OUT = ROOT / "public" / "downloads" / "production" / "kdp" / "companion-journal-drafts" / f"volume-{book.volume}"
    SOURCE_PAGE = ROOT / f"volume-{book.volume}-companion-journal-draft.html"
    PUBLIC_PAGE = ROOT / "public" / f"volume-{book.volume}-companion-journal-draft.html"
    LIBRARY_OUT = BOOK_ROOT / "06 Master Assembly" / "Companion Journal Full Interior Draft"


def current_commit() -> str:
    try:
        return subprocess.check_output(["git", "rev-parse", "--short", "HEAD"], cwd=ROOT, text=True).strip()
    except Exception:
        return "unknown"


def set_font(run, name: str = "Georgia", size: float | None = None, color=INK, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_bottom_rule(paragraph, color: str = "D9CCB4", size: str = "5") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)

    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
    normal.font.size = Pt(8.8)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.2)
    normal.paragraph_format.line_spacing = 1.02

    for name, size, before, after in [
        ("Heading 1", 16, 15, 7),
        ("Heading 2", 12.5, 10, 5),
        ("Heading 3", 10.8, 7, 3),
    ]:
        style = doc.styles[name]
        style.font.name = "Georgia"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
        style.font.size = Pt(size)
        style.font.color.rgb = DEEP
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.04

    header = section.header.paragraphs[0]
    header.text = f"{ACTIVE_BOOK.title} - Companion Journal"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(header.runs[0], size=7.2, color=MUTED, italic=True)

    footer = section.footer.paragraphs[0]
    footer.text = f"Volume {ACTIVE_BOOK.volume} Companion Journal Draft - not final KDP upload file"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.runs[0], size=7, color=MUTED)


def add_para(doc: Document, text: str = "", *, size: float = 9, bold: bool = False, italic: bool = False, color=INK, align=None, before: float = 0, after: float = 4, style: str | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.02
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_callout(doc: Document, label: str, body: str, fill: str = SOFT_FILL) -> None:
    if not body:
        return
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(4.85)
    cell = table.cell(0, 0)
    cell.width = Inches(4.85)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 0.98
    r = p.add_run(f"{label}: ")
    set_font(r, size=8, bold=True, color=DEEP)
    r = p.add_run(body)
    set_font(r, size=8, color=INK)


def add_ruled_lines(doc: Document, count: int, *, after: float = 3.2) -> None:
    for _ in range(count):
        add_para(doc, "_" * 60, size=8, color=RGBColor(210, 199, 178), after=after)


def strip_markdown(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^[-*]\s+", "", text)
    text = re.sub(r"^\d+[.)]\s+", "", text)
    text = text.replace("**", "")
    return re.sub(r"\s+", " ", text).strip()


def split_blocks(raw: str) -> list[str]:
    blocks: list[str] = []
    for part in re.split(r"\n\s*\n", raw.strip()):
        cleaned = "\n".join(line.strip() for line in part.splitlines() if line.strip() and not line.strip().startswith("<!--"))
        if cleaned:
            blocks.append(cleaned)
    return blocks


def parse_daily_body(raw: str) -> tuple[str, str, str, tuple[str, ...]]:
    focus = ""
    write = ""
    practice = ""
    extras: list[str] = []

    for block in split_blocks(raw):
        parsed_any = False
        for line in block.splitlines():
            line = line.strip()
            if not line:
                continue
            line = re.sub(r"^[-*]\s+", "", line)
            label_match = re.match(r"^(?:\*\*)?(Focus|Write|Practice):(?:\*\*)?\s*(.+)$", line, flags=re.I | re.S)
            if label_match:
                parsed_any = True
                label = label_match.group(1).lower()
                value = strip_markdown(label_match.group(2))
                if label == "focus":
                    focus = value
                elif label == "write":
                    write = value
                else:
                    practice = value
            else:
                extras.append(strip_markdown(line))

    if not write and extras:
        write = extras.pop(0)
    return focus, write, practice, tuple(extras)


def parse_extra_body(raw: str) -> tuple[str, ...]:
    return tuple(strip_markdown(block) for block in split_blocks(raw))


def parse_items() -> list[JournalItem]:
    text = MASTER_MD.read_text()
    week_matches = list(re.finditer(r"^## (.+)$", text, flags=re.M))
    items: list[JournalItem] = []
    source_index = 0

    for idx, week_match in enumerate(week_matches):
        week_title = week_match.group(1).strip()
        if week_title == "Journal Orientation":
            continue
        start = week_match.end()
        end = week_matches[idx + 1].start() if idx + 1 < len(week_matches) else len(text)
        section = text[start:end]
        heading_matches = list(re.finditer(r"^### (.+)$", section, flags=re.M))
        for h_idx, heading_match in enumerate(heading_matches):
            heading = heading_match.group(1).strip()
            body_start = heading_match.end()
            body_end = heading_matches[h_idx + 1].start() if h_idx + 1 < len(heading_matches) else len(section)
            body = section[body_start:body_end].strip()
            if not body:
                continue

            source_index += 1
            day_match = re.search(r"\bDay\s+(\d{3})\b", heading)
            bonus = bool(re.search(r"\b(Bonus|Leap Day)\b", heading, flags=re.I))
            if day_match:
                day = int(day_match.group(1))
                focus, write, practice, extras = parse_daily_body(body)
                items.append(
                    JournalItem(
                        kind="day",
                        sort_key=float(day),
                        week_title=week_title,
                        heading=heading,
                        focus=focus,
                        write=write,
                        practice=practice,
                        body=extras,
                    )
                )
            elif bonus:
                focus, write, practice, extras = parse_daily_body(body)
                items.append(
                    JournalItem(
                        kind="bonus",
                        sort_key=59.5,
                        week_title=week_title,
                        heading=heading,
                        focus=focus,
                        write=write,
                        practice=practice,
                        body=extras,
                    )
                )
            else:
                kind = "prayer" if "prayer" in heading.lower() else "review"
                items.append(
                    JournalItem(
                        kind=kind,
                        sort_key=source_index + 1000,
                        week_title=week_title,
                        heading=heading,
                        focus="",
                        write="",
                        practice="",
                        body=parse_extra_body(body),
                    )
                )
    return items


def ordered_items(items: list[JournalItem]) -> list[JournalItem]:
    return items


def daily_items(items: list[JournalItem]) -> list[JournalItem]:
    return [item for item in items if item.kind == "day"]


def bonus_items(items: list[JournalItem]) -> list[JournalItem]:
    return [item for item in items if item.kind == "bonus"]


def day_number(item: JournalItem) -> int:
    match = re.search(r"\bDay\s+(\d{3})\b", item.heading)
    if not match:
        raise ValueError(f"item is not a day: {item.heading}")
    return int(match.group(1))


def add_title_page(doc: Document, items: list[JournalItem]) -> None:
    add_para(doc, "LADY D DEVOTIONAL LIBRARY", size=8.5, bold=True, color=RULE, align=WD_ALIGN_PARAGRAPH.CENTER, before=52, after=22)
    add_para(doc, ACTIVE_BOOK.title, size=25, bold=True, color=DEEP, align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    add_para(doc, "Companion Journal", size=17, bold=True, color=SAGE, align=WD_ALIGN_PARAGRAPH.CENTER, after=7)
    add_para(doc, ACTIVE_BOOK.journal_subtitle, size=11.5, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=32)
    p = add_para(doc, AUTHOR, size=12.5, bold=True, color=DEEP, align=WD_ALIGN_PARAGRAPH.CENTER, after=42)
    add_bottom_rule(p, color="958463", size="8")
    add_para(
        doc,
        f"Full 6 x 9 companion journal review draft - {len(daily_items(items))} daily reflection pages plus {len(bonus_items(items))} February 29 bonus page",
        size=8.8,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=10,
    )
    add_para(
        doc,
        "Prepared by IDC Publishing for author review. This is a designed review draft, not the final KDP upload file.",
        size=8.8,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=0,
    )
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    doc.add_heading("How to Use This Companion Journal", level=1)
    for paragraph in [
        "Use this journal beside the devotional volume, one day at a time.",
        "Read the day's Scripture reference and devotional entry first, then bring your honest response here.",
        "The goal is not polished language. The goal is surrendered truth before God: what He is showing, what He is healing, and what He is asking you to practice today.",
        "Weekly prayers, Sabbath reflections, review pages, and bridge pages are included so the year can become a formed life rather than a stack of isolated mornings.",
    ]:
        add_para(doc, paragraph, size=9.4, after=6)

    doc.add_heading("Guardrail for This Draft", level=1)
    add_para(
        doc,
        "Sabbath language refers to the seventh-day/Saturday Sabbath. Obedience is presented as a response to God's grace, not a method of earning God's love.",
        size=9.2,
        bold=True,
        after=8,
    )
    add_para(
        doc,
        "Final publication still requires author approval, copyedit, theological proof, ISBN and permissions decisions, KDP Previewer review, and physical proof review.",
        size=9,
        after=0,
    )
    doc.add_page_break()


def add_daily_page(doc: Document, item: JournalItem) -> None:
    sabbath = "Sabbath" in item.heading
    small_label = f"{item.heading} / {item.week_title}"
    add_para(doc, small_label, size=7.2, bold=True, color=RULE, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=1.8)
    title = "Sabbath Reflection" if sabbath else "Daily Reflection"
    add_para(doc, title, size=14.2, bold=True, color=DEEP, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    if item.focus:
        add_callout(doc, "Focus", item.focus, fill=BOX_FILL)
    add_callout(doc, "Reflect", item.write, fill=SOFT_FILL)
    add_callout(doc, "Practice", item.practice, fill=SOFT_FILL)
    for extra in item.body:
        add_para(doc, extra, size=8.2, color=MUTED, italic=True, after=3)
    add_para(doc, "Journal Response", size=8, bold=True, color=SAGE, before=5, after=1.5)
    add_ruled_lines(doc, 13 if not sabbath else 12)
    doc.add_page_break()


def add_bonus_page(doc: Document, item: JournalItem) -> None:
    add_para(doc, f"{item.heading} / {item.week_title}", size=7.2, bold=True, color=RULE, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=1.8)
    add_para(doc, "February 29 Bonus Reflection", size=14.2, bold=True, color=DEEP, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    if item.focus:
        add_callout(doc, "Focus", item.focus, fill=BOX_FILL)
    add_callout(doc, "Reflect", item.write, fill=SOFT_FILL)
    add_callout(doc, "Practice", item.practice, fill=SOFT_FILL)
    for extra in item.body:
        add_para(doc, extra, size=8.2, color=MUTED, italic=True, after=3)
    add_para(doc, "Bonus-Day Response", size=8, bold=True, color=SAGE, before=5, after=1.5)
    add_ruled_lines(doc, 12)
    doc.add_page_break()


def add_extra_page(doc: Document, item: JournalItem) -> None:
    label = "Weekly Prayer" if item.kind == "prayer" else "Review and Bridge"
    add_para(doc, item.week_title, size=7.2, bold=True, color=RULE, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=1.8)
    add_para(doc, item.heading, size=14, bold=True, color=DEEP, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, label, size=8, bold=True, color=SAGE, align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    if item.kind == "prayer":
        for paragraph in item.body:
            add_para(doc, paragraph, size=9.4, italic=True, after=6)
        add_para(doc, "Prayer Notes", size=8, bold=True, color=SAGE, before=8, after=1.5)
        add_ruled_lines(doc, 12)
    else:
        for paragraph in item.body:
            add_para(doc, paragraph, size=8.8, after=3)
            add_ruled_lines(doc, 2, after=2.6)
    doc.add_page_break()


def add_back_matter(doc: Document, items: list[JournalItem]) -> None:
    doc.add_heading("Closing Note for Review", level=1)
    add_para(
        doc,
        f"This companion journal draft flows {len(daily_items(items))} daily reflection pages, {len(bonus_items(items))} February 29 bonus page, and weekly/monthly prayer and review pages through a 6 x 9 author-review rhythm.",
        size=9.2,
        after=6,
    )
    doc.add_heading("Before Final Upload", level=2)
    for item in [
        "Author-read the page rhythm beside the matching devotional volume.",
        "Complete copyedit and theological proof across prompts, prayers, Sabbath references, and practice lines.",
        "Lock paper type, page count, ISBN, barcode, and KDP category/metadata.",
        "Regenerate full-wrap covers from the final locked companion journal page count.",
        "Run KDP Previewer and order physical proof before public release.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def build_docx(items: list[JournalItem]) -> Path:
    doc = Document()
    setup_doc(doc)
    add_title_page(doc, items)
    add_front_matter(doc)
    for item in ordered_items(items):
        if item.kind == "day":
            add_daily_page(doc, item)
        elif item.kind == "bonus":
            add_bonus_page(doc, item)
        else:
            add_extra_page(doc, item)
    add_back_matter(doc, items)
    OUT.mkdir(parents=True, exist_ok=True)
    out = OUT / f"volume-{ACTIVE_BOOK.volume}-companion-journal-6x9-draft.docx"
    doc.save(out)
    return out


def normalize_pdf_trim(pdf_path: Path) -> None:
    reader = PdfReader(str(pdf_path))
    writer = PdfWriter()
    width = 432
    height = 648
    for page in reader.pages:
        page.mediabox.lower_left = (0, 0)
        page.mediabox.upper_right = (width, height)
        page.cropbox.lower_left = (0, 0)
        page.cropbox.upper_right = (width, height)
        page.trimbox.lower_left = (0, 0)
        page.trimbox.upper_right = (width, height)
        writer.add_page(page)
    with pdf_path.open("wb") as fh:
        writer.write(fh)


def convert_pdf(docx_path: Path) -> Path:
    subprocess.check_call(["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUT), str(docx_path)])
    pdf_path = OUT / f"{docx_path.stem}.pdf"
    normalize_pdf_trim(pdf_path)
    return pdf_path


def pdf_page_count(pdf_path: Path) -> int:
    return len(PdfReader(str(pdf_path)).pages)


def cover_width(page_count: int, paper: str) -> str:
    spine_factor = 0.002252 if paper == "white" else 0.0025
    return f"{12.25 + page_count * spine_factor:.3f} x 9.250 in"


def audit_payload(items: list[JournalItem], page_count: int, commit: str) -> dict[str, object]:
    days = [day_number(item) for item in daily_items(items)]
    missing = [n for n in range(1, 366) if n not in set(days)]
    duplicates = sorted({day for day in days if days.count(day) > 1})
    incomplete = []
    for item in daily_items(items) + bonus_items(items):
        if not item.write:
            incomplete.append(f"{item.heading}: reflect prompt")
        if not item.practice:
            incomplete.append(f"{item.heading}: practice prompt")
    return {
        "generated": GENERATED,
        "commit": commit,
        "volume": ACTIVE_BOOK.volume,
        "title": ACTIVE_BOOK.title,
        "source": str(MASTER_MD.relative_to(ROOT)),
        "daily_pages": len(days),
        "bonus_pages": len(bonus_items(items)),
        "prayer_pages": len([item for item in items if item.kind == "prayer"]),
        "review_pages": len([item for item in items if item.kind == "review"]),
        "total_content_pages": len(items),
        "rendered_pdf_pages": page_count,
        "pdf_page_size_points": "432 x 648",
        "missing_days": missing,
        "duplicate_days": duplicates,
        "incomplete_prompts": incomplete,
        "white_paper_full_wrap_size": cover_width(page_count, "white"),
        "cream_paper_full_wrap_size": cover_width(page_count, "cream"),
        "status": "pass" if not missing and not duplicates and not incomplete else "review",
    }


def summary_markdown(items: list[JournalItem], payload: dict[str, object]) -> str:
    return f"""# Volume {ACTIVE_BOOK.volume} Companion Journal 6x9 Draft

Generated: {GENERATED}

Base commit: `{payload["commit"]}`

Status: Full companion journal review draft. This is not a final KDP upload file.

## Book

- Title: {ACTIVE_BOOK.title}
- Journal subtitle: {ACTIVE_BOOK.journal_subtitle}
- Devotional lane: {ACTIVE_BOOK.lane_label}
- Author: {AUTHOR}
- Trim: 6 x 9 inches, no bleed interior draft
- Daily reflection pages: {payload["daily_pages"]}
- February 29 bonus pages: {payload["bonus_pages"]}
- Prayer pages: {payload["prayer_pages"]}
- Review / bridge pages: {payload["review_pages"]}
- Rendered PDF pages: {payload["rendered_pdf_pages"]}
- White-paper full-wrap working size from this draft: {payload["white_paper_full_wrap_size"]}
- Cream-paper full-wrap working size from this draft: {payload["cream_paper_full_wrap_size"]}
- Source: `{payload["source"]}`

## Judge And Auditor Loop Applied

1. Editorial judge: keep prompts plain, direct, and usable beside the devotional.
2. Theological auditor: preserve {ACTIVE_BOOK.lane_label}, grace-before-performance, and Adventist Sabbath frame.
3. Production auditor: 6 x 9 trim, no-bleed journal rhythm, exact 432 x 648 pt PDF page boxes.
4. Retail judge: capture a locked draft page count for later full-wrap cover regeneration.
5. Release auditor: mark review-draft status clearly and block final-upload language.

## Adventist Guardrail

Sabbath remains seventh-day/Saturday Sabbath. Obedience remains response to grace, not a method of earning God's love.

## Remaining Before Final Upload

- Author approval of the writing-space rhythm.
- Final copyedit and theological proof.
- Final ISBN, copyright, and permissions decisions.
- KDP Previewer pass.
- Physical proof review.
- Regenerated full-wrap cover from the final locked journal page count.
"""


def audit_markdown(payload: dict[str, object]) -> str:
    missing = payload["missing_days"]
    duplicates = payload["duplicate_days"]
    incomplete = payload["incomplete_prompts"]
    return f"""# Volume {ACTIVE_BOOK.volume} Companion Journal Draft Audit

Generated: {GENERATED}

## Result

{str(payload["status"]).upper()} for full companion journal review publication to the author-facing review site.

## Evidence

- Daily reflection pages parsed: {payload["daily_pages"]}
- Expected numbered daily reflections: 365
- February 29 bonus pages parsed: {payload["bonus_pages"]}
- Prayer pages parsed: {payload["prayer_pages"]}
- Review / bridge pages parsed: {payload["review_pages"]}
- Total parsed content pages: {payload["total_content_pages"]}
- Rendered PDF pages: {payload["rendered_pdf_pages"]}
- Exact PDF page size: 432 x 648 pt after export normalization
- White-paper full-wrap working size from this draft: {payload["white_paper_full_wrap_size"]}
- Cream-paper full-wrap working size from this draft: {payload["cream_paper_full_wrap_size"]}
- Missing daily reflections: {len(missing)}
- Duplicate daily reflections: {len(duplicates)}
- Incomplete reflect/practice prompts: {len(incomplete)}

## Coverage

{chr(10).join(f"- Missing Day {day:03d}" for day in missing) if missing else "- No missing numbered daily reflections."}
{chr(10).join(f"- Duplicate Day {day:03d}" for day in duplicates) if duplicates else "- No duplicate numbered daily reflections."}

## Prompt Completeness

{chr(10).join(f"- {item}" for item in incomplete) if incomplete else "- All daily and bonus reflection pages include reflect and practice prompts."}

## Guardrail Result

- Sabbath guardrail is explicit in the front matter.
- Obedience remains response to grace, not performance.
- Draft status is explicit and does not claim KDP upload readiness.
"""


def review_html(items: list[JournalItem], payload: dict[str, object]) -> str:
    volume = ACTIVE_BOOK.volume
    volume_dir = f"volume-{volume}"
    title = html.escape(ACTIVE_BOOK.title)
    docx_name = f"volume-{volume}-companion-journal-6x9-draft.docx"
    pdf_name = f"volume-{volume}-companion-journal-6x9-draft.pdf"
    notes_name = f"volume-{volume}-companion-journal-6x9-draft.md"
    audit_name = f"volume-{volume}-companion-journal-6x9-draft-audit.md"
    audit_json_name = f"volume-{volume}-companion-journal-6x9-draft-audit.json"
    zip_name = f"Lady-D-Volume-{volume}-Companion-Journal-Draft-Pack.zip"
    day_samples = daily_items(items)
    sample_source = [day_samples[0], day_samples[1], day_samples[58], bonus_items(items)[0], day_samples[59], day_samples[-2], day_samples[-1]]
    sample_cards = "\n".join(
        f"""<article class="entry-card">
          <span>{html.escape(item.heading)}</span>
          <h3>{html.escape("Sabbath Reflection" if "Sabbath" in item.heading else "Daily Reflection")}</h3>
          <p>{html.escape(item.write)}</p>
          <p><strong>Practice:</strong> {html.escape(item.practice)}</p>
        </article>"""
        for item in sample_source
    )
    stats = [
        ("Daily pages", payload["daily_pages"]),
        ("Bonus pages", payload["bonus_pages"]),
        ("Prayer pages", payload["prayer_pages"]),
        ("Review pages", payload["review_pages"]),
        ("PDF pages", payload["rendered_pdf_pages"]),
        ("Audit status", str(payload["status"]).upper()),
    ]
    stat_cards = "\n".join(f"<article class=\"stat\"><strong>{html.escape(str(value))}</strong><span>{html.escape(label)}</span></article>" for label, value in stats)
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='12' fill='%23182646'/%3E%3Ctext x='50%25' y='55%25' dominant-baseline='middle' text-anchor='middle' font-family='Georgia,serif' font-size='26' fill='%23fffdf8'%3ELD%3C/text%3E%3C/svg%3E">
  <title>Volume {volume} Companion Journal 6x9 Draft - Lady D</title>
  <style>
    :root {{
      --ink: #171717;
      --paper: #fffdf8;
      --mist: #f4efe4;
      --deep: #2d2925;
      --line: rgba(45, 41, 37, .16);
      --gold: #8b6c35;
      --sage: #4e6f62;
    }}
    * {{ box-sizing: border-box; }}
    body {{ margin: 0; color: var(--ink); font-family: Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; background: linear-gradient(180deg, var(--paper), var(--mist)); line-height: 1.55; }}
    nav {{ display: flex; flex-wrap: wrap; gap: 10px; padding: 14px 22px; background: #171717; }}
    nav a {{ color: white; text-decoration: none; font-size: 13px; font-weight: 850; padding: 7px 10px; border-radius: 999px; background: rgba(255,255,255,.08); }}
    header, main {{ max-width: 1160px; margin: 0 auto; padding: 38px 22px; }}
    h1, h2, h3 {{ font-family: Georgia, "Times New Roman", serif; letter-spacing: 0; line-height: 1.05; margin: 0 0 14px; }}
    h1 {{ font-size: clamp(42px, 7vw, 82px); max-width: 960px; }}
    h2 {{ font-size: clamp(30px, 4vw, 52px); }}
    h3 {{ font-size: 23px; }}
    p {{ margin: 0 0 14px; }}
    .lead {{ max-width: 840px; font-size: clamp(18px, 2vw, 23px); color: #343434; }}
    .kicker {{ color: var(--sage); text-transform: uppercase; letter-spacing: .14em; font-size: 12px; font-weight: 900; margin-bottom: 14px; }}
    .status, .download, .entry-card, .audit, .stat {{ border: 1px solid var(--line); border-radius: 8px; background: rgba(255, 253, 248, .94); box-shadow: 0 18px 50px rgba(45, 41, 37, .10); }}
    .status {{ padding: 18px; margin-top: 24px; }}
    .grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(230px, 1fr)); gap: 14px; }}
    section {{ border-top: 1px solid var(--line); padding: 34px 0; }}
    .download {{ display: block; padding: 16px; color: var(--deep); font-weight: 850; text-decoration: none; }}
    .entry-card, .audit, .stat {{ padding: 18px; }}
    .entry-card span {{ color: var(--gold); font-size: 12px; font-weight: 900; text-transform: uppercase; letter-spacing: .12em; }}
    .stat strong {{ display: block; color: var(--deep); font-family: Georgia, "Times New Roman", serif; font-size: 38px; line-height: 1; }}
    .stat span {{ color: var(--gold); font-weight: 900; }}
  </style>
</head>
<body>
  <nav aria-label="Companion journal draft navigation">
    <a href="production.html">Production Review</a>
    <a href="volume-1-companion-journal-draft.html">Volume 1 Journal</a>
    <a href="volume-2-companion-journal-draft.html">Volume 2 Journal</a>
    <a href="volume-3-companion-journal-draft.html">Volume 3 Journal</a>
    <a href="volume-1-full-interior-draft.html">V1 Devotional Draft</a>
    <a href="release-status.html">Release Dashboard</a>
    <a href="#downloads">Downloads</a>
  </nav>
  <header>
    <div class="kicker">Volume {volume} companion journal draft</div>
    <h1>Full 6x9 companion journal draft for <em>{title}</em></h1>
    <p class="lead">The companion journal is now flowed into a KDP-style 6x9 review rhythm with daily writing pages, the February 29 bonus page, weekly prayers, review pages, Sabbath framing, and a machine-readable audit for release tracking.</p>
    <div class="status">
      <p><strong>Status:</strong> Full companion journal review draft ready. Not final KDP upload file.</p>
      <p><strong>Base commit:</strong> {html.escape(str(payload["commit"]))}</p>
      <p><strong>Daily reflection pages:</strong> {payload["daily_pages"]}; <strong>bonus pages:</strong> {payload["bonus_pages"]}.</p>
      <p><strong>Rendered PDF pages:</strong> {payload["rendered_pdf_pages"]}; exact 432 x 648 pt page boxes.</p>
      <p><strong>Guardrail:</strong> Sabbath remains seventh-day/Saturday Sabbath; obedience remains response to grace.</p>
    </div>
  </header>
  <main>
    <section id="downloads">
      <h2>Companion Journal Downloads</h2>
      <div class="grid">
        <a class="download" href="downloads/production/kdp/companion-journal-drafts/{volume_dir}/{zip_name}">Download companion journal draft pack</a>
        <a class="download" href="downloads/production/kdp/companion-journal-drafts/{volume_dir}/{docx_name}">Download DOCX draft</a>
        <a class="download" href="downloads/production/kdp/companion-journal-drafts/{volume_dir}/{pdf_name}">Download PDF draft</a>
        <a class="download" href="downloads/production/kdp/companion-journal-drafts/{volume_dir}/{notes_name}">Download draft notes</a>
        <a class="download" href="downloads/production/kdp/companion-journal-drafts/{volume_dir}/{audit_name}">Download draft audit</a>
        <a class="download" href="downloads/production/kdp/companion-journal-drafts/{volume_dir}/{audit_json_name}">Download audit JSON</a>
      </div>
    </section>
    <section>
      <h2>Draft Coverage</h2>
      <div class="grid">{stat_cards}</div>
    </section>
    <section>
      <h2>Review Sample Points</h2>
      <div class="grid">{sample_cards}</div>
    </section>
    <section>
      <h2>Next Production Gate</h2>
      <div class="grid">
        <div class="audit"><h3>Editorial</h3><p>Read beside the matching devotional volume and tune the prompt rhythm where Susan wants a warmer or clearer page.</p></div>
        <div class="audit"><h3>Theological</h3><p>Proof Sabbath, obedience, Spirit language, and grace-before-performance across the journal prompts.</p></div>
        <div class="audit"><h3>Production</h3><p>Use the rendered page count to regenerate final companion journal full-wrap covers after approval.</p></div>
      </div>
    </section>
  </main>
</body>
</html>
"""


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    cleaned = "\n".join(line.rstrip() for line in content.rstrip().splitlines())
    path.write_text(cleaned + "\n")


def make_zip(paths: list[Path]) -> Path:
    zip_path = OUT / f"Lady-D-Volume-{ACTIVE_BOOK.volume}-Companion-Journal-Draft-Pack.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in paths:
            zf.write(path, path.name)
    return zip_path


def sync(paths: list[Path]) -> None:
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)
    for path in paths:
        shutil.copy2(path, PUBLIC_OUT / path.name)
        shutil.copy2(path, LIBRARY_OUT / path.name)


def validate_items(items: list[JournalItem]) -> None:
    days = [day_number(item) for item in daily_items(items)]
    missing = [n for n in range(1, 366) if n not in set(days)]
    duplicates = sorted({day for day in days if days.count(day) > 1})
    if missing or duplicates:
        raise SystemExit(f"Volume {ACTIVE_BOOK.volume} journal day coverage failed: missing={missing} duplicates={duplicates}")
    if len(bonus_items(items)) != 1:
        raise SystemExit(f"Volume {ACTIVE_BOOK.volume} expected one February 29 bonus reflection, found {len(bonus_items(items))}")


def build_book(book: BookConfig, commit: str) -> dict[str, object]:
    set_active_book(book)
    OUT.mkdir(parents=True, exist_ok=True)
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)

    items = parse_items()
    validate_items(items)
    docx_path = build_docx(items)
    pdf_path = convert_pdf(docx_path)
    page_count = pdf_page_count(pdf_path)
    payload = audit_payload(items, page_count, commit)
    if payload["status"] != "pass":
        raise SystemExit(f"Volume {ACTIVE_BOOK.volume} companion journal audit did not pass: {payload}")

    notes_path = OUT / f"volume-{ACTIVE_BOOK.volume}-companion-journal-6x9-draft.md"
    audit_path = OUT / f"volume-{ACTIVE_BOOK.volume}-companion-journal-6x9-draft-audit.md"
    audit_json_path = OUT / f"volume-{ACTIVE_BOOK.volume}-companion-journal-6x9-draft-audit.json"
    review_path = OUT / f"volume-{ACTIVE_BOOK.volume}-companion-journal-6x9-draft-review.html"

    write_text(notes_path, summary_markdown(items, payload))
    write_text(audit_path, audit_markdown(payload))
    write_text(audit_json_path, json.dumps(payload, indent=2))
    write_text(review_path, review_html(items, payload))

    paths = [docx_path, pdf_path, notes_path, audit_path, audit_json_path, review_path]
    zip_path = make_zip(paths)
    sync(paths + [zip_path])
    shutil.copy2(review_path, SOURCE_PAGE)
    shutil.copy2(review_path, PUBLIC_PAGE)
    return {
        "volume": ACTIVE_BOOK.volume,
        "daily_pages": payload["daily_pages"],
        "bonus_pages": payload["bonus_pages"],
        "content_pages": payload["total_content_pages"],
        "pages": page_count,
        "docx": str(docx_path.relative_to(ROOT)),
        "pdf": str(pdf_path.relative_to(ROOT)),
        "review": str(SOURCE_PAGE.relative_to(ROOT)),
        "zip": str(zip_path.relative_to(ROOT)),
    }


def main() -> None:
    commit = current_commit()
    results = [build_book(book, commit) for book in BOOKS]
    print(json.dumps(results, indent=2))


if __name__ == "__main__":
    main()
