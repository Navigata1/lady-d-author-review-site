#!/usr/bin/env python3
"""Build the DOCX/PDF version of the Lady D KDP finalization kit."""

from __future__ import annotations

import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from build_kdp_interior_finalization_kit import BOOKS, OUT, PUBLIC_OUT, SHARED_LIBRARY  # noqa: E402


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(17, 24, 39)
GOLD = RGBColor(122, 90, 0)
LIGHT_FILL = "E8EEF5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def apply_run_style(run, size: int | None = None, bold: bool = False, color=INK) -> None:
    run.font.name = "Calibri"
    run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    run.bold = bold


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)


def add_callout(doc: Document, label: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    r = p.add_run(label + ": ")
    apply_run_style(r, bold=True, color=GOLD)
    r = p.add_run(body)
    apply_run_style(r)


def add_status_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    headers = ["Volume", "Title", "Current files", "KDP working size"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(text)
        apply_run_style(run, bold=True, color=DARK_BLUE)
    widths = [Inches(0.75), Inches(1.85), Inches(2.0), Inches(1.9)]
    for book in BOOKS:
        row = table.add_row()
        values = [
            f"Vol. {book.volume}",
            book.title,
            f"{book.manuscript_pages} devotional pages\n{book.journal_pages} journal pages\n{book.manuscript_words:,} devotional words",
            f"Devotional white: {book.cover_width_white}\nDevotional cream: {book.cover_width_cream}\nJournal white: {book.journal_cover_width_white}\nJournal cream: {book.journal_cover_width_cream}",
        ]
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.width = widths[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            run = cell.paragraphs[0].add_run(value)
            apply_run_style(run)


def build_docx() -> Path:
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(10)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Lady D KDP Interior Finalization Kit")
    apply_run_style(r, size=26, color=DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.add_run("Prepared for Susan \"Lady D\" Damon - generated 2026-07-01")

    add_callout(
        doc,
        "Status",
        "This is a production finalization kit, not a final KDP upload package. It defines the remaining publishing gates and the standard front/back matter system.",
    )

    doc.add_heading("Current Book Matrix", level=1)
    add_status_table(doc)

    doc.add_heading("Completed Production Evidence", level=1)
    for item in [
        "1,098 devotional day cards across the three volumes.",
        "Three master manuscripts and three companion journals in Markdown, DOCX, and PDF.",
        "Master assembly audit result: pass.",
        "Sabbath guardrail preserved with zero Sunday mentions in master artifacts.",
        "Nine 6 x 9 front-cover candidates and first-pass Path-route full-wrap drafts.",
        "KDP trim and cover readiness worksheet generated from current official KDP guidance.",
        "Volume 1 6 x 9 interior prototype package generated for the first seven devotional entries.",
        "Full 6 x 9 interior review drafts generated for Volumes 1, 2, and 3 with 366 entries each.",
        "Full 6 x 9 companion journal review drafts generated for Volumes 1, 2, and 3.",
        "Public Vercel review page for author-facing review.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Standard Front Matter System", level=1)
    for item in [
        "Half-title page.",
        "Series page.",
        "Title page.",
        "Copyright page.",
        "Dedication.",
        "Author welcome.",
        "How to use this devotional.",
        "Scripture and permissions note.",
        "Sabbath and grace framing note.",
        "Reader orientation.",
        "January 1 devotional entry.",
    ]:
        add_number(doc, item)

    doc.add_heading("Standard Back Matter System", level=1)
    for item in [
        "Closing year-end reflection.",
        "Companion journal invitation.",
        "Prayer of surrender for the next season.",
        "About Susan \"Lady D\" Damon.",
        "About IDC Publishing.",
        "Other volumes in the Lady D Devotional Library.",
        "Acknowledgments.",
    ]:
        add_number(doc, item)

    doc.add_heading("Permissions Gate", level=1)
    for item in [
        "Keep references only and avoid full Bible quotation text.",
        "Or add full Scripture text only after selecting a translation and inserting its required copyright and permission notice.",
        "If using public-domain text, disclose the translation/source clearly and consistently.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Adventist Guardrail", level=1)
    doc.add_paragraph(
        "Sabbath means seventh-day/Saturday Sabbath. Obedience must remain response to grace, not a way to earn God's love. Any final copyedit, typesetting, metadata, or marketing copy must preserve that theological frame."
    )

    doc.add_heading("Judge And Auditor Loop", level=1)
    for item in [
        "Editorial judge: reader clarity, voice, flow, and author-heart fit.",
        "Theological auditor: Sabbath frame, grace/obedience frame, Jesus-centered Spirit language.",
        "Permissions auditor: Bible references/quotations, copyright notice, ISBN/imprint data.",
        "Production auditor: 6 x 9 margins, page count, running heads, page numbers, image resolution.",
        "Retail judge: cover thumbnail readability, Amazon metadata, description, categories, and series cohesion.",
        "Proof auditor: KDP Previewer, printed proof, spine alignment, trim, bleed, and barcode area.",
        "Final release judge: all blockers closed and no placeholder text remains.",
    ]:
        add_number(doc, item)

    doc.add_heading("Remaining Gates Before True KDP Upload", level=1)
    for item in [
        "Final paper type and trim choices.",
        "Final front matter and back matter approval.",
        "Final author bio, dedication, acknowledgments, and ISBN.",
        "Final Bible permissions statement.",
        "Final copyedit of all three manuscripts and journals.",
        "Final approved 6 x 9 devotional and journal upload interiors with locked page counts.",
        "Regenerated full-wrap covers from final page counts.",
        "KDP Previewer pass for each upload file.",
        "Physical proof review before public release.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Recommended Next Production Step", level=1)
    doc.add_paragraph(
        "Use the devotional and companion journal 6 x 9 review drafts as the next trilogy-wide copyedit and theological proof surface. Once each volume and journal is approved, regenerate the full-wrap covers from the locked page counts and paper type."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Lady D KDP Interior Finalization Kit")

    out = OUT / "lady-d-kdp-interior-finalization-kit.docx"
    doc.save(out)
    return out


def convert_pdf(docx_path: Path) -> Path:
    subprocess.check_call(["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUT), str(docx_path)])
    return OUT / f"{docx_path.stem}.pdf"


def update_zip(extra_paths: list[Path]) -> Path:
    zip_path = OUT / "Lady-D-KDP-Interior-Finalization-Kit.zip"
    existing = [p for p in OUT.glob("*") if p.is_file() and p.name != zip_path.name]
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(set(existing + extra_paths)):
            zf.write(path, path.name)
    return zip_path


def sync(paths: list[Path]) -> None:
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    SHARED_LIBRARY.mkdir(parents=True, exist_ok=True)
    for path in paths:
        shutil.copy2(path, PUBLIC_OUT / path.name)
        shutil.copy2(path, SHARED_LIBRARY / path.name)


def main() -> None:
    docx_path = build_docx()
    pdf_path = convert_pdf(docx_path)
    zip_path = update_zip([docx_path, pdf_path])
    sync([docx_path, pdf_path, zip_path])
    print({"docx": str(docx_path.relative_to(ROOT)), "pdf": str(pdf_path.relative_to(ROOT)), "zip": str(zip_path.relative_to(ROOT))})


if __name__ == "__main__":
    main()
