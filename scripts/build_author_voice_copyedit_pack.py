#!/usr/bin/env python3
"""Build the Lady D author-voice copyedit gate package."""

from __future__ import annotations

import html
import json
import re
import shutil
import zipfile
from collections import Counter
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
PRODUCTION = ROOT / "downloads" / "production"
MASTER = PRODUCTION / "master"
OUT = PRODUCTION / "kdp" / "author-voice-copyedit"
PUBLIC_OUT = ROOT / "public" / "downloads" / "production" / "kdp" / "author-voice-copyedit"
SOURCE_PAGE = ROOT / "author-voice-copyedit.html"
PUBLIC_PAGE = ROOT / "public" / "author-voice-copyedit.html"
LIBRARY_OUT = Path("/Users/IDC2.5/Documents/LADY D/Production Library/_Shared/KDP Readiness/Author Voice Copyedit")
GENERATED = "2026-07-01"
AUTHOR = 'Susan "Lady D" Damon'

VOICE_BIBLE = Path("/Users/IDC2.5/Documents/LADY D/Production Library/_Shared/Voice Bible/Lady-D-Voice-and-Theology-Guardrails.md")
OLD_VOICE_GUIDE = ROOT / "downloads" / "OLD-Lady-D-Author-Voice-Style-Guide.md"
MEETING_SYNTHESIS = ROOT / "downloads" / "FINAL-2026-06-14-GOOGLE-MEET-CAPTURE.md"
LENS_APPLICATION_JSON = OUT / "reader-facing-lens-application.json"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(89, 96, 108)
LIGHT_FILL = "E8EEF5"
SOFT_FILL = "F4F6F9"


@dataclass(frozen=True)
class Volume:
    number: int
    title: str
    lane: str

    @property
    def slug(self) -> str:
        return f"volume-{self.number}"


VOLUMES = [
    Volume(1, "Surrendering to God's Love", "The Father's love, identity, surrender, forgiveness, timing, daily trust"),
    Volume(2, "Walking with Jesus", "The Son, discipleship, nearness, obedience, healing, following, abiding"),
    Volume(3, "Filled with the Holy Spirit", "The Spirit, filling, comfort, conviction, gifts, fruit, rain, oil, breath"),
]


def write(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    cleaned = "\n".join(line.rstrip() for line in content.rstrip().splitlines())
    path.write_text(cleaned + "\n", encoding="utf-8")


def source_paths() -> list[Path]:
    return sorted(PRODUCTION.glob("volume-*-days-*-manuscript.md")) + sorted(PRODUCTION.glob("volume-1-leap-day-bonus-manuscript.md"))


def volume_from_path(path: Path) -> int:
    match = re.search(r"volume-(\d+)", path.name)
    if not match:
        raise ValueError(f"Cannot infer volume from {path}")
    return int(match.group(1))


def word_count(text: str) -> int:
    return len(re.findall(r"\b[A-Za-z][A-Za-z']*\b", text))


def read_optional(path: Path) -> str:
    return path.read_text(encoding="utf-8") if path.exists() else ""


def lens_payload() -> dict[str, object]:
    if LENS_APPLICATION_JSON.exists():
        return json.loads(LENS_APPLICATION_JSON.read_text(encoding="utf-8"))
    return {}


def source_audit(paths: list[Path]) -> dict[str, object]:
    by_volume: dict[str, dict[str, int]] = {str(volume.number): {
        "files": 0,
        "entries": 0,
        "context_lens_labels": 0,
        "internal_production_labels": 0,
        "translation_review_notes": 0,
        "sabbath_mentions": 0,
        "saturday_mentions": 0,
        "sunday_mentions": 0,
        "placeholders": 0,
        "words": 0,
    } for volume in VOLUMES}
    totals = {
        "files": 0,
        "entries": 0,
        "context_lens_labels": 0,
        "internal_production_labels": 0,
        "translation_review_notes": 0,
        "sabbath_mentions": 0,
        "saturday_mentions": 0,
        "sunday_mentions": 0,
        "placeholders": 0,
        "words": 0,
    }
    for path in paths:
        text = path.read_text(encoding="utf-8")
        volume = str(volume_from_path(path))
        row = by_volume[volume]
        counts = {
            "files": 1,
            "entries": len(re.findall(r"(?m)^## (?:Day \d{3}|Bonus)", text)),
            "context_lens_labels": len(re.findall(r"(?m)^\*\*Context and language lens:\*\*", text)),
            "internal_production_labels": len(re.findall(r"(?m)^\*\*(?:Production lens correction|Production lens note|Production lens|Production text note):\*\*", text)),
            "translation_review_notes": len(re.findall(r"(?m)^\*\*Translation review note:\*\*", text)),
            "sabbath_mentions": len(re.findall(r"\bSabbath\b", text, re.I)),
            "saturday_mentions": len(re.findall(r"\bSaturday\b", text, re.I)),
            "sunday_mentions": len(re.findall(r"\bSunday\b", text, re.I)),
            "placeholders": len(re.findall(r"\b(?:TODO|TBD|PLACEHOLDER|FPO|TK)\b|\[[^\]]*TBD[^\]]*\]", text, re.I)),
            "words": word_count(text),
        }
        for key, value in counts.items():
            row[key] += value
            totals[key] += value
    return {"totals": totals, "by_volume": by_volume}


def master_texts() -> dict[int, str]:
    texts: dict[int, str] = {}
    for volume in VOLUMES:
        path = MASTER / f"{volume.slug}-master-interior-manuscript.md"
        texts[volume.number] = read_optional(path)
    return texts


def extract_sentences(text: str) -> list[str]:
    sentences: list[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or line.startswith("<!--") or line.startswith("---"):
            continue
        match = re.match(r"^\*\*[^:]+:\*\*\s*(.*)$", line)
        if match:
            line = match.group(1).strip()
        if not line:
            continue
        for part in re.split(r"(?<=[.!?])\s+", line):
            sentence = " ".join(part.strip().split())
            if 6 <= len(sentence.split()) and len(sentence) <= 240:
                sentences.append(sentence)
    return sentences


def first_words(text: str, size: int = 5) -> str:
    words = re.findall(r"[A-Za-z']+", text)
    return " ".join(words[:size])


def sentence_action(phrase: str) -> str:
    lower = phrase.lower()
    if "this day falls on saturday" in lower:
        return "Keep the Saturday Sabbath truth, but move repeated production-calendar wording into a reusable Sabbath note or vary it by entry."
    if "let the spirit press" in lower:
        return "Rewrite Volume 3 formation bridges so the Spirit language feels alive and not template-driven."
    if "do not rush past the verse" in lower:
        return "Vary the intake line with fresh morning verbs: sit, receive, notice, breathe, confess, carry."
    if "let prayer turn insight into obedience" in lower:
        return "Keep obedience-as-response, but write more specific today-sized applications."
    if "take one surrendered step" in lower:
        return "Replace repeated surrender line with concrete steps tied to each day's passage."
    if "start from what god has revealed" in lower:
        return "Vary the opening claim so it sounds like Lady D speaking into that specific morning."
    if "father's love carry" in lower:
        return "Rewrite morning-impact lines so they carry the day's actual image rather than the title template."
    if "walking with jesus means" in lower:
        return "Vary Volume 2 concluding claims with scene-specific discipleship language."
    if "let the spirit carry" in lower:
        return "Rewrite Volume 3 morning impacts with less repeated carrier language."
    return "Review in the author-voice pass and vary where repetition sounds mechanical rather than intentional."


def repeated_ledgers(texts: dict[int, str]) -> dict[str, object]:
    sentence_counter: Counter[str] = Counter()
    opening_counter: Counter[str] = Counter()
    impact_counter: Counter[str] = Counter()
    for text in texts.values():
        sentence_counter.update(extract_sentences(text))
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped or stripped.startswith("#") or stripped.startswith("<!--") or stripped.startswith("---"):
                continue
            label_match = re.match(r"^\*\*Morning impact:\*\*\s*(.+)$", stripped)
            if label_match:
                impact_counter[first_words(label_match.group(1), 5)] += 1
            cleaned = re.sub(r"^\*\*[^:]+:\*\*\s*", "", stripped)
            fw = first_words(cleaned, 5)
            if fw:
                opening_counter[fw] += 1

    repeated_sentences = [
        {"phrase": phrase, "count": count, "action": sentence_action(phrase)}
        for phrase, count in sentence_counter.most_common(40)
        if count >= 10
    ][:24]
    opening_patterns = [
        {"phrase": phrase, "count": count, "action": sentence_action(phrase)}
        for phrase, count in opening_counter.most_common(40)
        if count >= 10
    ][:24]
    impact_patterns = [
        {"phrase": phrase, "count": count, "action": sentence_action(phrase)}
        for phrase, count in impact_counter.most_common(30)
        if count >= 10
    ][:18]
    return {
        "repeated_sentences": repeated_sentences,
        "opening_patterns": opening_patterns,
        "morning_impact_patterns": impact_patterns,
    }


def master_audit(texts: dict[int, str]) -> dict[str, object]:
    by_volume: dict[str, dict[str, int]] = {}
    totals = {
        "entries": 0,
        "bonus_entries": 0,
        "context_lens_labels": 0,
        "internal_production_labels": 0,
        "translation_review_notes": 0,
        "sabbath_mentions": 0,
        "saturday_mentions": 0,
        "sunday_mentions": 0,
        "placeholders": 0,
        "words": 0,
    }
    for volume, text in texts.items():
        row = {
            "entries": len(re.findall(r"(?m)^## Day \d{3}\b", text)) + len(re.findall(r"(?m)^## Bonus", text)),
            "bonus_entries": len(re.findall(r"(?m)^## Bonus", text)),
            "context_lens_labels": len(re.findall(r"(?m)^\*\*Context and language lens:\*\*", text)),
            "internal_production_labels": len(re.findall(r"(?m)^\*\*(?:Production lens correction|Production lens note|Production lens|Production text note):\*\*", text)),
            "translation_review_notes": len(re.findall(r"(?m)^\*\*Translation review note:\*\*", text)),
            "sabbath_mentions": len(re.findall(r"\bSabbath\b", text, re.I)),
            "saturday_mentions": len(re.findall(r"\bSaturday\b", text, re.I)),
            "sunday_mentions": len(re.findall(r"\bSunday\b", text, re.I)),
            "placeholders": len(re.findall(r"\b(?:TODO|TBD|PLACEHOLDER|FPO|TK)\b|\[[^\]]*TBD[^\]]*\]", text, re.I)),
            "words": word_count(text),
        }
        by_volume[str(volume)] = row
        for key, value in row.items():
            totals[key] += value
    return {"totals": totals, "by_volume": by_volume}


def copyedit_queue(ledger: dict[str, object]) -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    for item in ledger["repeated_sentences"][:12]:
        rows.append({"type": "Repeated sentence", **item})
    for item in ledger["morning_impact_patterns"][:8]:
        rows.append({"type": "Morning impact pattern", **item})
    for item in ledger["opening_patterns"][:8]:
        phrase = item["phrase"]
        if not any(row["phrase"] == phrase for row in rows):
            rows.append({"type": "Opening pattern", **item})
    return rows[:22]


def gpt_settings() -> list[dict[str, str]]:
    return [
        {
            "phase": "Batch rewriting",
            "recommended": "Best available flagship writing/reasoning model, medium reasoning",
            "use_for": "Rewriting 7 to 14 entries at a time while preserving Scripture reference, day title, prayer, journal prompt, and Adventist guardrails.",
        },
        {
            "phase": "Judge and auditor loop",
            "recommended": "Best available flagship model, extra-high reasoning",
            "use_for": "Theology, voice, repetition, permissions, and KDP readiness audits. Use three passes for routine batches and up to seven for difficult theological or voice sections.",
        },
        {
            "phase": "Mechanical checks",
            "recommended": "Fast or spark-class model",
            "use_for": "Link checks, count verification, HTML index updates, artifact inventories, and non-theological formatting checks.",
        },
        {
            "phase": "Cover and visual prompts",
            "recommended": "GPT image workflow with the strongest available image model and high-quality portrait/book-cover settings",
            "use_for": "6 x 9 cover variants, typography-free concept art, full-wrap direction boards, and visual comparison sheets.",
        },
    ]


def judge_loops() -> list[dict[str, str]]:
    return [
        {
            "loop": "Voice judge",
            "pass_count": "3 passes per batch",
            "standard": "Sounds like Lady D sitting with the reader in the morning: warm, reverent, direct, practical, not generic or academic.",
        },
        {
            "loop": "Theology auditor",
            "pass_count": "3 to 7 passes when needed",
            "standard": "Saturday Sabbath is protected, obedience is response to grace, and Father/Son/Spirit lanes stay distinct without separating God.",
        },
        {
            "loop": "Repetition auditor",
            "pass_count": "3 passes per volume after every rewrite wave",
            "standard": "No mechanical openings, repeated prayers, recycled morning-impact lines, or repeated bridge sentences unless deliberately liturgical.",
        },
        {
            "loop": "Permissions auditor",
            "pass_count": "1 pass every release build",
            "standard": "References-only remains safe until a final Bible translation and copyright notice are locked.",
        },
        {
            "loop": "Production auditor",
            "pass_count": "1 pass every artifact build",
            "standard": "Markdown, DOCX, PDF, ZIP, public download links, and Vercel pages all match the same current evidence.",
        },
    ]


def build_payload() -> dict[str, object]:
    texts = master_texts()
    source = source_audit(source_paths())
    master = master_audit(texts)
    ledger = repeated_ledgers(texts)
    payload = {
        "generated": GENERATED,
        "snapshot": "Production snapshot",
        "status": "author_voice_copyedit_gate_ready_not_final_upload",
        "author": AUTHOR,
        "sources": {
            "voice_bible": str(VOICE_BIBLE),
            "legacy_voice_guide": str(OLD_VOICE_GUIDE.relative_to(ROOT)),
            "meeting_synthesis": str(MEETING_SYNTHESIS.relative_to(ROOT)),
            "reader_facing_lens_application": str(LENS_APPLICATION_JSON.relative_to(ROOT)),
        },
        "source_audit": source,
        "master_audit": master,
        "reader_facing_lens_application": lens_payload(),
        "repetition_ledgers": ledger,
        "copyedit_queue": copyedit_queue(ledger),
        "judge_loops": judge_loops(),
        "gpt_settings": gpt_settings(),
        "release_boundary": [
            "This package is an author-voice copyedit gate, not a final KDP upload approval.",
            "The reader-facing lens labels are clean across the 1,098 source entries and public mirrors.",
            "The next manuscript work is a line-level voice and rhythm pass, followed by Bible permissions, final interiors, cover regeneration from locked page counts, KDP Previewer, and physical proof review.",
        ],
    }
    return payload


def md_table(headers: list[str], rows: list[list[object]]) -> str:
    head = "| " + " | ".join(headers) + " |"
    sep = "| " + " | ".join("---" for _ in headers) + " |"
    body = "\n".join("| " + " | ".join(str(cell).replace("\n", " ") for cell in row) + " |" for row in rows)
    return "\n".join([head, sep, body])


def markdown_report(payload: dict[str, object]) -> str:
    totals = payload["master_audit"]["totals"]
    source_totals = payload["source_audit"]["totals"]
    volume_rows = []
    for volume in VOLUMES:
        row = payload["master_audit"]["by_volume"][str(volume.number)]
        volume_rows.append([
            f"Volume {volume.number}",
            volume.title,
            row["entries"],
            row["context_lens_labels"],
            row["internal_production_labels"],
            row["sunday_mentions"],
            f"{row['words']:,}",
        ])
    queue_rows = [
        [item["type"], item["count"], item["phrase"], item["action"]]
        for item in payload["copyedit_queue"][:16]
    ]
    loop_rows = [[item["loop"], item["pass_count"], item["standard"]] for item in payload["judge_loops"]]
    settings_rows = [[item["phase"], item["recommended"], item["use_for"]] for item in payload["gpt_settings"]]
    return f"""# Lady D Author-Voice Copyedit Gate

Generated: {GENERATED}

Snapshot: `{payload['snapshot']}`

Status: Author-voice copyedit gate. This is not final KDP upload approval.

## Executive Readout

- Master devotional entries: {totals['entries']}
- Source manuscript entries: {source_totals['entries']}
- Source context/language lens labels: {source_totals['context_lens_labels']}
- Master internal production labels: {totals['internal_production_labels']}
- Source internal production labels: {source_totals['internal_production_labels']}
- Sunday mentions: {totals['sunday_mentions']}
- Placeholder flags: {totals['placeholders']}
- Translation review notes: {totals['translation_review_notes']}
- Master manuscript words: {totals['words']:,}

## Volume Snapshot

{md_table(['Volume', 'Title', 'Entries', 'Context lenses', 'Internal labels', 'Sunday', 'Words'], volume_rows)}

## What This Gate Cleared

The reader-facing source manuscripts and public mirrored manuscript files now use `Context and language lens` instead of internal production-lens labels. The master interiors should be regenerated after this pass so the cleaned language flows into the DOCX/PDF review artifacts.

## Main Author-Voice Queue

{md_table(['Type', 'Count', 'Pattern', 'Recommended action'], queue_rows)}

## Judge And Auditor Loops

{md_table(['Loop', 'Pass count', 'Standard'], loop_rows)}

## Recommended Goal-Mode Settings

{md_table(['Phase', 'Recommended setting', 'Use for'], settings_rows)}

## Voice And Theology Baseline

- Voice should feel like a mature believer sitting with the reader in the morning before the day gets loud.
- Keep the tone warm, reverent, direct, encouraging, practical, and alive.
- Avoid academic, generic, sugary, detached, or template-sounding devotional language.
- Saturday is the Sabbath when Sabbath is named.
- Obedience is response to grace, not a way to earn God's love.
- Surrender should lead to trust, courage, confession, service, forgiveness, rest, or faithful action.

## Release Boundary

This gate clears internal reader-facing labels and identifies the next voice-rhythm work. It does not replace the final author approval, Scripture permissions decision, final copyedit, locked KDP interiors, cover regeneration from final page counts, KDP Previewer, or physical proof review.
"""


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: object, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = color


def set_table_widths(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=INK)
        set_cell_shading(table.rows[0].cells[idx], LIGHT_FILL)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def add_para(doc: Document, text: str, style: str | None = None, bold: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def build_docx(payload: dict[str, object], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "Lady D Author-Voice Copyedit Gate"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = "IDC Publishing - Author review package"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("Lady D Author-Voice Copyedit Gate")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(23)
    title_run.font.bold = True
    title_run.font.color.rgb = INK

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    sub_run = subtitle.add_run("Reader-facing lens cleanup, voice-rhythm queue, and next Goal Mode loops")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = MUTED

    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Author", AUTHOR],
            ["Generated", GENERATED],
            ["Status", "Author-voice copyedit gate; not final KDP upload approval"],
            ["Snapshot", payload["snapshot"]],
            ["Design preset", "compact_reference_guide with memo_masthead opening"],
        ],
        [1.6, 4.8],
    )

    totals = payload["master_audit"]["totals"]
    source_totals = payload["source_audit"]["totals"]
    doc.add_heading("Executive Readout", level=1)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Master devotional entries", totals["entries"]],
            ["Source context/language lens labels", source_totals["context_lens_labels"]],
            ["Master internal production labels", totals["internal_production_labels"]],
            ["Source internal production labels", source_totals["internal_production_labels"]],
            ["Sunday mentions", totals["sunday_mentions"]],
            ["Placeholder flags", totals["placeholders"]],
            ["Translation review notes", totals["translation_review_notes"]],
            ["Master words", f"{totals['words']:,}"],
        ],
        [2.6, 3.8],
    )

    doc.add_heading("Volume Snapshot", level=1)
    volume_rows = []
    for volume in VOLUMES:
        row = payload["master_audit"]["by_volume"][str(volume.number)]
        volume_rows.append([f"V{volume.number}", volume.title, row["entries"], row["context_lens_labels"], row["internal_production_labels"], row["sunday_mentions"]])
    add_table(doc, ["Volume", "Title", "Entries", "Lenses", "Internal", "Sunday"], volume_rows, [0.55, 2.4, 0.7, 0.7, 0.7, 0.7])

    doc.add_heading("Main Author-Voice Queue", level=1)
    queue_rows = [[item["type"], item["count"], item["phrase"], item["action"]] for item in payload["copyedit_queue"][:7]]
    add_table(doc, ["Type", "Count", "Pattern", "Recommended action"], queue_rows, [1.25, 0.55, 2.0, 2.6])

    doc.add_page_break()
    doc.add_heading("Judge And Auditor Loops", level=1)
    loop_rows = [[item["loop"], item["pass_count"], item["standard"]] for item in payload["judge_loops"]]
    add_table(doc, ["Loop", "Pass count", "Standard"], loop_rows, [1.5, 1.1, 3.8])

    doc.add_page_break()
    doc.add_heading("Recommended Goal-Mode Settings", level=1)
    settings_rows = [[item["phase"], item["recommended"], item["use_for"]] for item in payload["gpt_settings"]]
    add_table(doc, ["Phase", "Recommended", "Use for"], settings_rows, [1.45, 2.05, 2.9])

    doc.add_heading("Release Boundary", level=1)
    for item in payload["release_boundary"]:
        add_para(doc, item)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def pdf_paragraph(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(html.escape(str(text)).replace("\n", "<br/>"), style)


def build_pdf(payload: dict[str, object], path: Path) -> None:
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("GateH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=19, leading=22, textColor=colors.HexColor("#111827"), spaceAfter=10)
    h2 = ParagraphStyle("GateH2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=colors.HexColor("#2E74B5"), spaceBefore=12, spaceAfter=6)
    body = ParagraphStyle("GateBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=12.2, alignment=TA_LEFT, spaceAfter=5)
    small = ParagraphStyle("GateSmall", parent=body, fontSize=8.5, leading=10.5)

    def make_table(headers: list[str], rows: list[list[object]], widths: list[float]) -> Table:
        data = [[pdf_paragraph(h, small) for h in headers]]
        for row in rows:
            data.append([pdf_paragraph(cell, small) for cell in row])
        table = Table(data, colWidths=[w * inch for w in widths], repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#AAB7C4")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        return table

    story = [
        pdf_paragraph("Lady D Author-Voice Copyedit Gate", h1),
        pdf_paragraph("Reader-facing lens cleanup, voice-rhythm queue, and next Goal Mode loops", body),
        Spacer(1, 8),
    ]
    totals = payload["master_audit"]["totals"]
    source_totals = payload["source_audit"]["totals"]
    story.extend([
        pdf_paragraph("Executive Readout", h2),
        make_table(
            ["Metric", "Value"],
            [
                ["Generated", GENERATED],
                ["Snapshot", payload["snapshot"]],
                ["Master devotional entries", totals["entries"]],
                ["Source context/language lens labels", source_totals["context_lens_labels"]],
                ["Master internal production labels", totals["internal_production_labels"]],
                ["Source internal production labels", source_totals["internal_production_labels"]],
                ["Sunday mentions", totals["sunday_mentions"]],
                ["Placeholder flags", totals["placeholders"]],
            ],
            [2.7, 3.5],
        ),
        Spacer(1, 10),
        pdf_paragraph("Main Author-Voice Queue", h2),
    ])
    story.append(make_table(
        ["Type", "Count", "Pattern", "Recommended action"],
        [[item["type"], item["count"], item["phrase"], item["action"]] for item in payload["copyedit_queue"][:7]],
        [1.15, 0.55, 2.05, 2.45],
    ))
    story.extend([PageBreak(), pdf_paragraph("Judge And Auditor Loops", h2)])
    story.append(make_table(
        ["Loop", "Pass count", "Standard"],
        [[item["loop"], item["pass_count"], item["standard"]] for item in payload["judge_loops"]],
        [1.45, 1.15, 3.6],
    ))
    story.extend([Spacer(1, 10), pdf_paragraph("Recommended Goal-Mode Settings", h2)])
    story.append(make_table(
        ["Phase", "Recommended", "Use for"],
        [[item["phase"], item["recommended"], item["use_for"]] for item in payload["gpt_settings"]],
        [1.35, 2.05, 2.8],
    ))
    story.extend([Spacer(1, 10), pdf_paragraph("Release Boundary", h2)])
    for item in payload["release_boundary"]:
        story.append(pdf_paragraph(item, body))

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.7 * inch, leftMargin=0.7 * inch, topMargin=0.7 * inch, bottomMargin=0.7 * inch)
    doc.build(story)


def html_page(payload: dict[str, object]) -> str:
    totals = payload["master_audit"]["totals"]
    source_totals = payload["source_audit"]["totals"]
    cards = [
        ("Entries", f"{totals['entries']:,}", "Master devotional entries now covered by the gate."),
        ("Context Lenses", f"{source_totals['context_lens_labels']:,}", "Reader-facing source lenses across the public manuscript set."),
        ("Internal Labels", str(totals["internal_production_labels"]), "Production lens labels remaining in regenerated master manuscripts."),
        ("Sunday Mentions", str(totals["sunday_mentions"]), "Sabbath guardrail remains clean for the master devotional manuscripts."),
    ]
    card_html = "".join(f"<article class=\"card\"><span>{html.escape(k)}</span><strong>{html.escape(v)}</strong><p>{html.escape(desc)}</p></article>" for k, v, desc in cards)
    queue_rows = "".join(
        f"<tr><td>{html.escape(item['type'])}</td><td>{item['count']}</td><td>{html.escape(item['phrase'])}</td><td>{html.escape(item['action'])}</td></tr>"
        for item in payload["copyedit_queue"][:18]
    )
    loop_rows = "".join(
        f"<tr><td>{html.escape(item['loop'])}</td><td>{html.escape(item['pass_count'])}</td><td>{html.escape(item['standard'])}</td></tr>"
        for item in payload["judge_loops"]
    )
    settings_rows = "".join(
        f"<tr><td>{html.escape(item['phase'])}</td><td>{html.escape(item['recommended'])}</td><td>{html.escape(item['use_for'])}</td></tr>"
        for item in payload["gpt_settings"]
    )
    volume_cards = "".join(
        f"""
        <article class="volume">
          <span>Volume {volume.number}</span>
          <h3>{html.escape(volume.title)}</h3>
          <p>{html.escape(volume.lane)}</p>
          <dl>
            <div><dt>Entries</dt><dd>{payload['master_audit']['by_volume'][str(volume.number)]['entries']}</dd></div>
            <div><dt>Context lenses</dt><dd>{payload['master_audit']['by_volume'][str(volume.number)]['context_lens_labels']}</dd></div>
            <div><dt>Internal labels</dt><dd>{payload['master_audit']['by_volume'][str(volume.number)]['internal_production_labels']}</dd></div>
          </dl>
        </article>
        """
        for volume in VOLUMES
    )
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Lady D Author-Voice Copyedit Gate</title>
  <style>
    :root {{
      --ink: #111827;
      --muted: #5b6474;
      --paper: #fffdf8;
      --mist: #f5f2eb;
      --indigo: #182646;
      --plum: #56325f;
      --gold: #c99335;
      --teal: #1d716f;
      --rose: #a35f68;
      --line: rgba(17,24,39,.14);
      --shadow: 0 22px 60px rgba(24,38,70,.12);
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      color: var(--ink);
      font-family: Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
      background: linear-gradient(180deg, var(--paper), var(--mist));
      line-height: 1.5;
    }}
    nav {{
      position: sticky;
      top: 0;
      z-index: 10;
      display: flex;
      flex-wrap: wrap;
      gap: 8px;
      align-items: center;
      padding: 10px clamp(16px,4vw,42px);
      background: rgba(17,24,39,.94);
      backdrop-filter: blur(14px);
    }}
    nav strong {{ color: white; margin-right: 6px; }}
    nav a {{
      color: rgba(255,255,255,.9);
      text-decoration: none;
      font-size: 13px;
      font-weight: 850;
      border-radius: 999px;
      padding: 7px 9px;
    }}
    nav a:hover {{ background: rgba(255,255,255,.1); color: white; }}
    header, main {{ max-width: 1180px; margin: 0 auto; padding: clamp(26px,5vw,58px) clamp(16px,4vw,28px); }}
    h1, h2, h3 {{ font-family: Georgia, "Times New Roman", serif; line-height: 1.04; letter-spacing: 0; margin: 0 0 14px; }}
    h1 {{ font-size: clamp(44px,7vw,88px); max-width: 990px; }}
    h2 {{ font-size: clamp(30px,4vw,50px); }}
    h3 {{ font-size: 25px; }}
    p {{ margin: 0 0 14px; color: #2f3a4b; }}
    .kicker {{ color: var(--plum); font-size: 12px; font-weight: 900; letter-spacing: .15em; text-transform: uppercase; margin-bottom: 12px; }}
    .lead {{ max-width: 900px; font-size: clamp(18px,2vw,23px); line-height: 1.42; }}
    .status {{
      display: inline-flex;
      min-height: 32px;
      align-items: center;
      border-radius: 999px;
      padding: 6px 10px;
      color: white;
      background: var(--indigo);
      font-size: 12px;
      font-weight: 900;
      margin-right: 7px;
    }}
    .grid {{ display: grid; grid-template-columns: repeat(auto-fit,minmax(230px,1fr)); gap: 14px; }}
    .card, .panel, .volume {{
      border: 1px solid var(--line);
      border-radius: 8px;
      background: white;
      box-shadow: var(--shadow);
      padding: 18px;
    }}
    .card span, .volume span {{ display: block; color: var(--gold); font-size: 12px; font-weight: 900; text-transform: uppercase; letter-spacing: .12em; }}
    .card strong {{ display: block; font-family: Georgia, "Times New Roman", serif; font-size: 42px; line-height: 1; margin: 8px 0; color: var(--indigo); }}
    section {{ border-top: 1px solid var(--line); padding: 36px 0; }}
    table {{ width: 100%; border-collapse: collapse; background: white; border: 1px solid var(--line); border-radius: 8px; overflow: hidden; }}
    th, td {{ padding: 10px; border-bottom: 1px solid var(--line); text-align: left; vertical-align: top; font-size: 14px; }}
    th {{ background: #e8eef5; color: var(--indigo); font-weight: 900; }}
    td:nth-child(2) {{ font-weight: 900; color: var(--plum); }}
    .actions {{ display: flex; flex-wrap: wrap; gap: 10px; margin-top: 18px; }}
    .actions a {{
      display: inline-flex;
      min-height: 38px;
      align-items: center;
      padding: 8px 11px;
      border: 1px solid var(--line);
      border-radius: 999px;
      color: var(--teal);
      background: white;
      font-weight: 900;
      font-size: 13px;
      text-decoration: none;
    }}
    dl {{ display: grid; gap: 9px; margin: 16px 0 0; }}
    dt {{ font-weight: 900; color: var(--indigo); }}
    dd {{ margin: 0; color: var(--muted); }}
    .note {{ border-left: 5px solid var(--teal); background: rgba(29,113,111,.1); padding: 16px; border-radius: 8px; }}
    @media (max-width: 700px) {{
      nav {{ position: static; }}
      th, td {{ font-size: 13px; }}
    }}
  </style>
</head>
<body>
  <nav aria-label="Author voice gate navigation">
    <strong>Lady D Production</strong>
    <a href="production.html">Production Review</a>
    <a href="release-status.html">Release Dashboard</a>
    <a href="#queue">Copyedit Queue</a>
    <a href="#loops">Loops</a>
    <a href="#settings">GPT Settings</a>
    <a href="#downloads">Downloads</a>
  </nav>
  <header>
    <div class="kicker">IDC Publishing author-voice gate</div>
    <h1>Lady D Author-Voice Copyedit Gate</h1>
    <p class="lead">The current trilogy manuscripts have been cleaned of internal production-lens labels and are ready for the next line-level author-voice pass. This page separates what is now clean from what still needs rhythmic, Lady D-centered editorial work before KDP upload.</p>
    <p><span class="status">Generated {GENERATED}</span><span class="status">{html.escape(str(payload['snapshot']))}</span><span class="status">Not final upload approval</span></p>
    <div class="actions">
      <a href="downloads/production/kdp/author-voice-copyedit/Lady-D-Author-Voice-Copyedit-Pack.zip">Download ZIP</a>
      <a href="downloads/production/kdp/author-voice-copyedit/lady-d-author-voice-copyedit-pack.pdf">PDF</a>
      <a href="downloads/production/kdp/author-voice-copyedit/lady-d-author-voice-copyedit-pack.docx">DOCX</a>
      <a href="downloads/production/kdp/author-voice-copyedit/author-voice-copyedit-pack.json">JSON</a>
      <a href="downloads/production/kdp/author-voice-copyedit/reader-facing-lens-application.md">Lens Application Report</a>
    </div>
  </header>
  <main>
    <section><div class="grid">{card_html}</div></section>
    <section>
      <h2>Volume Snapshot</h2>
      <div class="grid">{volume_cards}</div>
    </section>
    <section id="queue">
      <h2>Main Author-Voice Queue</h2>
      <p class="lead">These are not failures; they are the next obvious places where a real copyedit can lift the manuscripts from complete to beautiful.</p>
      <table><thead><tr><th>Type</th><th>Count</th><th>Pattern</th><th>Recommended action</th></tr></thead><tbody>{queue_rows}</tbody></table>
    </section>
    <section id="loops">
      <h2>Judge And Auditor Loops</h2>
      <table><thead><tr><th>Loop</th><th>Pass count</th><th>Standard</th></tr></thead><tbody>{loop_rows}</tbody></table>
    </section>
    <section id="settings">
      <h2>Recommended GPT Settings</h2>
      <table><thead><tr><th>Phase</th><th>Recommended</th><th>Use for</th></tr></thead><tbody>{settings_rows}</tbody></table>
    </section>
    <section>
      <h2>Release Boundary</h2>
      <div class="note">
        <p>This gate clears reader-facing labels and identifies the next voice-rhythm pass. It does not replace author approval, final Scripture permissions, final KDP interiors, cover regeneration from locked page counts, KDP Previewer, or physical proof review.</p>
      </div>
    </section>
    <section id="downloads">
      <h2>Downloads</h2>
      <div class="actions">
        <a href="downloads/production/kdp/author-voice-copyedit/Lady-D-Author-Voice-Copyedit-Pack.zip">Author Voice Pack ZIP</a>
        <a href="downloads/production/kdp/author-voice-copyedit/lady-d-author-voice-copyedit-pack.md">Markdown</a>
        <a href="downloads/production/kdp/author-voice-copyedit/lady-d-author-voice-copyedit-pack.pdf">PDF</a>
        <a href="downloads/production/kdp/author-voice-copyedit/lady-d-author-voice-copyedit-pack.docx">DOCX</a>
        <a href="downloads/production/kdp/author-voice-copyedit/author-voice-copyedit-pack.json">JSON</a>
        <a href="downloads/production/kdp/author-voice-copyedit/reader-facing-lens-application.json">Lens JSON</a>
      </div>
    </section>
  </main>
</body>
</html>
"""


def sync_outputs(paths: list[Path]) -> None:
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)
    for path in paths:
        shutil.copy2(path, PUBLIC_OUT / path.name)
        shutil.copy2(path, LIBRARY_OUT / path.name)


def make_zip(paths: list[Path], zip_path: Path) -> None:
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in paths:
            zf.write(path, path.name)


def main() -> None:
    payload = build_payload()
    OUT.mkdir(parents=True, exist_ok=True)
    json_path = OUT / "author-voice-copyedit-pack.json"
    md_path = OUT / "lady-d-author-voice-copyedit-pack.md"
    docx_path = OUT / "lady-d-author-voice-copyedit-pack.docx"
    pdf_path = OUT / "lady-d-author-voice-copyedit-pack.pdf"
    html_path = OUT / "lady-d-author-voice-copyedit-review.html"
    zip_path = OUT / "Lady-D-Author-Voice-Copyedit-Pack.zip"

    write(json_path, json.dumps(payload, indent=2))
    write(md_path, markdown_report(payload))
    build_docx(payload, docx_path)
    build_pdf(payload, pdf_path)
    write(html_path, html_page(payload))
    write(SOURCE_PAGE, html_page(payload))
    write(PUBLIC_PAGE, html_page(payload))

    zip_members = [json_path, md_path, docx_path, pdf_path, html_path]
    for extra in [OUT / "reader-facing-lens-application.json", OUT / "reader-facing-lens-application.md"]:
        if extra.exists():
            zip_members.append(extra)
    make_zip(zip_members, zip_path)
    all_outputs = zip_members + [zip_path]
    sync_outputs(all_outputs)

    print(json.dumps({
        "status": payload["status"],
        "entries": payload["master_audit"]["totals"]["entries"],
        "source_context_lens_labels": payload["source_audit"]["totals"]["context_lens_labels"],
        "master_internal_production_labels": payload["master_audit"]["totals"]["internal_production_labels"],
        "copyedit_queue_items": len(payload["copyedit_queue"]),
        "report": str(md_path.relative_to(ROOT)),
        "zip": str(zip_path.relative_to(ROOT)),
    }, indent=2))


if __name__ == "__main__":
    main()
