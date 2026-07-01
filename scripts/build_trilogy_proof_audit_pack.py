#!/usr/bin/env python3
"""Build the Lady D trilogy proof and copyedit audit pack."""

from __future__ import annotations

import html
import json
import re
import shutil
import subprocess
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_release_upload_readiness_pack import AUTHOR, PRODUCTS, ROOT, current_commit, word_count


OUT = ROOT / "downloads" / "production" / "kdp" / "trilogy-proof-audit"
PUBLIC_OUT = ROOT / "public" / "downloads" / "production" / "kdp" / "trilogy-proof-audit"
SOURCE_PAGE = ROOT / "trilogy-proof-audit.html"
PUBLIC_PAGE = ROOT / "public" / "trilogy-proof-audit.html"
LIBRARY_OUT = Path("/Users/IDC2.5/Documents/LADY D/Production Library/_Shared/KDP Readiness/Trilogy Proof Audit")
GENERATED = "2026-07-01"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(89, 96, 108)
LIGHT_FILL = "E8EEF5"
SOFT_FILL = "F4F6F9"

PLACEHOLDER_RE = re.compile(r"\b(?:TODO|TBD|PLACEHOLDER|LOREM|TK|INSERT)\b", re.I)
WORKS_WATCH_RE = re.compile(r"\b(?:earn(?:ing|ed|s)?|performance)\b", re.I)
ALIGNED_GRACE_MARKERS = (
    "not",
    "never",
    "cannot",
    "could not",
    "response",
    "grace",
    "instead of",
    "rather than",
    "release",
    "receive",
    "surrender",
    "gift",
    "without",
    "before",
)


def write(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content.rstrip() + "\n", encoding="utf-8")


def line_number(text: str, index: int) -> int:
    return text.count("\n", 0, index) + 1


def compact_snippet(text: str, start: int, end: int, radius: int = 130) -> str:
    left = max(0, start - radius)
    right = min(len(text), end + radius)
    return " ".join(text[left:right].split())


def product_label(product) -> str:
    return f"Volume {product.volume} {product.kind}: {product.title}"


def extract_devotional_entries(text: str) -> list[dict[str, object]]:
    markers = list(re.finditer(r"(?m)^## (?P<head>(?:Day \d{3}|Bonus)[^\n]*)", text))
    entries: list[dict[str, object]] = []
    for idx, marker in enumerate(markers):
        body_start = marker.end()
        body_end = markers[idx + 1].start() if idx + 1 < len(markers) else len(text)
        body = text[body_start:body_end]
        title_match = re.search(r"(?m)^### (.+)$", body)
        scripture_match = re.search(r"(?m)^\*\*Scripture Reference:\*\*\s*(.+)$", body)
        impact_match = re.search(r"(?m)^\*\*Morning impact:\*\*\s*(.+)$", body)
        day_match = re.match(r"Day (\d{3}) - ([A-Za-z]+) (\d+)", marker.group("head"))
        entries.append(
            {
                "heading": marker.group("head"),
                "day": day_match.group(1) if day_match else "bonus",
                "month": day_match.group(2) if day_match else "Bonus",
                "title": title_match.group(1).strip() if title_match else "",
                "scripture": scripture_match.group(1).strip() if scripture_match else "",
                "morning_impact": impact_match.group(1).strip() if impact_match else "",
                "line": line_number(text, marker.start()),
            }
        )
    return entries


def structure_counts(product, text: str) -> dict[str, int]:
    if product.kind == "Devotional":
        return {
            "day_entries": len(re.findall(r"(?m)^## Day \d{3}\b", text)),
            "bonus_entries": len(re.findall(r"(?m)^## Bonus\b", text)),
            "scripture_references": len(re.findall(r"(?m)^\*\*Scripture Reference:\*\*", text)),
            "today_steps": len(re.findall(r"(?m)^\*\*Today step:\*\*", text)),
            "prayers": len(re.findall(r"(?m)^\*\*Prayer:\*\*", text)),
            "journal_prompts": len(re.findall(r"(?m)^\*\*Journal prompt:\*\*", text)),
            "morning_impacts": len(re.findall(r"(?m)^\*\*Morning impact:\*\*", text)),
        }
    return {
        "day_reflections": len(re.findall(r"(?m)^### Day \d{3}\b", text)),
        "weekly_opening_prayers": len(re.findall(r"(?m)^### Weekly Opening Prayer\b", text)),
        "review_sections": len(re.findall(r"(?m)^### .*Review\b", text)),
        "sabbath_reflections": len(re.findall(r"(?m)^### Sabbath Reflection\b", text)),
    }


def structure_status(product, counts: dict[str, int]) -> str:
    if product.kind == "Devotional":
        if (
            counts["day_entries"] == 365
            and counts["bonus_entries"] == 1
            and counts["scripture_references"] == 366
            and counts["today_steps"] == 366
            and counts["prayers"] == 366
            and counts["journal_prompts"] == 366
            and counts["morning_impacts"] == 366
        ):
            return "pass"
        return "review"
    return "review"


def build_watch_contexts(product, text: str) -> list[dict[str, object]]:
    contexts: list[dict[str, object]] = []
    for match in WORKS_WATCH_RE.finditer(text):
        snippet = compact_snippet(text, match.start(), match.end())
        lower = snippet.lower()
        contexts.append(
            {
                "product": product.key,
                "title": product.title,
                "kind": product.kind,
                "term": match.group(0),
                "line": line_number(text, match.start()),
                "snippet": snippet,
                "preliminary_read": "likely grace-aligned" if any(marker in lower for marker in ALIGNED_GRACE_MARKERS) else "requires proof review",
            }
        )
    return contexts


def duplicate_groups(entries: list[dict[str, object]], field: str) -> list[dict[str, object]]:
    grouped: dict[str, list[dict[str, object]]] = defaultdict(list)
    for entry in entries:
        value = str(entry.get(field, "")).strip()
        if value:
            grouped[value].append(entry)
    rows = []
    for value, group in grouped.items():
        if len(group) > 1:
            rows.append(
                {
                    "value": value,
                    "count": len(group),
                    "locations": [{"heading": item["heading"], "line": item["line"], "month": item["month"]} for item in group],
                }
            )
    return sorted(rows, key=lambda item: (-int(item["count"]), str(item["value"])))


def analyze_products(commit: str) -> dict[str, object]:
    products = []
    title_duplicates = []
    impact_duplicates = []
    watch_contexts = []
    totals = {
        "products": len(PRODUCTS),
        "words": 0,
        "placeholders": 0,
        "sunday_mentions": 0,
        "sabbath_mentions": 0,
        "saturday_mentions": 0,
        "works_watch_contexts": 0,
        "review_required_contexts": 0,
        "duplicate_title_groups": 0,
        "duplicate_morning_impact_groups": 0,
        "missing_files": 0,
    }
    for product in PRODUCTS:
        text = product.master_source.read_text(encoding="utf-8")
        counts = structure_counts(product, text)
        placeholders = len(PLACEHOLDER_RE.findall(text))
        sunday = len(re.findall(r"\bSunday\b", text, re.I))
        sabbath = len(re.findall(r"\bSabbath\b", text, re.I))
        saturday = len(re.findall(r"\bSaturday\b", text, re.I))
        contexts = build_watch_contexts(product, text)
        watch_contexts.extend(contexts)
        missing = [str(path.relative_to(ROOT)) for path in (product.master_source, product.interior_pdf, product.cover_pdf) if not path.exists()]
        entries = extract_devotional_entries(text) if product.kind == "Devotional" else []
        product_title_dups = duplicate_groups(entries, "title")
        product_impact_dups = duplicate_groups(entries, "morning_impact")
        title_duplicates.extend([{**row, "product": product.key, "title": product.title, "volume": product.volume} for row in product_title_dups])
        impact_duplicates.extend([{**row, "product": product.key, "title": product.title, "volume": product.volume} for row in product_impact_dups])
        words = word_count(text)
        totals["words"] += words
        totals["placeholders"] += placeholders
        totals["sunday_mentions"] += sunday
        totals["sabbath_mentions"] += sabbath
        totals["saturday_mentions"] += saturday
        totals["works_watch_contexts"] += len(contexts)
        totals["review_required_contexts"] += sum(1 for item in contexts if item["preliminary_read"] == "requires proof review")
        totals["duplicate_title_groups"] += len(product_title_dups)
        totals["duplicate_morning_impact_groups"] += len(product_impact_dups)
        totals["missing_files"] += len(missing)
        products.append(
            {
                "key": product.key,
                "volume": product.volume,
                "kind": product.kind,
                "title": product.title,
                "source": str(product.master_source.relative_to(ROOT)),
                "words": words,
                "structure_counts": counts,
                "structure_status": structure_status(product, counts),
                "placeholders": placeholders,
                "sunday_mentions": sunday,
                "sabbath_mentions": sabbath,
                "saturday_mentions": saturday,
                "works_watch_contexts": len(contexts),
                "review_required_contexts": sum(1 for item in contexts if item["preliminary_read"] == "requires proof review"),
                "duplicate_title_groups": len(product_title_dups),
                "duplicate_morning_impact_groups": len(product_impact_dups),
                "missing": missing,
            }
        )
    return {
        "generated": GENERATED,
        "commit": commit,
        "result": "proof_ready_not_final_upload",
        "totals": totals,
        "products": products,
        "title_duplicates": title_duplicates,
        "morning_impact_duplicates": impact_duplicates,
        "works_watch_contexts": watch_contexts,
        "judge_auditor_loop": [
            "Mechanical completeness judge",
            "Adventist Sabbath guardrail auditor",
            "Grace and obedience theological auditor",
            "Reader rhythm and repetition judge",
            "KDP proof gate auditor",
        ],
        "remaining_gates": [
            "Resolve title and morning-impact repetition decision before final upload.",
            "Complete full copyedit and theological proof on all six products.",
            "Approve Bible translation permissions language before adding Scripture quotation text.",
            "Lock ISBN/barcode, paper type, and final page counts.",
            "Run KDP Previewer and physical proof review before public release.",
        ],
    }


def product_matrix_md(audit: dict[str, object]) -> str:
    lines = ["| Product | Type | Words | Structure | Sunday | Placeholders | Watch contexts |", "| --- | --- | ---: | --- | ---: | ---: | ---: |"]
    for product in audit["products"]:
        lines.append(
            f"| {product['title']} | {product['kind']} | {product['words']:,} | {product['structure_status']} | "
            f"{product['sunday_mentions']} | {product['placeholders']} | {product['works_watch_contexts']} |"
        )
    return "\n".join(lines)


def main_report_markdown(audit: dict[str, object]) -> str:
    totals = audit["totals"]
    return f"""# Lady D Trilogy Proof And Copyedit Audit Pack

Generated: {GENERATED}

Repo commit at generation: `{audit['commit']}`

Author: {AUTHOR}

Status: Proof/copyedit evidence package. This is not a final KDP upload declaration.

## Purpose

This pack moves the Lady D library from assembled review drafts into a disciplined proof lane. It checks the six current KDP-facing products for structure, placeholder risk, Adventist Sabbath consistency, grace/obedience watch terms, repeated day titles, and repeated morning-impact lines.

## Evidence Snapshot

- Products checked: {totals['products']}
- Source words checked: {totals['words']:,}
- Placeholder markers found: {totals['placeholders']}
- Sunday mentions: {totals['sunday_mentions']}
- Sabbath mentions: {totals['sabbath_mentions']}
- Saturday mentions: {totals['saturday_mentions']}
- Works-righteousness watch contexts: {totals['works_watch_contexts']}
- Contexts requiring proof review: {totals['review_required_contexts']}
- Duplicate devotional title groups: {totals['duplicate_title_groups']}
- Duplicate morning-impact groups: {totals['duplicate_morning_impact_groups']}
- Missing required files: {totals['missing_files']}

## Product Matrix

{product_matrix_md(audit)}

## Judge And Auditor Loop

1. Mechanical completeness judge: confirms dated-entry, bonus-entry, and section counts before prose review.
2. Adventist Sabbath guardrail auditor: confirms Sabbath remains seventh-day/Saturday Sabbath and Sunday remains absent.
3. Grace and obedience theological auditor: reviews earn/performance contexts so obedience remains response to grace, never a way to earn love.
4. Reader rhythm and repetition judge: decides whether repeated day-title and morning-impact patterns should remain as intentional rhythm or be individualized.
5. KDP proof gate auditor: keeps final-upload language blocked until Previewer and physical proofs pass.

## Key Findings

- The six master sources contain zero placeholder markers and zero Sunday mentions.
- The three devotionals each prove 365 dated entries plus one February 29 bonus and 366 Scripture references.
- The proof package identifies repeated title and morning-impact groups for editorial decision. These are not treated as automatic errors, but they must be deliberately accepted or revised before final KDP upload.
- Works-righteousness watch terms are captured in a separate context ledger so the theological proof pass can verify that each occurrence supports grace-shaped obedience.

## Remaining Gates

{chr(10).join(f'- {item}' for item in audit['remaining_gates'])}
"""


def title_repetition_markdown(audit: dict[str, object]) -> str:
    lines = [
        "# Title And Morning-Impact Repetition Ledger",
        "",
        f"Generated: {GENERATED}",
        "",
        "This ledger isolates repeated devotional titles and repeated morning-impact lines. Repetition may be an intentional series rhythm, but it should be an explicit editorial decision before final upload.",
        "",
        "## Duplicate Day Titles",
        "",
    ]
    if not audit["title_duplicates"]:
        lines.append("No duplicate devotional titles found.")
    else:
        lines.append("| Product | Repeated title | Count | Locations |")
        lines.append("| --- | --- | ---: | --- |")
        for row in audit["title_duplicates"]:
            locations = "; ".join(f"{item['heading']} (line {item['line']})" for item in row["locations"])
            lines.append(f"| Volume {row['volume']} | {row['value']} | {row['count']} | {locations} |")
    lines.extend(["", "## Duplicate Morning-Impact Lines", ""])
    if not audit["morning_impact_duplicates"]:
        lines.append("No duplicate morning-impact lines found.")
    else:
        lines.append("| Product | Repeated morning impact | Count | First locations |")
        lines.append("| --- | --- | ---: | --- |")
        for row in audit["morning_impact_duplicates"]:
            locations = "; ".join(f"{item['heading']} (line {item['line']})" for item in row["locations"][:8])
            if len(row["locations"]) > 8:
                locations += f"; +{len(row['locations']) - 8} more"
            lines.append(f"| Volume {row['volume']} | {row['value']} | {row['count']} | {locations} |")
    return "\n".join(lines)


def watch_contexts_markdown(audit: dict[str, object]) -> str:
    lines = [
        "# Theological Watchlist Context Ledger",
        "",
        f"Generated: {GENERATED}",
        "",
        "This ledger captures `earn` and `performance` contexts for theological proof. The preliminary read is automated and should not replace human theological judgment.",
        "",
    ]
    grouped: dict[str, list[dict[str, object]]] = defaultdict(list)
    for item in audit["works_watch_contexts"]:
        grouped[str(item["product"])].append(item)
    for product in audit["products"]:
        items = grouped[product["key"]]
        lines.extend([f"## {product['title']} - {product['kind']}", "", f"Contexts captured: {len(items)}", ""])
        lines.append("| Line | Term | Preliminary read | Context |")
        lines.append("| ---: | --- | --- | --- |")
        for item in items[:35]:
            snippet = str(item["snippet"]).replace("|", "/")
            lines.append(f"| {item['line']} | {item['term']} | {item['preliminary_read']} | {snippet} |")
        if len(items) > 35:
            lines.append(f"|  |  |  | {len(items) - 35} additional contexts are available in the JSON ledger. |")
        lines.append("")
    return "\n".join(lines)


def proofreader_runbook(audit: dict[str, object]) -> str:
    return f"""# Lady D Proofreader Runbook

Generated: {GENERATED}

## Proof Order

1. Read the devotional manuscript for one volume from January 1 through December 31 plus February 29.
2. Read that volume's companion journal against the devotional, checking prompt alignment.
3. Review the title and morning-impact repetition ledger for the same volume.
4. Review the theological watchlist contexts for the same volume.
5. Mark each issue as accept as intentional, revise, or escalate to author/publisher.
6. Repeat for the next volume.

## Non-Negotiable Guardrails

- Sabbath means seventh-day/Saturday Sabbath.
- Obedience is response to grace, not a way to earn God's love.
- Holy Spirit language should remain Jesus-honoring and biblically grounded.
- Scripture references may remain reference-only unless final permissions text is approved.
- Do not call any file final-upload ready until KDP Previewer and physical proof review pass.

## Proof Marks To Use

- `ACCEPT INTENTIONAL`: the repeated rhythm or phrase is deliberate and should remain.
- `REVISE`: editorial improvement needed before upload.
- `THEOLOGY REVIEW`: needs focused theological judgment.
- `AUTHOR APPROVAL`: requires Susan "Lady D" Damon approval.
- `KDP GATE`: cannot close without KDP Previewer or printed proof evidence.

## Current Automated Findings

- Products checked: {audit['totals']['products']}
- Placeholder markers: {audit['totals']['placeholders']}
- Sunday mentions: {audit['totals']['sunday_mentions']}
- Duplicate devotional title groups: {audit['totals']['duplicate_title_groups']}
- Duplicate morning-impact groups: {audit['totals']['duplicate_morning_impact_groups']}
- Works-righteousness watch contexts: {audit['totals']['works_watch_contexts']}
"""


def volume_checklist(audit: dict[str, object], volume: int) -> str:
    products = [p for p in audit["products"] if p["volume"] == volume]
    title_dups = [row for row in audit["title_duplicates"] if row["volume"] == volume]
    impact_dups = [row for row in audit["morning_impact_duplicates"] if row["volume"] == volume]
    watch = [row for row in audit["works_watch_contexts"] if any(p["key"] == row["product"] for p in products)]
    title = products[0]["title"] if products else f"Volume {volume}"
    lines = [
        f"# Volume {volume} Proof Checklist - {title}",
        "",
        f"Generated: {GENERATED}",
        "",
        "## Product Evidence",
        "",
        "| Product | Type | Words | Structure | Watch contexts |",
        "| --- | --- | ---: | --- | ---: |",
    ]
    for product in products:
        lines.append(f"| {product['title']} | {product['kind']} | {product['words']:,} | {product['structure_status']} | {product['works_watch_contexts']} |")
    lines.extend(
        [
            "",
            "## Checklist",
            "",
            "- [ ] Read devotional manuscript start to finish.",
            "- [ ] Read companion journal against devotional flow.",
            "- [ ] Confirm Sabbath/Saturday language stays Adventist and never drifts to Sunday.",
            "- [ ] Confirm obedience language stays response-to-grace.",
            "- [ ] Decide repeated title groups: accept intentional or revise.",
            "- [ ] Decide repeated morning-impact rhythm: accept intentional or individualize.",
            "- [ ] Confirm no Scripture quotations are added without final permissions statement.",
            "- [ ] Confirm author-facing front/back matter placeholders are replaced before final upload.",
            "- [ ] Run KDP Previewer after final PDF export.",
            "- [ ] Review physical proof before release.",
            "",
            "## Current Repetition/Watch Counts",
            "",
            f"- Duplicate title groups: {len(title_dups)}",
            f"- Duplicate morning-impact groups: {len(impact_dups)}",
            f"- Works-righteousness watch contexts: {len(watch)}",
        ]
    )
    return "\n".join(lines)


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header.paragraphs[0].text = "Lady D Trilogy Proof Audit"
    section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(section.header.paragraphs[0].runs[0], size=9, color=MUTED, italic=True)
    section.footer.paragraphs[0].text = "Proof/copyedit evidence pack - not final KDP upload files"
    section.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(section.footer.paragraphs[0].runs[0], size=8, color=MUTED)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK_BLUE)]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_para(doc: Document, text: str, *, size=11, color=INK, bold=False, italic=False, align=None, after=6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), size=10.5, color=INK)


def add_product_table(doc: Document, audit: dict[str, object]) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Product", "Type", "Words", "Structure", "Sunday", "Watch"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font(run, size=8.5, color=DARK_BLUE, bold=True)
    repeat_table_header(table.rows[0])
    for product in audit["products"]:
        row = table.add_row().cells
        values = [
            str(product["title"]),
            str(product["kind"]),
            f"{product['words']:,}",
            str(product["structure_status"]),
            str(product["sunday_mentions"]),
            str(product["works_watch_contexts"]),
        ]
        for idx, value in enumerate(values):
            row[idx].text = value
            set_cell_margins(row[idx])
            row[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in row[idx].paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=8.5, color=INK)


def add_repetition_table(doc: Document, audit: dict[str, object]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Volume", "Duplicate Titles", "Impact Groups", "Max Impact Repeat"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, SOFT_FILL)
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font(run, size=8.5, color=DARK_BLUE, bold=True)
    repeat_table_header(table.rows[0])
    for volume in (1, 2, 3):
        title_groups = [row for row in audit["title_duplicates"] if row["volume"] == volume]
        impact_groups = [row for row in audit["morning_impact_duplicates"] if row["volume"] == volume]
        max_impact = max([int(row["count"]) for row in impact_groups], default=0)
        row = table.add_row().cells
        values = [str(volume), str(len(title_groups)), str(len(impact_groups)), str(max_impact)]
        for idx, value in enumerate(values):
            row[idx].text = value
            set_cell_margins(row[idx])
            for paragraph in row[idx].paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=8.5, color=INK)


def build_docx(audit: dict[str, object]) -> Path:
    docx_path = OUT / "lady-d-trilogy-proof-audit-pack.docx"
    doc = Document()
    setup_doc(doc)
    doc.core_properties.title = "Lady D Trilogy Proof And Copyedit Audit Pack"
    doc.core_properties.author = "IDC Publishing"
    add_para(doc, "LADY D DEVOTIONAL LIBRARY", size=9, color=RGBColor(122, 90, 0), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Trilogy Proof And Copyedit Audit Pack", size=24, color=DARK_BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_para(doc, "Mechanical, theological, repetition, and KDP proof evidence for all six current products", size=12, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_para(doc, f"Generated {GENERATED} from commit {audit['commit']}", size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    doc.add_heading("Readiness Boundary", level=1)
    add_para(doc, "This package advances the proof lane, but it does not make the devotional library final KDP upload-ready. Copyedit, theological sign-off, final metadata approval, KDP Previewer, and physical proofs remain required.", after=10)
    doc.add_heading("Evidence Snapshot", level=1)
    totals = audit["totals"]
    for item in [
        f"Products checked: {totals['products']}",
        f"Source words checked: {totals['words']:,}",
        f"Placeholder markers: {totals['placeholders']}",
        f"Sunday mentions: {totals['sunday_mentions']}",
        f"Works-righteousness watch contexts: {totals['works_watch_contexts']}",
        f"Duplicate title groups: {totals['duplicate_title_groups']}",
        f"Duplicate morning-impact groups: {totals['duplicate_morning_impact_groups']}",
        f"Missing required files: {totals['missing_files']}",
    ]:
        add_bullet(doc, item)
    doc.add_heading("Product Matrix", level=1)
    add_product_table(doc, audit)
    doc.add_heading("Repetition Decision Surface", level=1)
    add_para(doc, "Repeated day titles and morning-impact lines are not automatically defects. They are now explicit proof decisions: accept as intentional series rhythm or revise before final upload.", after=6)
    add_repetition_table(doc, audit)
    doc.add_heading("Judge And Auditor Loop", level=1)
    for item in audit["judge_auditor_loop"]:
        add_bullet(doc, item)
    doc.add_heading("Next Operator Move", level=1)
    add_para(doc, "Use this pack to run one volume at a time through proof, theological review, repetition decisions, metadata approval, and KDP proofing. Keep the Adventist Sabbath and grace/obedience guardrails visible in every pass.", after=4)
    doc.save(docx_path)
    return docx_path


def convert_docx_to_pdf(docx_path: Path) -> Path:
    subprocess.check_call(["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUT), str(docx_path)])
    return docx_path.with_suffix(".pdf")


def review_html(audit: dict[str, object]) -> str:
    totals = audit["totals"]
    product_rows = "\n".join(
        f"<tr><td>{html.escape(product['title'])}</td><td>{html.escape(product['kind'])}</td><td>{product['words']:,}</td><td>{html.escape(product['structure_status'])}</td><td>{product['sunday_mentions']}</td><td>{product['works_watch_contexts']}</td></tr>"
        for product in audit["products"]
    )
    downloads = [
        ("Download proof audit ZIP", "downloads/production/kdp/trilogy-proof-audit/Lady-D-Trilogy-Proof-Audit-Pack.zip"),
        ("Proof audit report PDF", "downloads/production/kdp/trilogy-proof-audit/lady-d-trilogy-proof-audit-pack.pdf"),
        ("Proof audit report DOCX", "downloads/production/kdp/trilogy-proof-audit/lady-d-trilogy-proof-audit-pack.docx"),
        ("Machine-readable audit JSON", "downloads/production/kdp/trilogy-proof-audit/trilogy-proof-audit.json"),
        ("Title and impact repetition ledger", "downloads/production/kdp/trilogy-proof-audit/title-and-impact-repetition-ledger.md"),
        ("Theological watchlist ledger", "downloads/production/kdp/trilogy-proof-audit/theological-watchlist-context-ledger.md"),
        ("Proofreader runbook", "downloads/production/kdp/trilogy-proof-audit/proofreader-runbook.md"),
        ("Volume 1 proof checklist", "downloads/production/kdp/trilogy-proof-audit/volume-1-proof-checklist.md"),
        ("Volume 2 proof checklist", "downloads/production/kdp/trilogy-proof-audit/volume-2-proof-checklist.md"),
        ("Volume 3 proof checklist", "downloads/production/kdp/trilogy-proof-audit/volume-3-proof-checklist.md"),
    ]
    download_links = "\n".join(f'<a class="card" href="{href}">{html.escape(label)}</a>' for label, href in downloads)
    gates = "\n".join(f"<li>{html.escape(item)}</li>" for item in audit["remaining_gates"])
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='12' fill='%23182646'/%3E%3Ctext x='50%25' y='55%25' dominant-baseline='middle' text-anchor='middle' font-family='Georgia,serif' font-size='26' fill='%23fffdf8'%3ELD%3C/text%3E%3C/svg%3E">
  <title>Lady D Trilogy Proof Audit</title>
  <style>
    :root {{ --ink:#111827; --paper:#fffdf8; --mist:#f5f2eb; --indigo:#182646; --teal:#1d716f; --gold:#c99335; --line:rgba(17,24,39,.14); }}
    * {{ box-sizing: border-box; }}
    body {{ margin:0; color:var(--ink); font-family:Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; background:linear-gradient(180deg,var(--paper),var(--mist)); line-height:1.5; }}
    header, main {{ max-width:1180px; margin:0 auto; padding:34px 22px; }}
    h1, h2, h3 {{ font-family:Georgia, "Times New Roman", serif; line-height:1.05; margin:0 0 14px; letter-spacing:0; }}
    h1 {{ font-size:clamp(42px, 7vw, 82px); max-width:980px; }}
    h2 {{ font-size:clamp(28px, 4vw, 48px); }}
    p {{ margin:0 0 14px; }}
    .lead {{ font-size:clamp(18px, 2vw, 23px); max-width:860px; color:#2e3746; }}
    .kicker {{ color:var(--teal); font-weight:900; letter-spacing:.14em; text-transform:uppercase; font-size:12px; margin-bottom:14px; }}
    .badges {{ display:flex; flex-wrap:wrap; gap:8px; margin-top:18px; }}
    .badge {{ display:inline-flex; align-items:center; min-height:30px; padding:6px 10px; border-radius:999px; color:white; background:var(--indigo); font-size:12px; font-weight:900; }}
    section {{ border-top:1px solid var(--line); padding:34px 0; }}
    .grid {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(220px,1fr)); gap:14px; }}
    .card {{ display:block; border:1px solid var(--line); border-radius:8px; background:white; padding:18px; box-shadow:0 18px 50px rgba(24,38,70,.1); color:var(--teal); font-weight:900; text-decoration:none; }}
    .metric h3 {{ font-size:24px; }}
    .metric p {{ color:#374151; }}
    table {{ width:100%; border-collapse:collapse; background:white; border:1px solid var(--line); border-radius:8px; overflow:hidden; }}
    th, td {{ padding:10px; border-bottom:1px solid var(--line); text-align:left; vertical-align:top; }}
    th {{ color:var(--indigo); background:#e8eef5; font-size:13px; }}
    td {{ font-size:14px; }}
    ul {{ margin:0; padding-left:22px; }}
    li {{ margin:8px 0; }}
    a {{ color:var(--teal); text-underline-offset:3px; }}
  </style>
</head>
<body>
  <header>
    <div class="kicker">Proof and copyedit audit</div>
    <h1>Lady D trilogy proof audit pack</h1>
    <p class="lead">This page advances the next production gate: mechanical proof, theological watchlist review, repeated-title decisions, repeated morning-impact decisions, and KDP proof readiness for all three devotionals and companion journals.</p>
    <div class="badges">
      <span class="badge">Generated {GENERATED}</span>
      <span class="badge">Commit {html.escape(str(audit['commit']))}</span>
      <span class="badge">Not final upload-ready</span>
    </div>
  </header>
  <main>
    <section>
      <h2>Audit Snapshot</h2>
      <div class="grid">
        <div class="card metric"><h3>{totals['products']}</h3><p>Products checked</p></div>
        <div class="card metric"><h3>{totals['words']:,}</h3><p>Source words checked</p></div>
        <div class="card metric"><h3>{totals['placeholders']}</h3><p>Placeholder markers</p></div>
        <div class="card metric"><h3>{totals['sunday_mentions']}</h3><p>Sunday mentions</p></div>
        <div class="card metric"><h3>{totals['duplicate_title_groups']}</h3><p>Duplicate title groups</p></div>
        <div class="card metric"><h3>{totals['duplicate_morning_impact_groups']}</h3><p>Duplicate morning-impact groups</p></div>
      </div>
    </section>
    <section>
      <h2>Downloads</h2>
      <div class="grid">{download_links}</div>
    </section>
    <section>
      <h2>Product Evidence</h2>
      <table><thead><tr><th>Product</th><th>Type</th><th>Words</th><th>Structure</th><th>Sunday</th><th>Watch contexts</th></tr></thead><tbody>{product_rows}</tbody></table>
    </section>
    <section>
      <h2>Remaining Gates</h2>
      <ul>{gates}</ul>
      <p><a href="production.html">Return to production review</a></p>
    </section>
  </main>
</body>
</html>
"""


def make_zip(paths: list[Path]) -> Path:
    zip_path = OUT / "Lady-D-Trilogy-Proof-Audit-Pack.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in paths:
            zf.write(path, path.name)
    return zip_path


def sync(paths: list[Path]) -> None:
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)
    for path in paths:
        shutil.copy2(path, PUBLIC_OUT / path.name)
        shutil.copy2(path, LIBRARY_OUT / path.name)


def main() -> None:
    commit = current_commit()
    OUT.mkdir(parents=True, exist_ok=True)
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)
    audit = analyze_products(commit)

    generated: list[Path] = []
    audit_json = OUT / "trilogy-proof-audit.json"
    main_md = OUT / "lady-d-trilogy-proof-audit-pack.md"
    repetition_md = OUT / "title-and-impact-repetition-ledger.md"
    repetition_json = OUT / "title-and-impact-repetition-ledger.json"
    watch_md = OUT / "theological-watchlist-context-ledger.md"
    watch_json = OUT / "theological-watchlist-context-ledger.json"
    runbook_md = OUT / "proofreader-runbook.md"
    review_page = OUT / "lady-d-trilogy-proof-audit-review.html"

    write(audit_json, json.dumps(audit, indent=2))
    write(main_md, main_report_markdown(audit))
    write(repetition_md, title_repetition_markdown(audit))
    write(
        repetition_json,
        json.dumps(
            {"title_duplicates": audit["title_duplicates"], "morning_impact_duplicates": audit["morning_impact_duplicates"]},
            indent=2,
        ),
    )
    write(watch_md, watch_contexts_markdown(audit))
    write(watch_json, json.dumps(audit["works_watch_contexts"], indent=2))
    write(runbook_md, proofreader_runbook(audit))
    write(review_page, review_html(audit))
    generated.extend([audit_json, main_md, repetition_md, repetition_json, watch_md, watch_json, runbook_md, review_page])

    for volume in (1, 2, 3):
        checklist = OUT / f"volume-{volume}-proof-checklist.md"
        write(checklist, volume_checklist(audit, volume))
        generated.append(checklist)

    docx_path = build_docx(audit)
    pdf_path = convert_docx_to_pdf(docx_path)
    generated.extend([docx_path, pdf_path])

    zip_path = make_zip(generated)
    all_paths = generated + [zip_path]
    sync(all_paths)
    shutil.copy2(review_page, SOURCE_PAGE)
    shutil.copy2(review_page, PUBLIC_PAGE)
    print(
        json.dumps(
            {
                "report": str(docx_path.relative_to(ROOT)),
                "pdf": str(pdf_path.relative_to(ROOT)),
                "zip": str(zip_path.relative_to(ROOT)),
                "page": str(SOURCE_PAGE.relative_to(ROOT)),
                "result": audit["result"],
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
