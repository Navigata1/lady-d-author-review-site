#!/usr/bin/env python3
"""Build the Lady D proof decision application pack."""

from __future__ import annotations

import html
import json
import shutil
import subprocess
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_release_upload_readiness_pack import AUTHOR, ROOT, current_commit


GENERATED = "2026-07-01"
OUT = ROOT / "downloads" / "production" / "kdp" / "proof-decision-application"
PUBLIC_OUT = ROOT / "public" / "downloads" / "production" / "kdp" / "proof-decision-application"
SOURCE_PAGE = ROOT / "proof-decision-application.html"
PUBLIC_PAGE = ROOT / "public" / "proof-decision-application.html"
LIBRARY_OUT = Path("/Users/IDC2.5/Documents/LADY D/Production Library/_Shared/KDP Readiness/Proof Decision Application")

BASELINE_DECISION_PATH = OUT / "proof-decision-prior-queue-snapshot.json"
SOURCE_APPLICATION_PATH = OUT / "proof-decision-source-application.json"
CURRENT_DECISION_PATH = ROOT / "downloads" / "production" / "kdp" / "proof-decision-resolution" / "proof-decision-resolution.json"
CURRENT_AUDIT_PATH = ROOT / "downloads" / "production" / "kdp" / "trilogy-proof-audit" / "trilogy-proof-audit.json"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(89, 96, 108)
GOLD = RGBColor(122, 90, 0)
LIGHT_FILL = "E8EEF5"
SOFT_FILL = "F4F6F9"


def write(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    cleaned = "\n".join(line.rstrip() for line in content.rstrip().splitlines())
    path.write_text(cleaned + "\n", encoding="utf-8")


def load_json(path: Path) -> dict[str, object]:
    return json.loads(path.read_text(encoding="utf-8"))


def prior_queue_snapshot() -> dict[str, object]:
    if BASELINE_DECISION_PATH.exists():
        return load_json(BASELINE_DECISION_PATH)
    raw = subprocess.check_output(
        ["git", "show", f"HEAD:{CURRENT_DECISION_PATH.relative_to(ROOT)}"],
        cwd=ROOT,
        text=True,
    )
    snapshot = json.loads(raw)
    write(BASELINE_DECISION_PATH, json.dumps(snapshot, indent=2))
    return snapshot


def current_volume_counts(audit: dict[str, object]) -> dict[str, dict[str, int]]:
    counts = {str(volume): {"title_duplicate_groups": 0, "morning_impact_duplicate_groups": 0, "priority_theology_contexts": 0} for volume in (1, 2, 3)}
    for row in audit["title_duplicates"]:
        counts[str(row["volume"])]["title_duplicate_groups"] += 1
    for row in audit["morning_impact_duplicates"]:
        counts[str(row["volume"])]["morning_impact_duplicate_groups"] += 1
    product_to_volume = {product["key"]: str(product["volume"]) for product in audit["products"]}
    for item in audit["works_watch_contexts"]:
        if item["preliminary_read"] == "requires proof review":
            counts[product_to_volume[item["product"]]]["priority_theology_contexts"] += 1
    return counts


def build_payload(commit: str) -> dict[str, object]:
    prior = prior_queue_snapshot()
    current = load_json(CURRENT_DECISION_PATH)
    audit = load_json(CURRENT_AUDIT_PATH)
    source_application = load_json(SOURCE_APPLICATION_PATH)
    return {
        "generated": GENERATED,
        "commit": commit,
        "author": AUTHOR,
        "status": "proof_decisions_applied_queue_clear_not_final_upload",
        "prior_queue_totals": prior["totals"],
        "current_queue_totals": current["totals"],
        "current_proof_audit_totals": audit["totals"],
        "prior_volume_breakdown": prior["volume_breakdown"],
        "current_volume_proof_counts": current_volume_counts(audit),
        "source_repetition_check": source_application.get("post_source_repetition_check", {}),
        "remaining_release_gates": [
            "Full author voice copyedit and devotional flow pass.",
            "Final front matter, dedication, acknowledgments, bio, ISBN, and barcode.",
            "Bible permissions decision if quotation text is added.",
            "Locked page counts and regenerated final full-wrap covers.",
            "KDP Previewer pass and physical proof review before public release.",
        ],
    }


def report_markdown(payload: dict[str, object]) -> str:
    prior = payload["prior_queue_totals"]
    current = payload["current_queue_totals"]
    audit = payload["current_proof_audit_totals"]
    rows = []
    for volume in ("1", "2", "3"):
        prior_row = payload["prior_volume_breakdown"][volume]
        current_row = payload["current_volume_proof_counts"][volume]
        rows.append(
            f"| Volume {volume} | {prior_row['title_decisions']} | {prior_row['morning_impact_decisions']} | {prior_row['theology_decisions']} | "
            f"{current_row['title_duplicate_groups']} | {current_row['morning_impact_duplicate_groups']} | {current_row['priority_theology_contexts']} |"
        )
    return f"""# Lady D Proof Decision Application Pack

Generated: {GENERATED}

Repo commit at generation: `{payload['commit']}`

Author: {AUTHOR}

Status: Proof decisions applied; queue clear; not final KDP upload-ready.

## Purpose

This pack records the movement from the prior proof decision queue into the current clean proof-audit state. It preserves the prior 192-item decision snapshot, the source-application ledger, and the regenerated proof audit that now reports no duplicate-title groups, no duplicate morning-impact groups, and no priority theology contexts.

## Prior Queue

- Prior total decision items: {prior['decision_items']}
- Prior duplicate-title decisions: {prior['title_decisions']}
- Prior morning-impact decisions: {prior['morning_impact_decisions']}
- Prior theology decisions: {prior['theology_decisions']}

## Current Queue And Audit

- Current decision items: {current['decision_items']}
- Current duplicate-title decisions: {current['title_decisions']}
- Current morning-impact decisions: {current['morning_impact_decisions']}
- Current theology decisions: {current['theology_decisions']}
- Current audit duplicate-title groups: {audit['duplicate_title_groups']}
- Current audit duplicate morning-impact groups: {audit['duplicate_morning_impact_groups']}
- Current audit priority theology contexts: {audit['review_required_contexts']}
- Current audit Sunday mentions: {audit['sunday_mentions']}
- Current audit placeholder markers: {audit['placeholders']}

## Volume Movement

| Volume | Prior title decisions | Prior impact decisions | Prior theology decisions | Current title groups | Current impact groups | Current theology contexts |
| --- | ---: | ---: | ---: | ---: | ---: | ---: |
{chr(10).join(rows)}

## Remaining Release Gates

{chr(10).join(f"- {item}" for item in payload['remaining_release_gates'])}
"""


def set_font(run, size: float | None = None, bold: bool = False, italic: bool = False, color=INK) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.header.paragraphs[0].text = "Lady D Proof Decision Application"
    section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(section.header.paragraphs[0].runs[0], size=8.5, color=MUTED, italic=True)
    section.footer.paragraphs[0].text = "Proof decisions applied - not final KDP upload approval"
    section.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(section.footer.paragraphs[0].runs[0], size=8, color=MUTED)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK_BLUE)]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_para(doc: Document, text: str, *, size=11, color=INK, bold=False, italic=False, align=None, after=6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), size=10.5, color=INK)


def add_movement_table(doc: Document, payload: dict[str, object]) -> None:
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Vol", "Prior Titles", "Prior Impacts", "Prior Theology", "Now Titles", "Now Impacts", "Now Theology"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font(run, size=8, color=DARK_BLUE, bold=True)
    for volume in ("1", "2", "3"):
        prior = payload["prior_volume_breakdown"][volume]
        current = payload["current_volume_proof_counts"][volume]
        values = [
            volume,
            prior["title_decisions"],
            prior["morning_impact_decisions"],
            prior["theology_decisions"],
            current["title_duplicate_groups"],
            current["morning_impact_duplicate_groups"],
            current["priority_theology_contexts"],
        ]
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
            set_cell_margins(cells[idx])
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=8.5, color=INK)


def build_docx(payload: dict[str, object]) -> Path:
    docx_path = OUT / "lady-d-proof-decision-application-pack.docx"
    doc = Document()
    setup_doc(doc)
    doc.core_properties.title = "Lady D Proof Decision Application Pack"
    doc.core_properties.author = "IDC Publishing"
    add_para(doc, "LADY D DEVOTIONAL LIBRARY", size=9, color=GOLD, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Proof Decision Application Pack", size=24, color=DARK_BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_para(doc, "Evidence that the title, impact, and priority theology proof queue is now clear", size=12, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_para(doc, f"Generated {GENERATED} from commit {payload['commit']}", size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    prior = payload["prior_queue_totals"]
    current = payload["current_queue_totals"]
    audit = payload["current_proof_audit_totals"]
    doc.add_heading("Movement Summary", level=1)
    for item in [
        f"Prior proof decision items: {prior['decision_items']}",
        f"Current proof decision items: {current['decision_items']}",
        f"Current duplicate title groups: {audit['duplicate_title_groups']}",
        f"Current duplicate morning-impact groups: {audit['duplicate_morning_impact_groups']}",
        f"Current priority theology contexts: {audit['review_required_contexts']}",
        f"Current Sunday mentions: {audit['sunday_mentions']}",
        f"Current placeholder markers: {audit['placeholders']}",
    ]:
        add_bullet(doc, item)
    doc.add_heading("Volume Movement", level=1)
    add_movement_table(doc, payload)
    doc.add_heading("Release Boundary", level=1)
    add_para(doc, "This pack proves the proof-decision queue is clear. It does not make the library final KDP upload-ready; final copyedit, author approvals, ISBN/barcode, locked covers, KDP Previewer, and physical proofs still remain.", after=8)
    doc.add_heading("Next Gates", level=1)
    for item in payload["remaining_release_gates"]:
        add_bullet(doc, item)
    doc.save(docx_path)
    return docx_path


def convert_docx_to_pdf(docx_path: Path) -> Path:
    pdf_path = docx_path.with_suffix(".pdf")
    if pdf_path.exists():
        pdf_path.unlink()
    subprocess.check_call(["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUT), str(docx_path)])
    return pdf_path


def review_html(payload: dict[str, object]) -> str:
    prior = payload["prior_queue_totals"]
    current = payload["current_queue_totals"]
    audit = payload["current_proof_audit_totals"]
    volume_rows = "\n".join(
        f"<tr><td>Volume {volume}</td><td>{payload['prior_volume_breakdown'][volume]['title_decisions']}</td><td>{payload['prior_volume_breakdown'][volume]['morning_impact_decisions']}</td><td>{payload['prior_volume_breakdown'][volume]['theology_decisions']}</td><td>{payload['current_volume_proof_counts'][volume]['title_duplicate_groups']}</td><td>{payload['current_volume_proof_counts'][volume]['morning_impact_duplicate_groups']}</td><td>{payload['current_volume_proof_counts'][volume]['priority_theology_contexts']}</td></tr>"
        for volume in ("1", "2", "3")
    )
    downloads = [
        ("Download application ZIP", "downloads/production/kdp/proof-decision-application/Lady-D-Proof-Decision-Application-Pack.zip"),
        ("Application PDF", "downloads/production/kdp/proof-decision-application/lady-d-proof-decision-application-pack.pdf"),
        ("Application DOCX", "downloads/production/kdp/proof-decision-application/lady-d-proof-decision-application-pack.docx"),
        ("Application report MD", "downloads/production/kdp/proof-decision-application/lady-d-proof-decision-application-pack.md"),
        ("Application JSON", "downloads/production/kdp/proof-decision-application/proof-decision-application.json"),
        ("Source application ledger", "downloads/production/kdp/proof-decision-application/proof-decision-source-application.json"),
        ("Prior queue snapshot", "downloads/production/kdp/proof-decision-application/proof-decision-prior-queue-snapshot.json"),
    ]
    links = "\n".join(f'<a class="card" href="{href}">{html.escape(label)}</a>' for label, href in downloads)
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='12' fill='%23182646'/%3E%3Ctext x='50%25' y='55%25' dominant-baseline='middle' text-anchor='middle' font-family='Georgia,serif' font-size='26' fill='%23fffdf8'%3ELD%3C/text%3E%3C/svg%3E">
  <title>Lady D Proof Decision Application</title>
  <style>
    :root {{ --ink:#111827; --paper:#fffdf8; --mist:#f5f2eb; --indigo:#182646; --teal:#1d716f; --gold:#c99335; --line:rgba(17,24,39,.14); }}
    * {{ box-sizing:border-box; }}
    body {{ margin:0; color:var(--ink); font-family:Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; background:linear-gradient(180deg,var(--paper),var(--mist)); line-height:1.5; }}
    header, main {{ max-width:1180px; margin:0 auto; padding:34px 22px; }}
    h1, h2, h3 {{ font-family:Georgia, "Times New Roman", serif; line-height:1.05; margin:0 0 14px; letter-spacing:0; overflow-wrap:anywhere; }}
    h1 {{ font-size:clamp(42px, 7vw, 82px); max-width:980px; }}
    h2 {{ font-size:clamp(28px, 4vw, 48px); }}
    p {{ margin:0 0 14px; }}
    .lead {{ font-size:clamp(18px, 2vw, 23px); max-width:880px; color:#2e3746; }}
    .kicker {{ color:var(--teal); font-weight:900; letter-spacing:.14em; text-transform:uppercase; font-size:12px; margin-bottom:14px; }}
    .badges {{ display:flex; flex-wrap:wrap; gap:8px; margin-top:18px; }}
    .badge {{ display:inline-flex; align-items:center; min-height:30px; padding:6px 10px; border-radius:999px; color:white; background:var(--indigo); font-size:12px; font-weight:900; }}
    section {{ border-top:1px solid var(--line); padding:34px 0; }}
    .grid {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(220px,1fr)); gap:14px; }}
    .card {{ display:block; border:1px solid var(--line); border-radius:8px; background:white; padding:18px; box-shadow:0 18px 50px rgba(24,38,70,.1); color:var(--teal); font-weight:900; text-decoration:none; }}
    .metric h3 {{ font-size:28px; color:var(--indigo); }}
    table {{ width:100%; border-collapse:collapse; background:white; border:1px solid var(--line); border-radius:8px; overflow:hidden; }}
    th, td {{ padding:10px; border-bottom:1px solid var(--line); text-align:left; vertical-align:top; }}
    th {{ color:var(--indigo); background:#e8eef5; font-size:13px; }}
    td {{ font-size:14px; }}
    a {{ color:var(--teal); text-underline-offset:3px; }}
    @media (max-width:700px) {{ table {{ display:block; overflow-x:auto; }} }}
  </style>
</head>
<body>
  <header>
    <div class="kicker">Proof decision application</div>
    <h1>Lady D proof decisions applied</h1>
    <p class="lead">The prior 192-item proof decision queue has been applied into the source manuscripts and journals, then regenerated through the master/proof lane. The current proof audit now reports no duplicate-title groups, no repeated morning-impact groups, and no priority theology contexts.</p>
    <div class="badges">
      <span class="badge">Generated {GENERATED}</span>
      <span class="badge">Prior queue {prior['decision_items']}</span>
      <span class="badge">Current queue {current['decision_items']}</span>
      <span class="badge">Audit title groups {audit['duplicate_title_groups']}</span>
      <span class="badge">Audit impact groups {audit['duplicate_morning_impact_groups']}</span>
    </div>
  </header>
  <main>
    <section>
      <h2>Movement Snapshot</h2>
      <div class="grid">
        <div class="card metric"><h3>{prior['decision_items']}</h3><p>Prior proof decisions</p></div>
        <div class="card metric"><h3>{current['decision_items']}</h3><p>Current proof decisions</p></div>
        <div class="card metric"><h3>{audit['duplicate_title_groups']}</h3><p>Duplicate title groups</p></div>
        <div class="card metric"><h3>{audit['duplicate_morning_impact_groups']}</h3><p>Duplicate morning-impact groups</p></div>
        <div class="card metric"><h3>{audit['review_required_contexts']}</h3><p>Priority theology contexts</p></div>
        <div class="card metric"><h3>{audit['sunday_mentions']}</h3><p>Sunday mentions</p></div>
      </div>
    </section>
    <section>
      <h2>Downloads</h2>
      <div class="grid">{links}</div>
    </section>
    <section>
      <h2>Volume Movement</h2>
      <table><thead><tr><th>Volume</th><th>Prior title</th><th>Prior impact</th><th>Prior theology</th><th>Now title</th><th>Now impact</th><th>Now theology</th></tr></thead><tbody>{volume_rows}</tbody></table>
    </section>
    <section>
      <h2>Release Boundary</h2>
      <p class="lead">This proves the proof-decision queue is clear. It does not make the library final KDP upload-ready; final copyedit, author approvals, ISBN/barcode, locked covers, KDP Previewer, and physical proofs still remain.</p>
      <p><a href="production.html">Return to production review</a></p>
    </section>
  </main>
</body>
</html>
"""


def make_zip(paths: list[Path]) -> Path:
    zip_path = OUT / "Lady-D-Proof-Decision-Application-Pack.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in paths:
            zf.write(path, path.name)
    return zip_path


def sync(paths: list[Path]) -> None:
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)
    for path in paths:
        shutil.copy2(path, PUBLIC_OUT / path.name)
        shutil.copy2(path, LIBRARY_OUT / path.name)


def main() -> None:
    commit = current_commit()
    OUT.mkdir(parents=True, exist_ok=True)
    PUBLIC_OUT.mkdir(parents=True, exist_ok=True)
    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)
    payload = build_payload(commit)

    generated: list[Path] = []
    payload_json = OUT / "proof-decision-application.json"
    report_md = OUT / "lady-d-proof-decision-application-pack.md"
    review_page = OUT / "lady-d-proof-decision-application-review.html"
    write(payload_json, json.dumps(payload, indent=2))
    write(report_md, report_markdown(payload))
    write(review_page, review_html(payload))
    generated.extend([payload_json, report_md, review_page, BASELINE_DECISION_PATH, SOURCE_APPLICATION_PATH])

    docx_path = build_docx(payload)
    pdf_path = convert_docx_to_pdf(docx_path)
    generated.extend([docx_path, pdf_path])
    zip_path = make_zip(generated)
    sync(generated + [zip_path])
    shutil.copy2(review_page, SOURCE_PAGE)
    shutil.copy2(review_page, PUBLIC_PAGE)
    print(
        json.dumps(
            {
                "page": str(SOURCE_PAGE.relative_to(ROOT)),
                "zip": str(zip_path.relative_to(ROOT)),
                "pdf": str(pdf_path.relative_to(ROOT)),
                "status": payload["status"],
                "prior_decisions": payload["prior_queue_totals"]["decision_items"],
                "current_decisions": payload["current_queue_totals"]["decision_items"],
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
